"""Heading-related autofix helpers: formatting existing headings and promoting candidates."""
from __future__ import annotations

import logging
import re

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from app.rules_engine.autofix_helpers import enforce_run_font, is_field_code_run
from app.rules_engine.autofix_para_classify import is_heading_para
from app.rules_engine.style_resolve import detect_toc_paragraph_indices

logger = logging.getLogger(__name__)

_HEADING_LEVEL_RE = re.compile(r"\d+")


def _detect_heading_level(paragraph) -> int | None:
    """Detect heading level from Word style name (e.g. 'Heading 2' -> 2)."""
    name = getattr(paragraph.style, "name", "") or ""
    m = _HEADING_LEVEL_RE.search(name)
    return int(m.group()) if m else None


def _resolve_heading_style(paragraph, level: int):
    """Get or create the Heading N style in the document.

    Newly created styles intentionally use ``space_before=0`` /
    ``space_after=0``: paragraph spacing around chapter headings is
    governed by the per-paragraph values that ``fix_heading`` /
    ``promote_to_heading`` set explicitly. Hard-coding 12 pt / 6 pt here
    used to leak through promoted paragraphs whose direct paragraph
    properties contained nothing — the inherited style values then
    rendered as the «extra spacing above and below chapter titles» that
    reviewers keep flagging.
    """
    doc = paragraph.part.document
    target_name = f"Heading {level}"
    try:
        return doc.styles[target_name]
    except KeyError:
        pass
    style = doc.styles.add_style(target_name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles["Normal"]
    style.quick_style = True
    pf = style.paragraph_format
    pf.keep_with_next = True
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    return style


def _fix_alignment(paragraph, level: int | None, cfg, details: list[str], idx: int) -> bool:
    if level == 1 and cfg.heading_level1_center:
        if paragraph.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            details.append(f"Заголовок #{idx + 1}: выравнивание по центру")
            return True
    elif level is not None and level >= 2:
        if cfg.heading_level2plus_center:
            if paragraph.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph.paragraph_format.first_line_indent = Mm(0)
                details.append(
                    f"Заголовок #{idx + 1}: подзаголовок по центру"
                )
                return True
        elif paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            details.append(f"Заголовок #{idx + 1}: выравнивание по левому краю")
            return True
    return False


def fix_heading(paragraph, idx: int, cfg, details: list[str]) -> bool:
    changed = False
    for run in paragraph.runs:
        if is_field_code_run(run):
            continue
        if enforce_run_font(run, cfg.heading_font, cfg.heading_size_pt):
            changed = True
        if cfg.heading_bold and not run.bold:
            run.bold = True
            changed = True
        if run.font.underline:
            run.font.underline = False
            changed = True

    level = _detect_heading_level(paragraph)
    if _fix_alignment(paragraph, level, cfg, details, idx):
        changed = True

    pf = paragraph.paragraph_format
    target_ls = getattr(cfg, "line_spacing", 1.5)
    cur_ls = pf.line_spacing
    if cur_ls is None or abs(float(cur_ls) - target_ls) > 0.05:
        pf.line_spacing = target_ls
        changed = True

    from docx.shared import Pt as _Pt
    target_sb = getattr(cfg, "space_before_pt", 0)
    cur_sb = pf.space_before
    if cur_sb is not None and abs(cur_sb.pt - target_sb) > 0.2:
        pf.space_before = _Pt(target_sb)
        changed = True
    target_sa = getattr(cfg, "space_after_pt", 0)
    cur_sa = pf.space_after
    if cur_sa is not None and abs(cur_sa.pt - target_sa) > 0.2:
        pf.space_after = _Pt(target_sa)
        changed = True

    if changed:
        details.append(
            f"Заголовок #{idx + 1}: {cfg.heading_font}, {cfg.heading_size_pt} пт, полужирный"
        )
    return changed


def promote_to_heading(
    paragraph, level: int, idx: int, cfg, details: list[str],
) -> bool:
    """Assign Word Heading style to a paragraph detected as heading candidate by text."""
    try:
        style = _resolve_heading_style(paragraph, level)
        paragraph.style = style
    except Exception:
        logger.debug("Promote: cannot assign Heading %d", level, exc_info=True)
        return False

    for run in paragraph.runs:
        if is_field_code_run(run):
            continue
        enforce_run_font(run, cfg.heading_font, cfg.heading_size_pt)
        if cfg.heading_bold:
            run.bold = True
    pf = paragraph.paragraph_format
    pf.first_line_indent = Mm(0)

    target_sb = float(getattr(cfg, "space_before_pt", 0) or 0)
    target_sa = float(getattr(cfg, "space_after_pt", 0) or 0)
    pf.space_before = Pt(target_sb)
    pf.space_after = Pt(target_sa)

    target_ls = float(getattr(cfg, "line_spacing", 1.5) or 1.5)
    cur_ls = pf.line_spacing
    if cur_ls is None or abs(float(cur_ls) - target_ls) > 0.05:
        pf.line_spacing = target_ls

    if level == 1 and cfg.heading_level1_center:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif level >= 2:
        if cfg.heading_level2plus_center:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    details.append(f"Абзац #{idx + 1} → «Заголовок {level}»: {paragraph.text[:50]}")
    return True


def _para_heading_level(paragraph) -> int | None:
    """Return Word heading level (1..9) or None if paragraph is not a heading."""
    sid = getattr(paragraph.style, "style_id", "") or ""
    for i in range(1, 10):
        if sid == f"Heading{i}":
            return i
    sname = (getattr(paragraph.style, "name", "") or "").lower()
    for i in range(1, 10):
        if f"heading {i}" in sname or f"заголовок {i}" in sname:
            return i
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is not None:
        ol = pPr.find(qn("w:outlineLvl"))
        if ol is not None:
            try:
                val = int(ol.get(qn("w:val"), "9"))
                if val < 9:
                    return val + 1
            except (TypeError, ValueError):
                pass
    return None


def _para_has_page_break_before(paragraph) -> bool:
    if paragraph.paragraph_format.page_break_before:
        return True
    for br in paragraph._element.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def _build_empty_paragraph() -> "OxmlElement":
    """Build an empty <w:p/> with no runs — one blank line."""
    from docx.oxml import OxmlElement

    return OxmlElement("w:p")


def enforce_subheading_alignment(doc, cfg, details: list[str]) -> bool:
    """Force every heading to a specific alignment.

    * level 1 (chapters / ВВЕДЕНИЕ / ЗАКЛЮЧЕНИЕ / СПИСОК ЛИТЕРАТУРЫ) →
      center when ``heading_level1_center`` is True.
    * level 2+ (subsections 1.1, 1.2…) → center when
      ``heading_level2plus_center`` is True.

    Runs independently of ``skip_headings`` safety flag because changing only
    alignment (and wiping red-line indent) is safe and matches the client's
    explicit request «название параграфа / главы по середине».
    """
    center1 = bool(getattr(cfg, "heading_level1_center", True))
    center2plus = bool(getattr(cfg, "heading_level2plus_center", True))
    if not center1 and not center2plus:
        return False

    toc_elems = _collect_toc_paragraph_elements(doc)
    chapters_changed = 0
    subs_changed = 0
    for para in doc.paragraphs:
        if para._element in toc_elems:
            continue
        level = _para_heading_level(para)
        if level is None:
            continue
        if level == 1 and not center1:
            continue
        if level >= 2 and not center2plus:
            continue

        touched = False
        if para.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            touched = True
        pf = para.paragraph_format
        if pf.first_line_indent is not None and int(pf.first_line_indent) != 0:
            pf.first_line_indent = Mm(0)
            touched = True
        if pf.left_indent is not None and int(pf.left_indent) != 0:
            pf.left_indent = Mm(0)
            touched = True
        if touched:
            if level == 1:
                chapters_changed += 1
            else:
                subs_changed += 1

    if chapters_changed:
        details.append(
            f"Главы: {chapters_changed} шт. выровнены по центру"
        )
    if subs_changed:
        details.append(
            f"Подзаголовки: {subs_changed} шт. выровнены по центру"
        )
    return (chapters_changed + subs_changed) > 0


def ensure_blank_before_subheadings(doc, details: list[str]) -> bool:
    """Ensure exactly one empty paragraph precedes every level-2+ heading
    within a chapter.

    Chapters (heading level 1) usually start on a new page via
    ``page_break_before`` and are skipped here. For sub-headings ``1.1``,
    ``1.2`` etc. we guarantee a single blank paragraph above, matching the
    client's request «один пробел между параграфами внутри главы».
    """
    paragraphs = list(doc.paragraphs)
    if len(paragraphs) < 2:
        return False

    toc_elems = _collect_toc_paragraph_elements(doc)
    changed = False
    inserted = 0

    for para in paragraphs:
        if para._element in toc_elems:
            continue
        level = _para_heading_level(para)
        if level is None or level < 2:
            continue
        if _para_has_page_break_before(para):
            continue

        prev = para._element.getprevious()
        while prev is not None and prev.tag != qn("w:p"):
            prev = prev.getprevious()
        if prev is None:
            continue

        prev_text = "".join(
            (t.text or "") for t in prev.iter(qn("w:t"))
        ).strip()
        has_page_break = any(
            br.get(qn("w:type")) == "page" for br in prev.iter(qn("w:br"))
        )
        if has_page_break:
            continue

        if prev_text == "":
            continue

        blank = _build_empty_paragraph()
        para._element.addprevious(blank)
        inserted += 1
        changed = True

    if inserted:
        details.append(
            f"Отступы: добавлено {inserted} пустых абзаца(ев) перед подзаголовками"
        )
    return changed


_CHAPTER_PAGE_BREAK_RE = re.compile(
    r"^\s*(?:"
    r"глава\s+\d+"
    r"|\d+\s+глава"
    r"|введение"
    r"|заключение"
    r"|содержание"
    r"|оглавление"
    r"|список\s+(?:использованн|используем|литератур|источник)"
    r"|библиографическ(?:ий|ая)\s+список"
    r"|библиография"
    r"|приложение(?:\s+\S+)?"
    r"|аннотация"
    r"|реферат"
    r"|annotation"
    r"|abstract"
    r"|references"
    r")\b",
    re.IGNORECASE | re.UNICODE,
)


def _collect_toc_paragraph_elements(doc) -> set:
    """Return every ``<w:p>`` element that lives inside a TOC field.

    Covers both shapes a TOC field can take:
      * an inline ``<w:fldChar>``-based field that spans several body
        paragraphs between ``begin`` and ``end``;
      * an auto-TOC wrapped in ``<w:sdt>/<w:sdtContent>`` content control.

    Also includes ``<w:p>`` elements nested inside any ``<w:sdt>`` whose
    descendant ``<w:instrText>`` references ``TOC``.
    """
    result: set = set()
    body = doc.element.body
    fld_tag = qn("w:fldChar")
    instr_tag = qn("w:instrText")
    p_tag = qn("w:p")
    sdt_tag = qn("w:sdt")
    sdt_content_tag = qn("w:sdtContent")
    fld_type_attr = qn("w:fldCharType")

    field_stack: list[str] = []
    in_toc = False
    p_elements = list(body.iterchildren(p_tag))
    for p_elem in p_elements:
        was_in_toc = in_toc
        for elem in p_elem.iter():
            tag = elem.tag
            if tag == fld_tag:
                ftype = elem.get(fld_type_attr)
                if ftype == "begin":
                    field_stack.append("unknown")
                elif ftype == "separate":
                    if field_stack and field_stack[-1] == "unknown":
                        field_stack[-1] = "other"
                    if field_stack and field_stack[-1] == "TOC":
                        in_toc = True
                elif ftype == "end" and field_stack:
                    ended = field_stack.pop()
                    if ended == "TOC":
                        in_toc = False
            elif tag == instr_tag:
                text = elem.text or ""
                if field_stack and field_stack[-1] == "unknown" and "TOC" in text.upper():
                    field_stack[-1] = "TOC"
                    in_toc = True
        if in_toc or was_in_toc:
            result.add(p_elem)

    for sdt in body.iter(sdt_tag):
        instr_texts = "".join((el.text or "") for el in sdt.iter(instr_tag))
        if "TOC" not in instr_texts.upper():
            continue
        content = sdt.find(sdt_content_tag)
        if content is None:
            continue
        # Skip the very first paragraph of the SDT TOC — it's the
        # «Содержание»/«Оглавление» heading itself. Adding it to
        # ``toc_elements`` would prevent ``enforce_chapter_page_breaks``
        # from inserting the page-break before the auto-TOC, which is
        # exactly the «contents glued to the title page» complaint.
        sub_paras = list(content.iter(p_tag))
        first_skipped = False
        for sub in sub_paras:
            if not first_skipped:
                text_norm = "".join(
                    (t.text or "") for t in sub.iter(qn("w:t"))
                ).strip().lower().rstrip(":.;")
                if text_norm in ("содержание", "оглавление"):
                    first_skipped = True
                    continue
                first_skipped = True
            result.add(sub)
    return result


def _first_body_heading_paragraph_index(doc) -> int | None:
    """Index in ``doc.paragraphs`` of the first paragraph that opens the body."""
    for i, p in enumerate(doc.paragraphs):
        if is_heading_para(p):
            return i
    return None


def _paragraph_index_in_doc_paragraphs(doc, p_elem) -> int | None:
    for i, p in enumerate(doc.paragraphs):
        if p._element is p_elem:
            return i
    return None


def _iter_chapter_break_paragraphs(doc):
    """Yield every Paragraph eligible for chapter page-break enforcement.

    Returns body-level ``<w:p>`` elements plus the *first* paragraph of
    every ``<w:sdt>/<w:sdtContent>`` content control. The latter covers
    the auto-TOC case where «Содержание» is wrapped in an SDT and would
    otherwise be invisible to ``doc.paragraphs`` — only the first
    paragraph is yielded so subsequent TOC entries (which match the
    chapter regex too) do not get spurious page breaks inserted between
    them, which would scatter the table of contents across pages.

    Paragraphs nested inside tables are intentionally excluded — Word
    silently ignores ``pageBreakBefore`` for table-cell paragraphs.
    """
    from docx.text.paragraph import Paragraph

    body = doc.element.body
    p_tag = qn("w:p")
    sdt_tag = qn("w:sdt")
    sdt_content_tag = qn("w:sdtContent")

    for child in body.iterchildren():
        if child.tag == p_tag:
            yield Paragraph(child, doc)
        elif child.tag == sdt_tag:
            content = child.find(sdt_content_tag)
            if content is None:
                continue
            for sub in content.iterchildren():
                if sub.tag == p_tag:
                    yield Paragraph(sub, doc)
                    break


def enforce_chapter_page_breaks(doc, details: list[str]) -> bool:
    """Force ``page_break_before`` on every top-level heading (Heading 1) and
    on any paragraph whose text matches a well-known chapter/section label.

    This runs independently of the generic per-paragraph section-break fix
    which is gated by a strict length limit (``len(text) <= 100``) and
    therefore missed long chapter titles like
    «Глава 1 Теоретические основы организации работы государственного…».

    The iteration also descends into ``<w:sdt>`` content controls so the
    «Содержание» heading wrapped in an auto-TOC field is treated like any
    other top-level chapter title.
    """
    paragraphs = list(_iter_chapter_break_paragraphs(doc))
    try:
        toc_indices = detect_toc_paragraph_indices(doc)
    except Exception:
        toc_indices = set()

    # Build a set of <w:p> elements inside any TOC field (both inline
    # fldChar-based and SDT-wrapped) so we can skip them regardless of the
    # body index, which changes as the document is mutated elsewhere.
    toc_elements = _collect_toc_paragraph_elements(doc)

    # Map body-paragraph index (as seen by doc.paragraphs) to element, so we
    # can translate detect_toc_paragraph_indices into element checks.
    # IMPORTANT: the «Содержание»/«Оглавление» heading itself is always
    # part of these index sets (both ``_detect_manual_toc_indices`` and
    # ``detect_manual_toc_entry_indices`` start their range with the
    # heading paragraph). We deliberately do NOT add it to ``toc_elements``
    # — otherwise the heading would never receive ``page_break_before`` and
    # the table of contents would stay glued to the title page, which is
    # exactly the bug users keep reporting.
    body_paragraphs = doc.paragraphs
    for i in list(toc_indices):
        if 0 <= i < len(body_paragraphs):
            para = body_paragraphs[i]
            text_norm = re.sub(
                r"\s+", " ", (para.text or "")
            ).strip().lower().rstrip(":.;")
            if text_norm in ("содержание", "оглавление"):
                continue
            toc_elements.add(para._element)

    title_cutoff = _first_body_heading_paragraph_index(doc)

    changed = 0
    for idx, para in enumerate(paragraphs):
        if idx == 0:
            continue
        if para._element in toc_elements:
            continue
        if title_cutoff is not None:
            doc_idx = _paragraph_index_in_doc_paragraphs(doc, para._element)
            if doc_idx is not None and doc_idx < title_cutoff:
                continue
        text = (para.text or "").strip()
        if not text:
            continue

        level = _para_heading_level(para)
        matches_chapter = bool(_CHAPTER_PAGE_BREAK_RE.match(text))
        if level != 1 and not matches_chapter:
            continue
        if len(text) > 200:
            continue

        if _para_has_page_break_before(para):
            continue

        para.paragraph_format.page_break_before = True
        changed += 1

    if changed:
        details.append(
            f"Разрывы страниц: {changed} заголовок(ов) вынесено на новую страницу"
        )
    return changed > 0


# Canonical OOXML order of ``CT_RPr`` children (ECMA-376). Word
# silently ignores properties that appear in the wrong order, so
# whenever we touch a heading run's ``<w:rPr>`` we re-sort children
# to match this list.
_CT_RPR_CANONICAL_ORDER = (
    "rStyle",
    "rFonts",
    "b",
    "bCs",
    "i",
    "iCs",
    "caps",
    "smallCaps",
    "strike",
    "dstrike",
    "outline",
    "shadow",
    "emboss",
    "imprint",
    "noProof",
    "snapToGrid",
    "vanish",
    "webHidden",
    "color",
    "spacing",
    "w",
    "kern",
    "position",
    "sz",
    "szCs",
    "highlight",
    "u",
    "effect",
    "bdr",
    "shd",
    "fitText",
    "vertAlign",
    "rtl",
    "cs",
    "em",
    "lang",
    "eastAsianLayout",
    "specVanish",
    "oMath",
)
_CT_RPR_ORDER_INDEX = {name: i for i, name in enumerate(_CT_RPR_CANONICAL_ORDER)}


def _normalize_rpr_order(rPr) -> bool:
    """Re-sort ``<w:rPr>`` children into canonical OOXML order.

    Returns ``True`` when at least one element was re-positioned.
    """
    children = list(rPr)
    if not children:
        return False
    children_with_keys = []
    seen_unknown = False
    for child in children:
        local = child.tag.split("}")[-1]
        if local in _CT_RPR_ORDER_INDEX:
            children_with_keys.append((_CT_RPR_ORDER_INDEX[local], child))
        else:
            seen_unknown = True
            children_with_keys.append((len(_CT_RPR_ORDER_INDEX), child))
    sorted_pairs = sorted(
        enumerate(children_with_keys), key=lambda p: (p[1][0], p[0])
    )
    new_order = [pair[1][1] for pair in sorted_pairs]
    if [c for c in children] == new_order and not seen_unknown:
        return False
    if [c for c in children] == new_order:
        return False
    for child in children:
        rPr.remove(child)
    for child in new_order:
        rPr.append(child)
    return True


def _ensure_heading_styles_bold(doc) -> tuple[int, dict[int, str]]:
    """Add ``<w:b/>``/``<w:bCs/>`` to ``Heading 1/2/3`` (and "Заголовок 1/2/3")
    style ``<w:rPr>``s so headings keep displaying bold even after we
    strip run-level ``<w:b/>`` from heading paragraphs (which is needed
    so Word's TOC refresh doesn't carry that bold into TOC entries).

    Returns a tuple ``(style_changed, level_to_styleId)`` where
    ``level_to_styleId`` maps each detected heading level (1..3) to the
    *real* ``w:styleId`` of the style that carries that level. The
    map is needed because document templates routinely ship a style
    named "heading 1" whose ``styleId`` is something opaque like
    ``"1"`` or ``"Заголовок1"`` — setting ``<w:pStyle w:val="Heading1"/>``
    would point at a non-existent style and Word would silently render
    the paragraph in the default font.
    """
    from docx.oxml import OxmlElement
    styles_root = doc.styles.element
    canonical_names = {f"heading {i}": i for i in range(1, 4)}
    canonical_names.update({f"заголовок {i}": i for i in range(1, 4)})
    canonical_ids = {f"Heading{i}": i for i in range(1, 4)}

    def _force_bold_rpr(style_el) -> bool:
        """Ensure ``<w:b/>``/``<w:bCs/>`` are present (and non-zero) inside
        the style's ``<w:rPr>``."""
        rPr = style_el.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            pPr = style_el.find(qn("w:pPr"))
            if pPr is not None:
                pPr.addnext(rPr)
            else:
                style_el.append(rPr)
        local = False
        for tag in ("w:b", "w:bCs"):
            el = rPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rPr.append(el)
                local = True
            elif el.get(qn("w:val")) in ("0", "false"):
                el.set(qn("w:val"), "1")
                local = True
        return local

    by_id: dict[str, "OxmlElement"] = {}
    for st in styles_root.iter(qn("w:style")):
        sid = st.get(qn("w:styleId"))
        if sid:
            by_id[sid] = st

    level_to_styleid: dict[int, str] = {}
    style_changed = 0
    for st in styles_root.iter(qn("w:style")):
        sid = st.get(qn("w:styleId")) or ""
        sname_el = st.find(qn("w:name"))
        sname_val = (sname_el.get(qn("w:val")) if sname_el is not None else "") or ""
        sname_lower = sname_val.lower()
        level = canonical_ids.get(sid) or canonical_names.get(sname_lower)
        if level is None:
            continue
        # Prefer the first style we see for each level so deterministic
        # behaviour when the document has duplicates.
        level_to_styleid.setdefault(level, sid)
        if _force_bold_rpr(st):
            style_changed += 1
        # Word's "linked" character style (``<w:link w:val="…"/>``) is
        # automatically mirrored from the paragraph style's run
        # properties. Some renderers (notably the one Word uses when
        # exporting to PDF) prefer the linked character style's rPr
        # over the paragraph style's rPr for runs, which makes our
        # paragraph-style ``<w:b/>`` invisible. Update the linked
        # character style as well so the bold survives the round-trip.
        link_el = st.find(qn("w:link"))
        if link_el is not None:
            link_target = link_el.get(qn("w:val")) or ""
            link_st = by_id.get(link_target)
            if link_st is not None and _force_bold_rpr(link_st):
                style_changed += 1
    return style_changed, level_to_styleid


def enforce_heading_bold(doc, cfg, details: list[str]) -> bool:
    """Ensure every Heading 1/2/3 paragraph DISPLAYS bold while keeping
    auto-TOC entries non-bold.

    Strategy:
        * Add ``<w:b/>``/``<w:bCs/>`` to the Heading 1/2/3 (Заголовок 1/2/3)
          style ``<w:rPr>`` so the bold flag travels via style inheritance.
        * Strip run-level ``<w:b/>``/``<w:bCs/>`` from every heading run
          and from the paragraph's default ``<w:rPr>``. Run-level bold
          is what Word's TOC field-refresh copies into the regenerated
          TOC entries; once it's gone the TOC1/TOC2/TOC3 styles' explicit
          ``<w:b w:val="0"/>`` takes effect and the TOC renders non-bold,
          which is what the customer asked for.

    The previous behaviour (force run-level ``<w:b w:val="1"/>``) made
    the body look correct but caused «Содержание» entries to render
    bold after Word silently refreshed the TOC field on PDF export.
    """
    if not getattr(cfg, "heading_bold", True):
        return False

    from docx.oxml import OxmlElement
    style_changed, level_to_styleid = _ensure_heading_styles_bold(doc)

    toc_elems = _collect_toc_paragraph_elements(doc)
    paras_changed = 0
    pstyles_changed = 0
    bold_runs_added = 0
    canonical_names_lower = (
        {f"heading {i}" for i in range(1, 4)}
        | {f"заголовок {i}" for i in range(1, 4)}
    )
    canonical_ids_set = {f"Heading{i}" for i in range(1, 4)}
    canonical_target_ids = set(level_to_styleid.values())
    for para in doc.paragraphs:
        if para._element in toc_elems:
            continue
        level = _para_heading_level(para)
        if level is None or level > 3:
            continue
        sid = (getattr(para.style, "style_id", "") or "")
        sname = (getattr(para.style, "name", "") or "").lower()
        # "Canonical" means the paragraph is on a style that we
        # marked bold via ``_ensure_heading_styles_bold``: the
        # official ``HeadingN`` styleIds, anything whose name matches
        # ``heading N`` / ``заголовок N``, or any styleId we already
        # noted as the level's primary style.
        is_canonical = (
            sid in canonical_ids_set
            or sname in canonical_names_lower
            or sid in canonical_target_ids
        )
        touched = False
        if is_canonical:
            # Body heading is on a canonical style (e.g. avto's
            # ``Heading1`` / voina's ``Heading 1``). The style's
            # ``<w:b/>`` (added by ``_ensure_heading_styles_bold``)
            # makes the paragraph render bold via inheritance —
            # strip explicit run-level ``<w:b/>`` so Word's TOC
            # refresh doesn't carry that bold into regenerated
            # entries.
            pPr = para._element.find(qn("w:pPr"))
            if pPr is not None:
                rPr_def = pPr.find(qn("w:rPr"))
                if rPr_def is not None:
                    for tag in ("w:b", "w:bCs"):
                        el = rPr_def.find(qn(tag))
                        if el is not None:
                            rPr_def.remove(el)
                            touched = True
            for r_elem in para._element.iter(qn("w:r")):
                rPr = r_elem.find(qn("w:rPr"))
                if rPr is None:
                    continue
                for tag in ("w:b", "w:bCs"):
                    el = rPr.find(qn(tag))
                    if el is not None:
                        rPr.remove(el)
                        touched = True
        else:
            # Non-canonical heading paragraph (custom style like
            # ref6's ``af0`` / "основной" + explicit ``<w:outlineLvl/>``).
            # The original document encodes bold at the run level
            # — leave the paragraph alone here. Switching ``pStyle``
            # to a different style ID empirically breaks bold
            # rendering in some Word templates because ``<w:rFonts>``
            # / ``<w:sz>`` overrides we add later trigger Word to
            # use the linked character style's rPr instead of the
            # paragraph style's rPr, and the linked char style is
            # rebuilt without bold. The ``run.bold = True`` guard
            # below makes sure heading runs that lost their bold to
            # a sibling autofix still display bold.
            for run in para.runs:
                if is_field_code_run(run):
                    continue
                if not (run.text or "").strip():
                    continue
                if run.bold is True:
                    continue
                run.bold = True
                bold_runs_added += 1
                touched = True

        if touched:
            paras_changed += 1

    if style_changed or paras_changed or pstyles_changed or bold_runs_added:
        details.append(
            f"Заголовки: жирность нормализована "
            f"(стилей: {style_changed}, абзацев: {paras_changed}, "
            f"переведено на «Heading N»: {pstyles_changed}, "
            f"добавлено run-level bold: {bold_runs_added}) — "
            f"оглавление при этом остаётся не жирным"
        )
    return bool(style_changed or paras_changed or pstyles_changed or bold_runs_added)


def enforce_heading_font(doc, cfg, details: list[str]) -> bool:
    """Жёстко проставить шрифт и размер всем runs заголовков.

    Аналог ``enforce_heading_bold``. Работает независимо от safety-флага
    ``skip_headings``, потому что унаследованные стили заголовков в
    pandoc-выгрузках обычно используют majorHAnsi (Cambria) и крупные
    кегли (16/14/12pt с цветом темы), а клиент требует Times New Roman
    14pt. Заодно гарантирует, что у CJK-символов внутри заголовка
    проставлен ``w:eastAsia`` — без этого они продолжали бы рендериться
    Calibri через тему.
    """
    if not getattr(cfg, "normalize_font", True):
        return False
    target_font = getattr(cfg, "heading_font", "") or ""
    target_size = float(getattr(cfg, "heading_size_pt", 0) or 0)
    if not target_font and target_size <= 0:
        return False

    toc_elems = _collect_toc_paragraph_elements(doc)
    changed = 0
    for para in doc.paragraphs:
        if para._element in toc_elems:
            continue
        level = _para_heading_level(para)
        if level is None:
            continue
        touched = False
        for run in para.runs:
            if is_field_code_run(run):
                continue
            if enforce_run_font(run, target_font, target_size or None):
                touched = True
        if touched:
            changed += 1
    if changed:
        details.append(
            f"Заголовки: шрифт/размер выровнены ({target_font} {target_size:g} пт) "
            f"в {changed} заголовке(ах)"
        )
    return changed > 0


def enforce_heading_spacing(doc, cfg, details: list[str]) -> bool:
    """Hard-set ``space_before`` / ``space_after`` to ``cfg`` values on every
    ``HeadingN`` paragraph **and** style.

    Why:
        Pandoc/Word templates routinely ship Heading 1 with
        ``<w:spacing w:before="480"/>`` (= 24 pt) and Heading 2..9 with
        ``200`` (= 10 pt). Heading paragraphs in the source document
        usually have *no* explicit paragraph-format spacing, so the value
        cascades from the style — and the existing ``fix_heading`` pass
        skipped the override (``cur_sb is not None`` guard) because the
        per-paragraph ``space_before`` reads as ``None``. The result: a
        big visible gap above each chapter title that reviewers keep
        flagging as «лишний отступ перед заголовком».

    The fix is intentionally split into two passes:
        1. **Style-level**: walk ``doc.styles``, find every ``HeadingN``,
           force ``space_before`` / ``space_after`` to the configured
           values. This is enough on its own for any heading whose
           paragraph-format does not override the style.
        2. **Paragraph-level**: walk every heading paragraph and write
           the values explicitly via ``<w:spacing w:before="0"
           w:after="0"/>``. Required for headings that DID inherit
           explicit non-zero spacing somewhere along the chain
           (manual override, theme, custom linked style).

    Runs unconditionally — independent of ``safety.skip_headings`` —
    because the operation only adjusts paragraph spacing and never
    touches font / alignment / bold.
    """
    target_sb = float(getattr(cfg, "space_before_pt", 0) or 0)
    target_sa = float(getattr(cfg, "space_after_pt", 0) or 0)

    styles_changed = 0
    for style in doc.styles:
        sid = getattr(style, "style_id", "") or ""
        if not sid.startswith("Heading"):
            continue
        try:
            spf = style.paragraph_format
        except AttributeError:
            continue
        touched = False
        cur_sb = spf.space_before
        if cur_sb is None or abs(cur_sb.pt - target_sb) > 0.2:
            spf.space_before = Pt(target_sb)
            touched = True
        cur_sa = spf.space_after
        if cur_sa is None or abs(cur_sa.pt - target_sa) > 0.2:
            spf.space_after = Pt(target_sa)
            touched = True
        if touched:
            styles_changed += 1

    paragraphs_changed = 0
    toc_elems = _collect_toc_paragraph_elements(doc)
    for para in doc.paragraphs:
        if para._element in toc_elems:
            continue
        if _para_heading_level(para) is None:
            continue
        pf = para.paragraph_format
        touched = False
        cur_sb = pf.space_before
        if cur_sb is None or abs(cur_sb.pt - target_sb) > 0.2:
            pf.space_before = Pt(target_sb)
            touched = True
        cur_sa = pf.space_after
        if cur_sa is None or abs(cur_sa.pt - target_sa) > 0.2:
            pf.space_after = Pt(target_sa)
            touched = True
        if touched:
            paragraphs_changed += 1

    if styles_changed:
        details.append(
            f"Стили заголовков: интервал до/после обнулён ({styles_changed} шт.)"
        )
    if paragraphs_changed:
        details.append(
            f"Заголовки: интервал до/после = {target_sb:g}/{target_sa:g} пт "
            f"({paragraphs_changed} шт.)"
        )
    return (styles_changed + paragraphs_changed) > 0


_CHAPTER_DECOR_PREFIX_RE = re.compile(
    r"^[\s\xa0]*[\u2580-\u259f\u25b6\u25c0\u25c6\u25c7\u25cf\u25cb]+[\s\xa0]*",
    re.UNICODE,
)


def strip_chapter_decoration_chars(doc, details: list[str]) -> bool:
    """Strip decorative block-element symbols (▌▍▎▏█▓▒░◀▶◆◇●○) from
    the beginning of paragraphs.

    Some templates / copy-paste sources prepend these symbols to chapter
    titles (e.g. «▌Введение», «▌Глава 1»). The leftover prefix
    prevents the heading detection from recognising the paragraph as a
    real chapter title — without this strip pass headings would stay
    as plain ``Normal`` paragraphs and the auto-TOC would miss them
    entirely. We only touch paragraphs whose REMAINING text matches a
    typical heading shape (chapter keyword, well-known section title,
    «N.» numeric prefix, …) so generic body bullets stay intact.
    """
    from app.rules_engine.heading_detection import detect_heading_candidate

    changed = 0
    for para in doc.paragraphs:
        text = para.text or ""
        m = _CHAPTER_DECOR_PREFIX_RE.match(text)
        if not m:
            continue
        remainder = text[m.end():].strip()
        if not remainder:
            continue
        # Only strip when the cleaned text actually looks like a heading
        # (the safe-list prevents accidental edits to body bullet lists
        # that happen to use these block characters).
        if detect_heading_candidate(remainder) is None:
            continue
        # Walk runs and shave off the matching prefix.
        consumed = 0
        target = m.end()
        for r in para._element.iter(qn("w:r")):
            for t in r.findall(qn("w:t")):
                if consumed >= target:
                    break
                txt = t.text or ""
                if not txt:
                    continue
                take = min(len(txt), target - consumed)
                t.text = txt[take:]
                consumed += take
            if consumed >= target:
                break
        changed += 1

    if changed:
        details.append(
            f"Заголовки: убраны декоративные символы (▌, ◀, …) перед {changed} заголовком(ами)"
        )
    return changed > 0


def fix_remove_underline(paragraph, para_label: str, details: list[str]) -> bool:
    p_elem = paragraph._element
    changed = False
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is not None:
        rPr_default = pPr.find(qn("w:rPr"))
        if rPr_default is not None:
            el = rPr_default.find(qn("w:u"))
            if el is not None:
                rPr_default.remove(el)
                changed = True
    for r_elem in p_elem.iter(qn("w:r")):
        rPr = r_elem.find(qn("w:rPr"))
        if rPr is None:
            continue
        el = rPr.find(qn("w:u"))
        if el is not None:
            rPr.remove(el)
            changed = True
    if changed:
        details.append(f"{para_label}: подчёркивание убрано")
    return changed
