"""Autofix: sort bibliography alphabetically and number entries."""
from __future__ import annotations

import logging
import re

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from docx.text.run import Run


def _force_run_underline_none(r_el) -> bool:
    """Stamp ``<w:u w:val="none"/>`` on the run's ``<w:rPr>``.

    Required because character styles like ``Hyperlink`` add underline
    via style inheritance — only an explicit run-level ``w:u`` override
    disables them. Returns True if anything changed.
    """
    rPr = r_el.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r_el.insert(0, rPr)
    u = rPr.find(qn("w:u"))
    if u is None:
        u = OxmlElement("w:u")
        rPr.append(u)
        u.set(qn("w:val"), "none")
        return True
    if u.get(qn("w:val")) != "none":
        u.set(qn("w:val"), "none")
        return True
    return False

_HYPERLINK_STYLE_IDS = frozenset({"Hyperlink", "FollowedHyperlink"})

logger = logging.getLogger(__name__)

_BIBLIOGRAPHY_HEADINGS = frozenset({
    "список литературы",
    "список использованной литературы",
    "список используемой литературы",
    "список используемых источников",
    "список использованных источников",
    "список использованных источников и литературы",
    "список используемых источников и литературы",
    "список использованных источников информации",
    "библиографический список",
    "библиография",
    "литература",
    "references",
    "bibliography",
    "list of references",
})

# Extra phrases to recognise bibliography headings with minor variations
# (punctuation, extra words after the core phrase, etc.).
_BIBLIOGRAPHY_PATTERN = re.compile(
    r"^\s*(?:"
    r"список(?:\s+(?:использованн(?:ых|ой)|используем(?:ых|ой)))?\s+(?:источник(?:ов|и)|литератур[ыа])"
    r"|библиографическ(?:ий|ая)\s+список"
    r"|библиография"
    r"|references"
    r"|bibliography"
    r")\s*[:.]?\s*$",
    re.IGNORECASE | re.UNICODE,
)

_NUM_PREFIX_RE = re.compile(r"^\s*(?:\[\d{1,3}\][\.)\s]?|\d{1,3}[\.)])\s*")
_NUMBERED_ENTRY_RE = re.compile(r"^\s*(?:\[\d{1,3}\][\.)\s]?|\d{1,3}[\.)])\s+\S")
_LEADING_NONWORD_RE = re.compile(r"^[\s\W_]+", re.UNICODE)
_HEADING_STYLE_IDS = frozenset({f"Heading{i}" for i in range(1, 10)})

# Recognise common bibliography subsection titles ("Нормативно-правовые
# акты", "Судебная практика", "Научная литература", "Учебная литература",
# "Иные источники", "Электронные ресурсы", …). The whole paragraph must
# match — these are short label rows, not actual reference entries. Used
# to keep the multi-section structure intact instead of flattening
# everything alphabetically (which would, for example, push «Конституция»
# into the middle of «Учебная литература»).
_SUBSECTION_TITLE_RE = re.compile(
    r"^\s*(?:"
    r"нормативн[оы]\s*[\-–—]?\s*правов(?:ые|ой|о)\s+акт\w*"
    r"|нормативн[оы]\s+(?:и\s+иные\s+)?(?:правовые\s+)?акт\w*"
    r"|нпа"
    r"|официальн\w+\s+документ\w*"
    r"|международн\w+\s+(?:акт\w*|договор\w*)"
    r"|законы(?:\s+и\s+подзаконные\s+акты)?"
    r"|подзаконные\s+акты"
    r"|конституц\w+(?:\s+и\s+(?:федеральн\w+\s+)?закон\w*)?"
    r"|(?:материал\w+\s+)?судебн(?:ая|ой)\s+практик\w*"
    r"|акт\w*\s+(?:высш\w+\s+)?судебн\w+\s+орган\w*"
    r"|прав\w+\s+позиц\w+\s+(?:судов|судебн\w+)"
    r"|научн(?:ая|ые|о[\s\-]?метод\w*)\s+(?:литератур\w*|публикац\w+|источник\w+|труд\w+|стат\w+|исследован\w+)"
    r"|научно[\s\-]?учебн\w+\s+литератур\w*"
    r"|учебн(?:ая|ые|о[\s\-]?метод\w*)\s+литератур\w*"
    r"|(?:учебная\s+)?литератур\w+(?:\s+и\s+учебн\w+\s+пособ\w+)?"
    r"|учебники?(?:\s+и\s+учебные\s+пособ\w+)?"
    r"|учебн\w+\s+пособ\w+"
    r"|специальн(?:ая|ые)\s+литератур\w*"
    r"|монограф\w+(?:\s+и\s+стат\w+)?"
    r"|периодическ\w+(?:\s+(?:издани\w+|печат\w+))?"
    r"|стат\w+\s+в\s+период\w+\s+издан\w+"
    r"|диссертац\w+(?:\s+и\s+автореферат\w+)?"
    r"|автореферат\w+"
    r"|иные?\s+(?:источник\w+|документ\w+|материал\w+|публикац\w+)"
    r"|прочие?\s+(?:источник\w+|документ\w+|материал\w+)"
    r"|дополнительн\w+\s+(?:источник\w+|литератур\w+)"
    r"|электронн(?:ые|ый|ой)\s+(?:ресурс\w*|источник\w*)"
    r"|интернет[\s\-]ресурс\w*"
    r"|интернет[\s\-]источник\w*"
    r"|сайт\w+\s+в\s+интернет\w*"
    r"|справочно\W*правов\w+\s+систем\w+"
    r"|references"
    r"|primary\s+sources"
    r"|secondary\s+sources"
    r")\s*[:.;]?\s*$",
    re.IGNORECASE | re.UNICODE,
)


def _is_subsection_heading(text: str) -> bool:
    """Return True if *text* is a bibliography subsection title.

    The check is text-only (we cannot rely on bold any more — earlier
    autofix passes strip bold from body paragraphs before this function
    runs). Works on the trimmed paragraph text.
    """
    if not text:
        return False
    if len(text) > 80:
        return False
    if _NUMBERED_ENTRY_RE.match(text):
        return False
    return bool(_SUBSECTION_TITLE_RE.match(text.strip()))

_RU_ALPHABET = (
    "абвгдеёжзийклмнопрстуфхцчшщъыьэюя"
)
_RU_ORDER = {ch: i for i, ch in enumerate(_RU_ALPHABET)}


def _is_heading_para(paragraph) -> bool:
    sid = getattr(paragraph.style, "style_id", "") or ""
    if sid in _HEADING_STYLE_IDS:
        return True
    sname = (getattr(paragraph.style, "name", "") or "").lower()
    return "heading" in sname or "заголов" in sname


def _find_bibliography_range(doc) -> tuple[int, int, int] | None:
    """Return ``(heading_idx, start_idx, end_idx)`` for the bibliography section.

    ``heading_idx`` is the paragraph index of the «Список литературы» heading,
    ``start_idx`` is the first paragraph after it and ``end_idx`` is exclusive.

    Prefers a paragraph whose style is a real ``Heading``/``Заголовок`` over
    plain-text matches so we ignore the «Список литературы» row inside the
    auto-generated table of contents at the start of the document.
    """
    paragraphs = doc.paragraphs
    heading_match: int | None = None
    fallback_match: int | None = None

    for idx, para in enumerate(paragraphs):
        text = (para.text or "").strip()
        low = text.lower()
        if not (low in _BIBLIOGRAPHY_HEADINGS or _BIBLIOGRAPHY_PATTERN.match(low)):
            continue
        if _is_heading_para(para):
            heading_match = idx
            break
        if fallback_match is None:
            fallback_match = idx

    bib_heading_idx = heading_match if heading_match is not None else fallback_match
    if bib_heading_idx is None:
        return None

    start = bib_heading_idx + 1
    end = len(paragraphs)
    for idx in range(start, len(paragraphs)):
        para = paragraphs[idx]
        text = (para.text or "").strip()
        if not text:
            continue
        if _is_heading_para(para):
            end = idx
            break
        if text == text.upper() and len(text) > 3 and not re.match(r"^\s*[\[\d]", text):
            end = idx
            break

    return (bib_heading_idx, start, end)


def _strip_number_prefix(text: str) -> str:
    """Strip leading numeric prefixes ("1.", "[1]", "1)") one or more times.

    Bibliography entries sometimes contain double-numbering like "11.\t13. Ivanov"
    when they were manually copy-pasted; we drop every leading number so the
    alphabetical sort sees the author surname as the first significant token.
    """
    result = text
    for _ in range(4):
        stripped = _NUM_PREFIX_RE.sub("", result)
        if stripped == result:
            break
        result = stripped
    return result.strip()


def _sort_key(text: str) -> tuple:
    """Locale-friendly alphabetical key for Russian/Latin bibliography entries.

    Russian goes first (а–я, with «ё» sorted right after «е»), Latin second.
    Digits, brackets and other punctuation are stripped from the head so
    GOST-style entries like "[1] Иванов..." sort by the actual author name.
    """
    stripped = _strip_number_prefix(text)
    stripped = _LEADING_NONWORD_RE.sub("", stripped)
    lower = stripped.lower()
    key: list[tuple[int, int]] = []
    for ch in lower:
        if ch in _RU_ORDER:
            key.append((0, _RU_ORDER[ch]))
        elif "a" <= ch <= "z":
            key.append((1, ord(ch)))
        elif ch.isdigit():
            key.append((2, ord(ch)))
        elif ch.isspace():
            key.append((3, 0))
    return tuple(key)


def _set_paragraph_text(paragraph, new_text: str) -> None:
    """Replace paragraph text while preserving formatting of the first run."""
    runs = paragraph.runs
    if not runs:
        return

    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


def _collect_subsection_indices(
    paragraphs, start: int, end: int,
) -> list[int]:
    """Return paragraph indices (in the bibliography range) whose text
    looks like a subsection title («Нормативно-правовые акты», …).
    """
    out: list[int] = []
    for idx in range(start, end):
        text = (paragraphs[idx].text or "").strip()
        if _is_subsection_heading(text):
            out.append(idx)
    return out


def _enforce_subsection_bold(paragraph) -> None:
    """Re-stamp bold on a subsection heading paragraph.

    Earlier autofix passes strip bold from body-text paragraphs, so by
    the time we get here the «Нормативно-правовые акты» row is plain
    text. Customer expects subsection labels to stay bold.
    """
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is not None:
        rPr_def = pPr.find(qn("w:rPr"))
        if rPr_def is None:
            rPr_def = OxmlElement("w:rPr")
            pPr.append(rPr_def)
        if rPr_def.find(qn("w:b")) is None:
            rPr_def.append(OxmlElement("w:b"))
        if rPr_def.find(qn("w:bCs")) is None:
            rPr_def.append(OxmlElement("w:bCs"))

    for r_el in paragraph._element.iter(qn("w:r")):
        rPr = r_el.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            r_el.insert(0, rPr)
        if rPr.find(qn("w:b")) is None:
            rPr.append(OxmlElement("w:b"))
        if rPr.find(qn("w:bCs")) is None:
            rPr.append(OxmlElement("w:bCs"))


def _build_grouped_entries(
    paragraphs, start: int, end: int, subsection_indices: list[int],
) -> list[dict]:
    """Group bibliography paragraphs into per-subsection entry groups.

    Returns ``[{"subsection_idx": int|None, "entries": [entry, …]}]``
    where each *entry* is ``{"idxs": [...], "text": "..."}``. Subsection
    titles are used only as group dividers — they are NOT added to
    entry indices, so they can later be re-formatted independently
    (kept bold, no numbering).
    """
    sub_set = set(subsection_indices)
    groups: list[dict] = [{"subsection_idx": None, "entries": []}]
    current = groups[0]

    for idx in range(start, end):
        text = (paragraphs[idx].text or "").strip()
        if not text:
            continue
        if idx in sub_set:
            current = {"subsection_idx": idx, "entries": []}
            groups.append(current)
            continue
        if _NUMBERED_ENTRY_RE.match(text):
            current["entries"].append({"idxs": [idx], "text": text})
        else:
            if len(text) < 5:
                continue
            if current["entries"]:
                current["entries"][-1]["idxs"].append(idx)
                current["entries"][-1]["text"] = (
                    current["entries"][-1]["text"] + " " + text
                )
            else:
                current["entries"].append({"idxs": [idx], "text": text})

    # Drop the leading "no subsection" group if it contains no entries
    if not groups[0]["entries"] and len(groups) > 1:
        groups.pop(0)

    return groups


def fix_bibliography_order_and_numbering(
    doc, details: list[str],
) -> bool:
    """Sort/renumber bibliography entries.

    Two operating modes:

    * **No subsections detected** — flatten the list, sort
      alphabetically and renumber 1.…N. (Original behaviour.)
    * **Subsections detected** (НПА / Судебная практика / Научная
      литература / Учебная литература / Иные источники, …) — group
      entries by subsection, **preserve the order students wrote
      inside each section** (NPA hierarchy beats alphabetical order:
      Конституция must stay on top), and renumber continuously
      across the whole bibliography. Subsection labels stay where
      they are and have their bold re-applied.
    """
    rng = _find_bibliography_range(doc)
    if rng is None:
        return False

    heading_idx, start, end = rng
    paragraphs = doc.paragraphs

    subsection_indices = _collect_subsection_indices(paragraphs, start, end)

    if subsection_indices:
        return _fix_with_subsections(
            doc, details, heading_idx, start, end, subsection_indices,
        )

    return _fix_flat_alphabetical(doc, details, heading_idx, start, end)


def _fix_flat_alphabetical(
    doc, details: list[str], heading_idx: int, start: int, end: int,
) -> bool:
    paragraphs = doc.paragraphs

    entries: list[dict] = []
    for idx in range(start, end):
        text = (paragraphs[idx].text or "").strip()
        if not text:
            continue
        if _NUMBERED_ENTRY_RE.match(text):
            entries.append({"idxs": [idx], "text": text})
        else:
            if len(text) < 5:
                continue
            if entries:
                entries[-1]["idxs"].append(idx)
                entries[-1]["text"] = entries[-1]["text"] + " " + text

    # Fallback: when the bibliography is completely un-numbered (typical
    # student manuscript that just lists references one per paragraph
    # without «1. 2. 3.» prefixes), the loop above produces zero
    # entries. Treat every non-trivial paragraph in the bibliography
    # range as a separate entry so we still sort and number them.
    if not entries:
        for idx in range(start, end):
            text = (paragraphs[idx].text or "").strip()
            if len(text) < 10:
                continue
            entries.append({"idxs": [idx], "text": text})

    if len(entries) < 2:
        return False

    sorted_entries = sorted(entries, key=lambda e: _sort_key(e["text"]))

    already_sorted = all(
        _sort_key(entries[i]["text"]) <= _sort_key(entries[i + 1]["text"])
        for i in range(len(entries) - 1)
    )

    already_numbered = all(
        re.match(r"^\s*\d{1,3}[\.)]\s", e["text"]) for e in entries
    )

    if already_sorted and already_numbered:
        return False

    if not already_sorted:
        heading_elem = paragraphs[heading_idx]._element
        parent = heading_elem.getparent()
        if parent is None:
            return False

        all_entry_elements: list = []
        sorted_elements_groups: list[list] = []
        for e in entries:
            all_entry_elements.extend(paragraphs[i]._element for i in e["idxs"])
        for e in sorted_entries:
            sorted_elements_groups.append(
                [paragraphs[i]._element for i in e["idxs"]]
            )

        for elem in all_entry_elements:
            el_parent = elem.getparent()
            if el_parent is not None:
                el_parent.remove(elem)

        heading_pos = list(parent).index(heading_elem)
        insert_point = heading_pos + 1
        while insert_point < len(parent):
            child = parent[insert_point]
            if child.tag != qn("w:p"):
                break
            child_text = "".join(
                (t.text or "") for t in child.iter(qn("w:t"))
            ).strip()
            if child_text:
                break
            insert_point += 1

        cursor = insert_point
        for group in sorted_elements_groups:
            for elem in group:
                parent.insert(cursor, elem)
                cursor += 1

        details.append(
            f"Библиография: источники отсортированы по алфавиту ({len(entries)} шт.)"
        )

    refreshed = doc.paragraphs
    rng2 = _find_bibliography_range(doc)
    if rng2 is None:
        return True

    _, s2, e2 = rng2
    num = 1
    prev_blank = True
    for idx in range(s2, e2):
        text = (refreshed[idx].text or "").strip()
        if not text:
            prev_blank = True
            continue
        is_numbered = bool(_NUMBERED_ENTRY_RE.match(text))
        is_head = is_numbered or (prev_blank and len(text) >= 25) or len(text) >= 25
        if not is_head:
            prev_blank = False
            continue
        clean = _strip_number_prefix(text)
        clean = _NUM_PREFIX_RE.sub("", clean).strip()
        new_text = f"{num}. {clean}"
        _set_paragraph_text(refreshed[idx], new_text)
        num += 1
        prev_blank = False

    if num > 1:
        details.append(
            f"Библиография: источники пронумерованы 1–{num - 1}"
        )
    return True


def _fix_with_subsections(
    doc, details: list[str], heading_idx: int, start: int, end: int,
    subsection_indices: list[int],
) -> bool:
    """Renumber bibliography entries while keeping subsection structure.

    Order inside each subsection is intentionally preserved — students
    write NPA in legal-force hierarchy («Конституция» → ГК → ЗК → ЖК →
    федеральные законы), and alphabetical sorting would break that.
    """
    paragraphs = doc.paragraphs
    groups = _build_grouped_entries(paragraphs, start, end, subsection_indices)

    total_entries = sum(len(g["entries"]) for g in groups)
    if total_entries < 2:
        return False

    # Renumber in place: per-subsection order stays the same, numbers
    # run continuously across the whole bibliography.
    num = 1
    for group in groups:
        for entry in group["entries"]:
            head_idx = entry["idxs"][0]
            head_para = paragraphs[head_idx]
            text = (head_para.text or "").strip()
            clean = _strip_number_prefix(text)
            clean = _NUM_PREFIX_RE.sub("", clean).strip()
            new_text = f"{num}. {clean}"
            if new_text != text:
                _set_paragraph_text(head_para, new_text)
            num += 1

    # Re-stamp bold on subsection rows so they visually stand out
    # again after general bold-stripping passes ran.
    for sub_idx in subsection_indices:
        _enforce_subsection_bold(paragraphs[sub_idx])

    details.append(
        f"Библиография: разделы сохранены ({len(subsection_indices)} шт.), "
        f"источники пронумерованы сквозной нумерацией 1–{num - 1}"
    )
    return True


def enforce_bibliography_entry_formatting(
    doc,
    details: list[str],
    *,
    line_spacing: float,
    first_line_indent_mm: float,
    space_after_pt: float,
    font_name: str,
    font_size_pt: float,
) -> bool:
    """Force every bibliography entry to follow body-text formatting rules.

    Customer requirement: bibliography entries must use the same paragraph
    settings as body text (1.25 cm first-line indent, 1.5 line spacing,
    Times New Roman 14 pt, no bold, no underline) and stay justified. The
    «Список литературы» heading itself is left untouched — it is promoted
    to ``Heading 1`` elsewhere and inherits its own formatting.
    """
    rng = _find_bibliography_range(doc)
    if rng is None:
        return False

    _, start, end = rng
    paragraphs = doc.paragraphs
    touched = 0

    for idx in range(start, end):
        para = paragraphs[idx]
        text = (para.text or "").strip()
        if not text:
            continue
        if _is_heading_para(para):
            continue
        # Subsection rows («Нормативно-правовые акты», «Судебная
        # практика», …) must keep their bold and stay un-indented;
        # they are dividers, not regular reference entries.
        if _is_subsection_heading(text):
            continue
        if _apply_entry_formatting(
            para, line_spacing, first_line_indent_mm, space_after_pt,
            font_name, font_size_pt,
        ):
            touched += 1

    if touched:
        details.append(
            f"Библиография: {touched} записей приведено к телу "
            f"(абзац {first_line_indent_mm} мм, интервал {line_spacing}, без жирного)"
        )
    return touched > 0


def _apply_entry_formatting(
    para,
    line_spacing: float,
    first_line_indent_mm: float,
    space_after_pt: float,
    font_name: str,
    font_size_pt: float,
) -> bool:
    pf = para.paragraph_format
    changed = False

    if para.alignment != WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        changed = True
    if pf.first_line_indent is None or abs(float(pf.first_line_indent.mm) - first_line_indent_mm) > 0.5:
        pf.first_line_indent = Mm(first_line_indent_mm)
        changed = True
    if pf.left_indent is not None and int(pf.left_indent) != 0:
        pf.left_indent = Mm(0)
        changed = True
    if pf.line_spacing is None or abs(float(pf.line_spacing) - line_spacing) > 0.05:
        pf.line_spacing = line_spacing
        changed = True
    if pf.space_after is None or abs(float(pf.space_after.pt) - space_after_pt) > 0.2:
        pf.space_after = Pt(space_after_pt)
        changed = True
    if pf.space_before is not None and int(pf.space_before) != 0:
        pf.space_before = Pt(0)
        changed = True

    pPr = para._element.find(qn("w:pPr"))
    if pPr is not None:
        rPr_default = pPr.find(qn("w:rPr"))
        if rPr_default is not None:
            for tag_name in ("w:b", "w:bCs", "w:u"):
                el = rPr_default.find(qn(tag_name))
                if el is not None:
                    rPr_default.remove(el)
                    changed = True

    for r_el in para._element.iter(qn("w:r")):
        if _normalize_run(r_el, para, font_name, font_size_pt):
            changed = True

    return changed


def _normalize_run(
    r_el,
    parent,
    font_name: str,
    font_size_pt: float,
) -> bool:
    """Strip hyperlink/bold/underline styling from ``<w:r>`` and enforce font.

    Covers run elements nested inside ``<w:hyperlink>`` (which ``Paragraph.runs``
    does not expose) and removes ``w:rStyle`` references to the ``Hyperlink``
    character style together with any direct color override, so the entry is
    rendered as plain body text even for URL links.
    """
    run = Run(r_el, parent)
    changed = False

    if run.bold:
        run.bold = False
        changed = True
    # Always force ``<w:u w:val="none"/>`` — ``run.font.underline`` is
    # ``None`` when the underline is inherited from a character style,
    # and truthy only for explicit run-level overrides. Testing the
    # truthiness would skip inherited cases and leave bibliography URLs
    # underlined. We write ``None`` to drop any run-level element and
    # then explicitly stamp ``none`` so inheritance cannot take effect.
    run.font.underline = None
    if _force_run_underline_none(r_el):
        changed = True
    if font_name and run.font.name != font_name:
        run.font.name = font_name
        changed = True
    if run.font.size is None or abs(float(run.font.size.pt) - font_size_pt) > 0.2:
        run.font.size = Pt(font_size_pt)
        changed = True

    rPr = r_el.find(qn("w:rPr"))
    if rPr is None:
        return changed

    rStyle = rPr.find(qn("w:rStyle"))
    if rStyle is not None and rStyle.get(qn("w:val")) in _HYPERLINK_STYLE_IDS:
        rPr.remove(rStyle)
        changed = True

    color = rPr.find(qn("w:color"))
    if color is not None:
        rPr.remove(color)
        changed = True

    bCs = rPr.find(qn("w:bCs"))
    if bCs is not None:
        val = bCs.get(qn("w:val"))
        if val not in ("0", "false"):
            bCs.set(qn("w:val"), "0")
            changed = True

    return changed
