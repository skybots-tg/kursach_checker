from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm, Pt

from app.rules_engine.autofix_config import AutoFixConfig as _AutoFixConfig
from app.rules_engine.autofix_refresh import refresh_fields_via_libreoffice
from app.rules_engine.autofix_headings import (
    ensure_blank_before_subheadings,
    enforce_chapter_page_breaks,
    enforce_heading_bold,
    enforce_subheading_alignment,
    fix_heading as _fix_heading,
    fix_remove_underline,
    promote_to_heading as _promote_to_heading,
)
from app.rules_engine.autofix_helpers import (
    clamp_overflow_table_widths,
    fix_caption_trailing_dot,
    fix_dashes_in_text,
    fix_font_color_runs,
    fix_font_color_styles,
    fix_italic_styles,
    fix_list_indent,
    fix_markers_text,
    fix_numbering_bullets,
    fix_page_break_before,
    fix_remove_highlight,
    fix_remove_italic,
    fix_remove_strange_chars,
    fix_strip_markdown_artifacts,
    fix_section_margins,
    is_field_code_run,
    is_manual_list_para,
    postprocess_fixed_docx,
    preflight_margins_safe,
    remove_empty_paras_before_page_breaks,
    remove_manual_page_breaks,
)
from app.rules_engine.autofix_bibliography import (
    enforce_bibliography_entry_formatting,
    fix_bibliography_order_and_numbering,
)
from app.rules_engine.autofix_captions import fix_caption_positions
from app.rules_engine.autofix_table_pass import process_table_cells
from app.rules_engine.autofix_lists import convert_informal_lists
from app.rules_engine.autofix_toc import insert_toc_field, detect_manual_toc_entry_indices
from app.rules_engine.autofix_split_breaks import split_soft_break_paragraphs
from app.rules_engine.autofix_redundant_breaks import remove_redundant_manual_page_breaks
from app.rules_engine.autofix_toc_normalize import normalize_toc_heading_formatting
from app.rules_engine.autofix_whitespace import (
    collapse_excessive_empty_paras,
    fix_normalize_left_indent,
    fix_strip_leading_whitespace,
    normalize_doc_defaults_spacing,
    normalize_source_line_spacing,
    normalize_title_page_spacing,
)
from app.rules_engine.checks_content import ALLOWED_CHARS_RE as _ALLOWED_CHARS_RE
from app.rules_engine.autofix_para_classify import (
    collect_toc_heading_levels,
    is_heading_para,
    is_list_para,
    should_skip_para,
)
from app.rules_engine.heading_detection import (
    CHAPTER_RE as _CHAPTER_RE,
    TOC_LINE_TAIL_RE as _TOC_LINE_TAIL_RE,
    detect_heading_candidate,
    detect_heading_via_toc,
)
from app.rules_engine.findings import Finding
from app.rules_engine.style_resolve import (
    detect_toc_paragraph_indices,
    effective_alignment,
    effective_first_line_indent_mm,
    effective_font_name,
    effective_font_size_pt,
    effective_line_spacing,
    effective_space_after_pt,
    effective_space_before_pt,
)

logger = logging.getLogger(__name__)


@dataclass(slots=True)
class AutoFixResult:
    output_file_path: str | None
    details: list[str]

def apply_safe_autofixes(
    file_path: str, rules: dict | None, findings: list[Finding],
    admin_autofix_config: dict | None = None,
) -> AutoFixResult:
    source = Path(file_path)
    if source.suffix.lower() != ".docx" or not source.exists():
        return AutoFixResult(output_file_path=None, details=[])

    admin_cfg = admin_autofix_config or {}
    cfg = _AutoFixConfig.from_rules(rules, admin_defaults=admin_cfg.get("defaults"))
    if not cfg.enabled:
        return AutoFixResult(output_file_path=None, details=[])

    safety = admin_cfg.get("safety_limits", {})
    max_paragraphs = int(safety.get("max_paragraphs_touched", 2000))
    skip_toc = bool(safety.get("skip_toc", True))
    skip_headings_safety = bool(safety.get("skip_headings", True))
    allow_promote = bool(safety.get("allow_promote_heading_candidates", cfg.promote_heading_candidates))
    skip_tables_safety = bool(safety.get("skip_tables", True))
    skip_margins_safety = bool(safety.get("skip_margin_normalization", True))
    use_libreoffice_refresh = bool(safety.get("libreoffice_toc_refresh", False))

    try:
        doc = Document(str(source))
    except Exception:
        logger.warning("Autofix: cannot open DOCX %s", file_path)
        return AutoFixResult(output_file_path=None, details=[])

    changed = False
    details: list[str] = []
    change_count = 0

    # Convert soft line breaks into proper paragraph marks BEFORE we compute
    # any indices (the splitting renumbers paragraphs). Students often press
    # Shift+Enter between sentences, which prevents the body-text indent and
    # paragraph-style passes from working correctly because the entire
    # introduction lives in a single ``<w:p>``.
    if split_soft_break_paragraphs(doc, details):
        changed = True

    toc_indices: set[int] = set()
    if skip_toc:
        toc_indices = detect_toc_paragraph_indices(doc)

    # Manual TOC lines usually don't carry a page-number tail (students just
    # type the section titles). ``detect_toc_paragraph_indices`` only catches
    # entries that do, so we augment ``toc_indices`` with a more permissive
    # heuristic pass before running any heading/normalization logic. This is
    # essential: without it, the subsequent TOC-based heading promotion would
    # mark the manual TOC lines themselves as headings and ``_remove_manual_
    # toc_entries`` would then refuse to delete them, leaving two tables of
    # contents glued together in the output.
    toc_indices |= detect_manual_toc_entry_indices(doc)

    toc_heading_levels: dict[str, int] = collect_toc_heading_levels(doc, toc_indices)

    if cfg.normalize_line_spacing or cfg.normalize_spacing_before_after:
        if normalize_doc_defaults_spacing(
            doc, cfg.line_spacing, cfg.space_after_pt, details,
        ):
            changed = True

    if cfg.normalize_font_color and fix_font_color_styles(doc, details):
        changed = True
    if cfg.remove_italic and fix_italic_styles(doc, details):
        changed = True
    if cfg.normalize_list_markers and fix_numbering_bullets(doc, cfg.font_name, details, cfg.list_marker_char):
        changed = True

    if cfg.convert_informal_lists:
        if convert_informal_lists(
            doc, cfg.informal_list_markers, cfg.list_marker_char,
            cfg.informal_list_min_consecutive, toc_indices, details,
        ):
            changed = True

    if cfg.normalize_margins and not skip_margins_safety:
        if preflight_margins_safe(doc, cfg.margins_mm):
            for sec_idx, section in enumerate(doc.sections):
                if fix_section_margins(section, cfg.margins_mm, sec_idx, details):
                    changed = True
                    change_count += 1
        else:
            logger.info("Autofix: margin normalization skipped — tables would overflow")

    for idx, paragraph in enumerate(doc.paragraphs):
        text = (paragraph.text or "").strip()
        if not text:
            continue
        para_label = f"\u0410\u0431\u0437\u0430\u0446 #{idx + 1}"
        if cfg.remove_highlight:
            if fix_remove_highlight(paragraph, para_label, details):
                changed = True
        if cfg.remove_strange_chars:
            if fix_remove_strange_chars(paragraph, para_label, details, _ALLOWED_CHARS_RE):
                changed = True
        if cfg.strip_leading_whitespace and idx not in toc_indices:
            if fix_strip_leading_whitespace(paragraph, para_label, details):
                changed = True
        if (
            cfg.fix_section_breaks
            and idx > 0
            and len(text) <= 100
            and idx not in toc_indices
        ):
            if not _TOC_LINE_TAIL_RE.search(text):
                needs_break = any(s in text.lower() for s in cfg.section_break_sections)
                if not needs_break:
                    needs_break = bool(_CHAPTER_RE.match(text))
                if needs_break:
                    prev_idx = idx - 1
                    while prev_idx >= 0:
                        prev_para = doc.paragraphs[prev_idx]
                        if (prev_para.text or "").strip():
                            break
                        if remove_manual_page_breaks(prev_para):
                            changed = True
                        prev_idx -= 1
                    if remove_manual_page_breaks(paragraph):
                        changed = True
                    if fix_page_break_before(paragraph, para_label, details):
                        changed = True

    body_start = 0
    for i, p in enumerate(doc.paragraphs):
        if is_heading_para(p):
            body_start = i
            break

    if body_start > 0 and cfg.normalize_spacing_before_after:
        if normalize_title_page_spacing(doc, body_start, details):
            changed = True

    para_count = 0
    for idx, paragraph in enumerate(doc.paragraphs):
        if para_count >= max_paragraphs:
            break

        text = (paragraph.text or "").strip()
        if not text:
            continue

        para_label = f"\u0410\u0431\u0437\u0430\u0446 #{idx + 1}"
        para_touched = False

        if cfg.normalize_font_color:
            if fix_font_color_runs(paragraph, para_label, details):
                changed = True
                para_touched = True

        if cfg.remove_italic:
            if fix_remove_italic(paragraph, para_label, details):
                changed = True
                para_touched = True

        if cfg.remove_underline:
            if fix_remove_underline(paragraph, para_label, details):
                changed = True
                para_touched = True

        if cfg.remove_caption_trailing_dot:
            if fix_caption_trailing_dot(paragraph, para_label, details):
                changed = True
                para_touched = True

        if idx in toc_indices:
            if para_touched:
                para_count += 1
            continue

        is_heading = is_heading_para(paragraph)
        candidate_level = detect_heading_candidate(text) if not is_heading else None
        if not is_heading and candidate_level is None and toc_heading_levels:
            candidate_level = detect_heading_via_toc(text, toc_heading_levels)

        if is_heading:
            if not skip_headings_safety and cfg.normalize_headings:
                if _fix_heading(paragraph, idx, cfg, details):
                    changed = True
                    para_touched = True
            if cfg.normalize_dashes:
                if fix_dashes_in_text(paragraph, para_label, details):
                    changed = True
                    para_touched = True
            if para_touched:
                para_count += 1
            continue

        if candidate_level is not None and allow_promote:
            # «Содержание»/«Оглавление» must stay as a plain centered paragraph
            # (client's explicit request). Skip heading promotion for it —
            # normalize_toc_heading_formatting will format it correctly later.
            low = text.strip().lower().rstrip(":.")
            if low in ("содержание", "оглавление"):
                if para_touched:
                    para_count += 1
                continue
            if _promote_to_heading(paragraph, candidate_level, idx, cfg, details):
                changed = True
                para_touched = True
            if para_touched:
                para_count += 1
            continue

        if candidate_level is not None:
            if para_touched:
                para_count += 1
            continue

        if should_skip_para(paragraph):
            if para_touched:
                para_count += 1
            continue

        if idx < body_start:
            if para_touched:
                para_count += 1
            continue

        eff_align = effective_alignment(paragraph)
        is_center_like = eff_align in (
            WD_PARAGRAPH_ALIGNMENT.CENTER,
            WD_PARAGRAPH_ALIGNMENT.RIGHT,
        )

        is_word_list = is_list_para(paragraph)
        is_manual = not is_word_list and is_manual_list_para(text)
        is_list = is_word_list or is_manual
        pf = paragraph.paragraph_format

        if not is_center_like and is_list and cfg.normalize_list_markers:
            if fix_markers_text(paragraph, para_label, details, cfg.list_marker_char):
                changed = True
                para_touched = True

        if not is_center_like and is_word_list and cfg.normalize_list_indent:
            if fix_list_indent(paragraph, para_label, details):
                changed = True
                para_touched = True

        if cfg.normalize_dashes:
            if fix_dashes_in_text(paragraph, para_label, details):
                changed = True
                para_touched = True

        if fix_strip_markdown_artifacts(paragraph, para_label, details):
            changed = True
            para_touched = True

        if not is_center_like and cfg.normalize_alignment:
            if eff_align != WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                changed = True
                para_touched = True
                details.append(f"{para_label}: выравнивание по ширине")

        if cfg.normalize_line_spacing:
            eff_ls = effective_line_spacing(paragraph)
            if eff_ls is not None and abs(eff_ls - cfg.line_spacing) > 0.05:
                pf.line_spacing = cfg.line_spacing
                changed = True
                para_touched = True
                details.append(f"{para_label}: межстрочный интервал {cfg.line_spacing}")

        if not is_center_like and not is_word_list and cfg.normalize_first_line_indent:
            eff_indent = effective_first_line_indent_mm(paragraph)
            if abs(eff_indent - cfg.first_line_indent_mm) > 0.5:
                pf.first_line_indent = Mm(cfg.first_line_indent_mm)
                changed = True
                para_touched = True
                details.append(f"{para_label}: абзацный отступ {cfg.first_line_indent_mm} мм")

        if not is_center_like and not is_list and cfg.normalize_body_left_indent:
            if fix_normalize_left_indent(paragraph, para_label, details):
                changed = True
                para_touched = True

        if not is_center_like and cfg.normalize_spacing_before_after:
            eff_sb = effective_space_before_pt(paragraph)
            if abs(eff_sb - cfg.space_before_pt) > 0.2:
                pf.space_before = Pt(cfg.space_before_pt)
                changed = True
                para_touched = True
                details.append(f"{para_label}: интервал до {cfg.space_before_pt} пт")
            eff_sa = effective_space_after_pt(paragraph)
            if abs(eff_sa - cfg.space_after_pt) > 0.2:
                pf.space_after = Pt(cfg.space_after_pt)
                changed = True
                para_touched = True
                details.append(f"{para_label}: интервал после {cfg.space_after_pt} пт")

        if cfg.normalize_font:
            font_changed = False
            for run in paragraph.runs:
                if is_field_code_run(run):
                    continue
                eff_name = effective_font_name(run, paragraph)
                if cfg.font_name and eff_name and eff_name != cfg.font_name:
                    run.font.name = cfg.font_name
                    font_changed = True
                eff_size = effective_font_size_pt(run, paragraph)
                if eff_size is not None and abs(eff_size - cfg.font_size_pt) > 0.2:
                    run.font.size = Pt(cfg.font_size_pt)
                    font_changed = True
            if font_changed:
                changed = True
                para_touched = True
                details.append(f"{para_label}: шрифт {cfg.font_name}, {cfg.font_size_pt} пт")

        if para_touched:
            para_count += 1

    table_changed, para_count = process_table_cells(
        doc, cfg, max_paragraphs, para_count, details,
    )
    changed |= table_changed

    if cfg.normalize_table_width and not skip_tables_safety and clamp_overflow_table_widths(doc, details):
        changed = True

    if cfg.generate_toc and insert_toc_field(doc, toc_indices, details):
        changed = True
    if cfg.normalize_toc_heading and normalize_toc_heading_formatting(doc, details, cfg=cfg):
        changed = True

    if cfg.fix_bibliography:
        changed |= fix_bibliography_order_and_numbering(doc, details)
        if enforce_bibliography_entry_formatting(
            doc,
            details,
            line_spacing=cfg.line_spacing,
            first_line_indent_mm=cfg.first_line_indent_mm,
            space_after_pt=cfg.space_after_pt,
            font_name=cfg.font_name,
            font_size_pt=cfg.font_size_pt,
        ):
            changed = True

    if getattr(cfg, "fix_caption_positions", True) and fix_caption_positions(doc, details):
        changed = True
    changed |= normalize_source_line_spacing(doc, details)
    changed |= enforce_subheading_alignment(doc, cfg, details)
    changed |= enforce_heading_bold(doc, cfg, details)
    if cfg.fix_section_breaks and enforce_chapter_page_breaks(doc, details):
        changed = True
    if cfg.ensure_subheading_spacing and ensure_blank_before_subheadings(doc, details):
        changed = True
    if cfg.collapse_empty_paras and collapse_excessive_empty_paras(doc, cfg.max_consecutive_empty_paras, details):
        changed = True

    if changed:
        if remove_redundant_manual_page_breaks(doc, details):
            pass
        if remove_empty_paras_before_page_breaks(doc, details):
            pass

    if not changed:
        return AutoFixResult(output_file_path=None, details=[])

    output = source.with_name(f"{source.stem}.fixed.docx")
    try:
        doc.save(str(output))
    except Exception:
        logger.exception("Autofix: failed to save fixed DOCX to %s", output)
        return AutoFixResult(output_file_path=None, details=[])

    try:
        postprocess_fixed_docx(source, output)
    except Exception:
        logger.warning("Autofix: postprocessing failed for %s, using python-docx output", output)

    if cfg.generate_toc and use_libreoffice_refresh:
        if refresh_fields_via_libreoffice(output, details):
            logger.info("Autofix: LibreOffice refreshed fields for %s", output)

    findings.append(
        Finding(
            title="\u0410\u0432\u0442\u043e\u0438\u0441\u043f\u0440\u0430\u0432\u043b\u0435\u043d\u0438\u044f",
            category="autofix",
            severity="advice",
            expected="",
            found="",
            location="\u0434\u043e\u043a\u0443\u043c\u0435\u043d\u0442",
            recommendation="\u0421\u043a\u0430\u0447\u0430\u0439\u0442\u0435 \u0438\u0441\u043f\u0440\u0430\u0432\u043b\u0435\u043d\u043d\u044b\u0439 DOCX",
            auto_fixed=True,
            auto_fix_details="; ".join(details[:30]),
        )
    )

    return AutoFixResult(output_file_path=str(output), details=details)


