"""Autofix: position and align figure/table captions according to GOST.

Rules applied:
    * «Рисунок N — …» caption belongs **below** the figure and is centered.
    * «Таблица N — …» caption belongs **above** the table, left-aligned,
      without a red-line indent.

If the caption already sits in the right place we only normalise the
alignment/indent. If the caption is on the wrong side of its target
(e.g. «Рисунок 2» typed above the chart) the paragraph element is moved
next to the image / table in the document body.
"""
from __future__ import annotations

import logging
import re

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Mm

logger = logging.getLogger(__name__)

_FIGURE_CAPTION_RE = re.compile(
    r"^\s*(?:рисунок|рис\.?|figure|fig\.?)\s*\d+",
    re.IGNORECASE | re.UNICODE,
)
_TABLE_CAPTION_RE = re.compile(
    r"^\s*таблица\s*\d+",
    re.IGNORECASE | re.UNICODE,
)
_SOURCE_LINE_RE = re.compile(
    r"^\s*(?:источник|source|примечание|note|составлено)\s*[:\u2014\u2013-]",
    re.IGNORECASE | re.UNICODE,
)

_DRAWING_TAGS = (qn("w:drawing"), qn("w:pict"), qn("w:object"))


def _para_has_image(p_elem) -> bool:
    """Return True if paragraph element contains inline/anchored image objects."""
    if p_elem is None or p_elem.tag != qn("w:p"):
        return False
    for tag in _DRAWING_TAGS:
        if p_elem.find(f".//{tag}") is not None:
            return True
    return False


def _prev_sibling(elem):
    sib = elem.getprevious()
    while sib is not None and sib.tag is etree_comment_tag():
        sib = sib.getprevious()
    return sib


def _next_sibling(elem):
    sib = elem.getnext()
    while sib is not None and sib.tag is etree_comment_tag():
        sib = sib.getnext()
    return sib


def etree_comment_tag():
    from lxml import etree
    return etree.Comment


def _prev_nonempty_sibling(elem):
    """Skip over empty paragraphs until we find a meaningful element."""
    sib = elem.getprevious()
    while sib is not None:
        if sib.tag == qn("w:p"):
            text = "".join(
                (t.text or "") for t in sib.iter(qn("w:t"))
            ).strip()
            if text or _para_has_image(sib):
                return sib
            sib = sib.getprevious()
            continue
        if sib.tag == qn("w:tbl"):
            return sib
        sib = sib.getprevious()
    return None


def _next_nonempty_sibling(elem):
    sib = elem.getnext()
    while sib is not None:
        if sib.tag == qn("w:p"):
            text = "".join(
                (t.text or "") for t in sib.iter(qn("w:t"))
            ).strip()
            if text or _para_has_image(sib):
                return sib
            sib = sib.getnext()
            continue
        if sib.tag == qn("w:tbl"):
            return sib
        sib = sib.getnext()
    return None


def _format_figure_caption(para) -> bool:
    """Center the caption paragraph and wipe red-line / left indents."""
    changed = False
    if para.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        changed = True
    pf = para.paragraph_format
    if pf.first_line_indent is not None and int(pf.first_line_indent) != 0:
        pf.first_line_indent = Mm(0)
        changed = True
    if pf.left_indent is not None and int(pf.left_indent) != 0:
        pf.left_indent = Mm(0)
        changed = True
    return changed


def _format_table_caption(para) -> bool:
    """Apply customer-mandated layout to a table caption row.

    Customer rules (verbatim):
        * «в одну строчку» — collapse soft breaks, drop leading tabs,
          single line spacing.
        * «без выделения» — strip highlight + bold/italic/underline.
        * «без абзаца» — wipe red-line and left indent.
        * «одинарный интервал» — line spacing 1.0.
    """
    from docx.oxml import OxmlElement
    changed = False
    if para.alignment not in (WD_PARAGRAPH_ALIGNMENT.LEFT, None):
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        changed = True
    pf = para.paragraph_format
    if pf.first_line_indent is not None and int(pf.first_line_indent) != 0:
        pf.first_line_indent = Mm(0)
        changed = True
    if pf.left_indent is not None and int(pf.left_indent) != 0:
        pf.left_indent = Mm(0)
        changed = True

    p_elem = para._element

    # Collapse soft breaks («Таблица 1\nХарактеристика выборки» split via
    # Shift+Enter) to a single space so the caption literally fits in
    # one row of source text. Word-level wrapping still kicks in if the
    # caption is wider than the page, but the single-line constraint
    # the customer asked for is about the source markup, not visual
    # wrap.
    for br in list(p_elem.iter(qn("w:br"))):
        if br.get(qn("w:type")) in (None, "textWrapping"):
            parent = br.getparent()
            if parent is not None:
                parent.remove(br)
                changed = True

    # Drop leading tab characters that students sometimes type before
    # «Таблица N» to align with body text — they push the caption off
    # to the right and break the «без абзаца» rule.
    runs = list(p_elem.iter(qn("w:r")))
    for r in runs:
        first_child = next(iter(r), None)
        if first_child is None:
            continue
        if first_child.tag == qn("w:tab"):
            r.remove(first_child)
            changed = True
        t_first = r.find(qn("w:t"))
        if t_first is not None and (t_first.text or "").startswith("\t"):
            t_first.text = (t_first.text or "").lstrip("\t ")
            changed = True
        break

    # Force ``<w:spacing w:line="240" w:lineRule="auto" w:before="0"
    # w:after="0"/>`` so the caption renders at single line height.
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_elem.insert(0, pPr)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    if spacing.get(qn("w:line")) != "240":
        spacing.set(qn("w:line"), "240")
        changed = True
    if spacing.get(qn("w:lineRule")) != "auto":
        spacing.set(qn("w:lineRule"), "auto")
        changed = True
    if spacing.get(qn("w:before")) not in (None, "0"):
        spacing.set(qn("w:before"), "0")
        changed = True
    if spacing.get(qn("w:after")) not in (None, "0"):
        spacing.set(qn("w:after"), "0")
        changed = True

    # Strip highlight, bold, italic, underline from every run so the
    # caption renders as plain body text. Default-paragraph rPr too.
    if _strip_caption_text_decoration(pPr.find(qn("w:rPr"))):
        changed = True
    for r in p_elem.iter(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            continue
        if _strip_caption_text_decoration(rPr):
            changed = True

    return changed


def _strip_caption_text_decoration(rPr) -> bool:
    """Remove highlight / bold / italic / underline from a ``<w:rPr>``."""
    if rPr is None:
        return False
    changed = False
    for tag in ("w:highlight", "w:b", "w:bCs", "w:i", "w:iCs", "w:u"):
        el = rPr.find(qn(tag))
        if el is not None:
            rPr.remove(el)
            changed = True
    shd = rPr.find(qn("w:shd"))
    if shd is not None and shd.get(qn("w:fill")) not in (None, "auto"):
        shd.set(qn("w:fill"), "auto")
        changed = True
    return changed


def _merge_table_caption_with_description(p_elem, details: list[str]) -> bool:
    """Merge a bare «Таблица N» paragraph with the descriptive line that
    immediately follows it.

    Many student documents look like:

        Таблица 1
        Характеристика выборки исследования
        <w:tbl>…</w:tbl>

    The caption rule «в одну строчку» requires that we render this as
    a single paragraph «Таблица 1 – Характеристика выборки
    исследования». We mutate *p_elem* in place: append « – <description>»
    and remove the description paragraph from the body.

    Returns True when a merge happened.
    """
    text = "".join((t.text or "") for t in p_elem.iter(qn("w:t"))).strip()
    if not text:
        return False
    if not _TABLE_CAPTION_RE.match(text):
        return False
    # Already has a description (text after the leading «Таблица N»)?
    head_match = re.match(r"^\s*таблица\s*[\dА-ЯЁA-Z\.]+", text, re.IGNORECASE | re.UNICODE)
    if head_match is None:
        return False
    tail = text[head_match.end():].strip()
    if tail:
        # Already «Таблица N – …», nothing to merge with.
        return False

    nxt = _next_nonempty_sibling(p_elem)
    if nxt is None or nxt.tag != qn("w:p"):
        return False
    desc_text = "".join((t.text or "") for t in nxt.iter(qn("w:t"))).strip()
    if not desc_text:
        return False
    # Don't pull in another caption / heading / source line.
    if _TABLE_CAPTION_RE.match(desc_text) or _FIGURE_CAPTION_RE.match(desc_text) or _SOURCE_LINE_RE.match(desc_text):
        return False
    nxt_pPr = nxt.find(qn("w:pPr"))
    if nxt_pPr is not None:
        ps = nxt_pPr.find(qn("w:pStyle"))
        if ps is not None and (ps.get(qn("w:val")) or "").startswith("Heading"):
            return False
    # Don't merge with a paragraph that contains an image / drawing.
    if _para_has_image(nxt):
        return False

    # The next sibling AFTER the description paragraph should be a
    # ``<w:tbl>`` for the merge to make sense; otherwise we'd be
    # gluing unrelated body text onto «Таблица N».
    after_desc = _next_nonempty_sibling(nxt)
    if after_desc is None or after_desc.tag != qn("w:tbl"):
        return False

    # Append « – description» to the caption's last text run.
    t_elements = list(p_elem.iter(qn("w:t")))
    if not t_elements:
        return False
    last_t = t_elements[-1]
    last_t.text = (last_t.text or "") + " \u2013 " + desc_text

    parent = nxt.getparent()
    if parent is not None:
        parent.remove(nxt)
    details.append(
        "Подписи: «Таблица N» объединена с описанием в одну строчку"
    )
    return True


def fix_caption_positions(doc, details: list[str]) -> bool:
    """Move figure captions below images, table captions above tables; align them."""
    changed = False
    fig_moved = 0
    fig_formatted = 0
    tbl_moved = 0
    tbl_formatted = 0

    # Merge bare «Таблица N» rows with the description sitting on the
    # following paragraph BEFORE the alignment / spacing pass — once
    # the rows are merged the existing logic handles them correctly.
    for para in list(doc.paragraphs):
        p_elem = para._element
        if p_elem.getparent() is None:
            continue
        if _merge_table_caption_with_description(p_elem, details):
            changed = True

    for para in list(doc.paragraphs):
        text = (para.text or "").strip()
        if not text:
            continue
        p_elem = para._element
        if p_elem.getparent() is None:
            continue

        if _FIGURE_CAPTION_RE.match(text):
            # Figure caption: should be right below an image paragraph.
            prev = _prev_nonempty_sibling(p_elem)
            if prev is not None and prev.tag == qn("w:p") and _para_has_image(prev):
                if _format_figure_caption(para):
                    changed = True
                    fig_formatted += 1
                continue

            nxt = _next_nonempty_sibling(p_elem)
            if nxt is not None and nxt.tag == qn("w:p") and _para_has_image(nxt):
                # Caption is above the image; move it below.
                parent = p_elem.getparent()
                if parent is not None:
                    parent.remove(p_elem)
                    nxt.addnext(p_elem)
                    changed = True
                    fig_moved += 1
                if _format_figure_caption(para):
                    changed = True
                continue

            # No adjacent image — just align the caption (the image may sit
            # further away, but GOST layout still wants a centered caption).
            if _format_figure_caption(para):
                changed = True
                fig_formatted += 1
            continue

        if _TABLE_CAPTION_RE.match(text):
            # Table caption: should be right above a table (w:tbl).
            nxt = _next_nonempty_sibling(p_elem)
            if nxt is not None and nxt.tag == qn("w:tbl"):
                if _format_table_caption(para):
                    changed = True
                    tbl_formatted += 1
                continue

            prev = _prev_nonempty_sibling(p_elem)
            if prev is not None and prev.tag == qn("w:tbl"):
                parent = p_elem.getparent()
                if parent is not None:
                    parent.remove(p_elem)
                    prev.addprevious(p_elem)
                    changed = True
                    tbl_moved += 1
                if _format_table_caption(para):
                    changed = True
                continue

            if _format_table_caption(para):
                changed = True
                tbl_formatted += 1
            continue

    if fig_moved:
        details.append(f"Подписи: {fig_moved} подпись(ей) «Рисунок» перемещено под изображение")
    if fig_formatted:
        details.append(f"Подписи: {fig_formatted} подпись(ей) «Рисунок» выровнено по центру")
    if tbl_moved:
        details.append(f"Подписи: {tbl_moved} подпись(ей) «Таблица» перемещено над таблицу")
    if tbl_formatted:
        details.append(f"Подписи: {tbl_formatted} подпись(ей) «Таблица» выровнено по левому краю")
    return changed


def _zero_paragraph_spacing(p_elem) -> bool:
    """Force ``space_before=0`` / ``space_after=0`` and ``line=auto`` (1.0)
    on *p_elem*'s ``<w:spacing>``. Returns True if anything changed.

    Used to compress the visual block around a figure / table so the
    caption sits flush with the image / data and does not bleed into
    the surrounding body text.
    """
    from docx.oxml import OxmlElement
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_elem.insert(0, pPr)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    changed = False
    if spacing.get(qn("w:before")) not in (None, "0"):
        spacing.set(qn("w:before"), "0")
        changed = True
    if spacing.get(qn("w:after")) not in (None, "0"):
        spacing.set(qn("w:after"), "0")
        changed = True
    if spacing.get(qn("w:beforeAutospacing")) not in (None, "0"):
        spacing.set(qn("w:beforeAutospacing"), "0")
        changed = True
    if spacing.get(qn("w:afterAutospacing")) not in (None, "0"):
        spacing.set(qn("w:afterAutospacing"), "0")
        changed = True
    return changed


def _drop_trailing_period(para) -> bool:
    """Strip a single trailing «.» from the last text run.

    Mirrors :func:`fix_caption_trailing_dot` from autofix_helpers but
    applies to the «Источник:» style of caption that we detect here
    rather than the figure/table caption itself. Keeps the period when
    it's actually an ellipsis («…»/«...»).
    """
    text = para.text.rstrip()
    if not text.endswith(".") or text.endswith(".."):
        return False
    t_elements = list(para._element.iter(qn("w:t")))
    for t_el in reversed(t_elements):
        s = (t_el.text or "").rstrip()
        if s and s.endswith(".") and not s.endswith(".."):
            tail = (t_el.text or "")[len(s):]
            t_el.text = s[:-1] + tail
            return True
        if s:
            break
    return False


def fix_source_caption_lines(doc, details: list[str]) -> bool:
    """Justify «Источник: …» / «Примечание: …» / «Составлено …» captions
    that follow a table and strip any trailing «.».

    Customer requirement (МЭО reference document):
        * Lines such as «Источник: составлено автором …» that sit
          underneath a table must be aligned **по ширине** (justified)
          rather than centered like the table caption above the table.
        * The trailing period at the end of that line must be removed.
        * The block must be visually tight — no extra space below the
          table itself nor between the «Источник» line and the next
          body paragraph.
    """
    changed = False
    fixed_align = 0
    stripped_dot = 0

    for para in list(doc.paragraphs):
        text = (para.text or "").strip()
        if not text or not _SOURCE_LINE_RE.match(text):
            continue
        # Only treat it as a table source caption when the previous
        # non-empty sibling is a table (or the caption directly follows
        # one). Otherwise it might be a regular bibliography note that
        # we shouldn't touch.
        prev = _prev_nonempty_sibling(para._element)
        if prev is None or prev.tag != qn("w:tbl"):
            continue

        if para.alignment != WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            fixed_align += 1
            changed = True
        pf = para.paragraph_format
        if pf.first_line_indent is not None and int(pf.first_line_indent) != 0:
            pf.first_line_indent = Mm(0)
            changed = True
        if pf.left_indent is not None and int(pf.left_indent) != 0:
            pf.left_indent = Mm(0)
            changed = True

        if _drop_trailing_period(para):
            stripped_dot += 1
            changed = True

        if _zero_paragraph_spacing(para._element):
            changed = True

    if fixed_align:
        details.append(
            f"Подписи: {fixed_align} строка(и) «Источник» выровнено по ширине"
        )
    if stripped_dot:
        details.append(
            f"Подписи: убрана конечная точка в {stripped_dot} строке(ах) «Источник»"
        )
    return changed


def ensure_blank_after_caption_blocks(doc, details: list[str]) -> bool:
    """Insert (or trim) exactly one empty paragraph after every figure/table/
    «Источник:» block.

    Customer rule: «После таблицы и рисунков или фразы источник: один пробел».
    The earlier passes ``tighten_caption_block_layout`` and
    ``fix_source_caption_lines`` zero ``space_after`` on the figure
    caption / source row, which makes the next body paragraph stick
    flush to the bottom of the table. We compensate by ensuring a
    single empty ``<w:p>`` between the block and the next body
    paragraph.

    A "block end" is detected as one of:
        * a ``<w:tbl>`` not followed by an «Источник:» row (the table
          itself is the end);
        * a ``<w:p>`` whose text matches ``_SOURCE_LINE_RE`` and which
          immediately follows a table (the source row is the end);
        * a ``<w:p>`` whose text matches ``_FIGURE_CAPTION_RE`` and
          which sits below a paragraph with an inline image.
    """
    from docx.oxml import OxmlElement
    body = doc.element.body
    inserted = 0
    trimmed = 0

    def _is_text_para(el):
        if el is None or el.tag != qn("w:p"):
            return False
        txt = "".join((t.text or "") for t in el.iter(qn("w:t"))).strip()
        return bool(txt) or _para_has_image(el)

    def _is_blank_para(el):
        if el is None or el.tag != qn("w:p"):
            return False
        txt = "".join((t.text or "") for t in el.iter(qn("w:t"))).strip()
        return not txt and not _para_has_image(el)

    def _ends_a_block(el) -> bool:
        if el.tag == qn("w:tbl"):
            nxt = _next_nonempty_sibling(el)
            if nxt is None:
                return True
            if nxt.tag == qn("w:p"):
                txt = "".join((t.text or "") for t in nxt.iter(qn("w:t"))).strip()
                if _SOURCE_LINE_RE.match(txt):
                    return False
            return True
        if el.tag != qn("w:p"):
            return False
        txt = "".join((t.text or "") for t in el.iter(qn("w:t"))).strip()
        if not txt:
            return False
        if _SOURCE_LINE_RE.match(txt):
            prev = _prev_nonempty_sibling(el)
            return prev is not None and prev.tag == qn("w:tbl")
        if _FIGURE_CAPTION_RE.match(txt):
            prev = _prev_nonempty_sibling(el)
            return prev is not None and prev.tag == qn("w:p") and _para_has_image(prev)
        return False

    children = list(body)
    for el in children:
        parent = el.getparent()
        if parent is None:
            continue
        if not _ends_a_block(el):
            continue

        nxt = el.getnext()
        # Walk past existing empties; if there are 2+ blanks, drop the extras.
        empties: list = []
        cursor = nxt
        while cursor is not None and _is_blank_para(cursor):
            empties.append(cursor)
            cursor = cursor.getnext()

        # If next non-blank sibling is something other than body text
        # (e.g., another caption / table / sectPr / heading), don't add
        # spacing — the caller's other rules govern that boundary.
        if cursor is None:
            for extra in empties[1:]:
                parent.remove(extra)
                trimmed += 1
            continue

        # Heading / page-break-before paragraphs already create the
        # vertical break we want; no extra blank required.
        cursor_pPr = cursor.find(qn("w:pPr")) if cursor.tag == qn("w:p") else None
        starts_section = False
        if cursor_pPr is not None:
            if cursor_pPr.find(qn("w:pageBreakBefore")) is not None:
                starts_section = True
            ps = cursor_pPr.find(qn("w:pStyle"))
            if ps is not None and (ps.get(qn("w:val")) or "").startswith("Heading"):
                starts_section = True

        if starts_section:
            for extra in empties:
                parent.remove(extra)
                trimmed += 1
            continue

        # Need exactly one blank between block end and *cursor*.
        if not empties:
            new_p = OxmlElement("w:p")
            el.addnext(new_p)
            inserted += 1
        elif len(empties) > 1:
            for extra in empties[1:]:
                parent.remove(extra)
                trimmed += 1

    if inserted:
        details.append(
            f"Подписи: после таблиц/рисунков/«Источник:» добавлен пустой абзац ({inserted} шт.)"
        )
    if trimmed:
        details.append(
            f"Подписи: убраны лишние пустые абзацы после таблиц/рисунков ({trimmed} шт.)"
        )
    return bool(inserted or trimmed)


def tighten_caption_block_layout(doc, details: list[str]) -> bool:
    """Compress the vertical space around figures, captions and tables.

    For every paragraph that:
        * contains an inline ``<w:drawing>`` (figure body), or
        * matches ``_FIGURE_CAPTION_RE`` (figure caption «Рисунок N…»),
          or
        * matches ``_TABLE_CAPTION_RE`` (table caption «Таблица N…»),
    we hard-zero ``<w:spacing w:before w:after>`` and turn off Word's
    *autospacing* flags. The line spacing of the figure paragraph
    itself is also forced to single (``line="240"``, ``lineRule="auto"``)
    so the inline image doesn't get a 30 % padding from the inherited
    1.3 multiplier.

    The «Источник» caption gets the same treatment via
    :func:`fix_source_caption_lines`.
    """
    from docx.oxml import OxmlElement
    changed = False
    fig_paragraphs = 0
    capt_paragraphs = 0

    for para in list(doc.paragraphs):
        p_elem = para._element
        text = (para.text or "").strip()
        is_figure = _para_has_image(p_elem)
        is_caption = bool(text) and (
            _FIGURE_CAPTION_RE.match(text) or _TABLE_CAPTION_RE.match(text)
        )
        if not is_figure and not is_caption:
            continue
        touched = _zero_paragraph_spacing(p_elem)
        if is_figure:
            # Force single line spacing on the figure-bearing paragraph
            # so the inline drawing renders without the 1.3 multiplier
            # padding, which used to add a visible gap above the
            # caption.
            pPr = p_elem.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                p_elem.insert(0, pPr)
            spacing = pPr.find(qn("w:spacing"))
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                pPr.append(spacing)
            if spacing.get(qn("w:line")) != "240":
                spacing.set(qn("w:line"), "240")
                spacing.set(qn("w:lineRule"), "auto")
                touched = True
        if touched:
            changed = True
            if is_figure:
                fig_paragraphs += 1
            if is_caption:
                capt_paragraphs += 1

    if fig_paragraphs:
        details.append(
            f"Рисунки: убраны лишние отступы вокруг изображений ({fig_paragraphs} шт.)"
        )
    if capt_paragraphs:
        details.append(
            f"Подписи рисунков/таблиц: интервал до/после обнулён ({capt_paragraphs} шт.)"
        )
    return changed
