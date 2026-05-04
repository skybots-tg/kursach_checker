"""Normalize formatting of an existing «Содержание» heading and its TOC
entries (auto-TOC fields wrapped in ``<w:sdt>`` content controls included).

Split out from ``autofix_toc.py`` to keep both modules under the project's
500-line cap.
"""
from __future__ import annotations

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx.text.paragraph import Paragraph

from app.rules_engine.autofix_toc import (
    _TOC_HEADING_RE,
    _clear_bold_underline_in_paragraph,
    _normalize_existing_toc_heading,
)
from app.rules_engine.style_resolve import detect_toc_paragraph_indices


def _normalize_toc_entry_run_font(p_elem: OxmlElement, font_name: str, font_size_pt: float) -> bool:
    """Force every ``<w:r>`` in *p_elem* to use ``font_name`` / ``font_size_pt``
    (incl. complex script size) so TOC entries match body text instead of
    rendering in the theme/heading font.
    """
    changed = False
    size_half_pt = str(int(round(font_size_pt * 2)))
    for r_elem in p_elem.iter(qn("w:r")):
        rPr = r_elem.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            r_elem.insert(0, rPr)
            changed = True
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
            changed = True
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            if rFonts.get(qn(attr)) != font_name:
                rFonts.set(qn(attr), font_name)
                changed = True
        for theme_attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
            if rFonts.get(qn(theme_attr)) is not None:
                del rFonts.attrib[qn(theme_attr)]
                changed = True
        for sz_tag in ("w:sz", "w:szCs"):
            sz = rPr.find(qn(sz_tag))
            if sz is None:
                sz = OxmlElement(sz_tag)
                rPr.append(sz)
                changed = True
            if sz.get(qn("w:val")) != size_half_pt:
                sz.set(qn("w:val"), size_half_pt)
                changed = True
    return changed


def _normalize_toc_entry_paragraph_format(p_elem: OxmlElement, line_spacing: float) -> bool:
    """Set 1.5 (or whatever cfg specifies) line spacing on a TOC entry, and
    drop any explicit space-before/space-after that would push entries apart.
    """
    changed = False
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_elem.insert(0, pPr)
        changed = True

    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
        changed = True

    line_240 = str(int(round(line_spacing * 240)))
    if spacing.get(qn("w:line")) != line_240:
        spacing.set(qn("w:line"), line_240)
        changed = True
    if spacing.get(qn("w:lineRule")) != "auto":
        spacing.set(qn("w:lineRule"), "auto")
        changed = True
    if spacing.get(qn("w:before")) not in (None, "0"):
        spacing.set(qn("w:before"), "0")
        changed = True
    if spacing.get(qn("w:after")) not in (None, "0"):
        spacing.set(qn("w:after"), "0")
        changed = True
    return changed


def _iter_sdt_toc_paragraphs(doc):
    """Yield ``(Paragraph, is_first_in_sdt)`` for every ``<w:p>`` nested in
    a ``<w:sdt>/<w:sdtContent>`` content control under the document body.

    Auto-generated tables of contents in Word are typically wrapped in such
    an SDT, and python-docx's ``doc.paragraphs`` silently skips paragraphs
    nested inside SDT — so without this helper neither the «Содержание»
    heading nor any of the TOC entries would be visible to the formatting
    pass.
    """
    body = doc.element.body
    p_tag = qn("w:p")
    sdt_content_tag = qn("w:sdtContent")

    for sdt in body.iter(qn("w:sdt")):
        content = sdt.find(sdt_content_tag)
        if content is None:
            continue
        is_first = True
        for sub in content.iterchildren():
            if sub.tag != p_tag:
                continue
            yield Paragraph(sub, doc), is_first
            is_first = False


def _ensure_toc_styles_font(
    doc, font_name: str, font_size_pt: float, line_spacing: float | None,
) -> bool:
    """Make sure ``TOC1``/``TOC2``/``TOC3`` paragraph styles exist in
    ``styles.xml`` and inherit ``font_name`` / ``font_size_pt`` /
    ``line_spacing`` from explicit ``<w:rFonts>`` / ``<w:sz>`` /
    ``<w:spacing>`` properties (no theme attributes).

    Why: when our auto-TOC field is later refreshed by Word or
    LibreOffice the renderer creates fresh ``<w:r>`` runs for the page
    numbers and dot leaders. Those runs only carry minimal local
    formatting and inherit everything else from the linked TOCx style.
    If the TOCx style is absent or themed, the inherited font becomes
    Calibri / theme-default and the page numbers drift away from the
    body's «Times New Roman 14». Setting the styles up explicitly fixes
    that drift permanently.

    Returns ``True`` when ``styles.xml`` was modified.
    """
    if not font_name or not font_size_pt:
        return False
    styles_part = doc.part.styles_part if hasattr(doc.part, "styles_part") else None
    if styles_part is None:
        try:
            styles_part = doc.styles.element
        except AttributeError:
            return False
        styles_root = styles_part
    else:
        styles_root = styles_part.element

    size_half = str(int(round(font_size_pt * 2)))
    line_240 = str(int(round((line_spacing or 1.0) * 240)))
    changed = False

    for level in (1, 2, 3):
        style_id = f"TOC{level}"
        target_name = f"toc {level}"
        # Word documents that already contain an auto-TOC frequently
        # ship a pre-existing style with name="toc N" but a custom
        # styleId (e.g. "11"). The TOC paragraphs reference that
        # numeric id directly, so creating a parallel ``TOC1`` style
        # alone wouldn't take effect on refresh. We therefore
        # normalise EVERY style whose name matches ``toc N`` plus
        # ensure the canonical ``TOCx`` exists for newly inserted
        # entries.
        targets: list = []
        for st in styles_root.iter(qn("w:style")):
            sid = st.get(qn("w:styleId"))
            sname_el = st.find(qn("w:name"))
            sname_val = (sname_el.get(qn("w:val")) if sname_el is not None else "") or ""
            if sid == style_id or sname_val.lower() == target_name:
                targets.append(st)
        if not any(st.get(qn("w:styleId")) == style_id for st in targets):
            style_el = OxmlElement("w:style")
            style_el.set(qn("w:type"), "paragraph")
            style_el.set(qn("w:styleId"), style_id)
            name_el = OxmlElement("w:name")
            name_el.set(qn("w:val"), target_name)
            style_el.append(name_el)
            based_on = OxmlElement("w:basedOn")
            based_on.set(qn("w:val"), "Normal")
            style_el.append(based_on)
            next_el = OxmlElement("w:next")
            next_el.set(qn("w:val"), "Normal")
            style_el.append(next_el)
            uiPriority = OxmlElement("w:uiPriority")
            uiPriority.set(qn("w:val"), "39")
            style_el.append(uiPriority)
            style_el.append(OxmlElement("w:unhideWhenUsed"))
            styles_root.append(style_el)
            targets.append(style_el)
            changed = True

        for style_el in targets:
            if _apply_toc_style_props(style_el, font_name, size_half, line_240):
                changed = True

    return changed


def _apply_toc_style_props(style_el, font_name: str, size_half: str, line_240: str) -> bool:
    """Force pPr/rPr inside *style_el* to point to TNR / 14 pt / 1.0
    spacing without theme attributes. Returns True on changes."""
    changed = False
    pPr = style_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        style_el.append(pPr)
        changed = True
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
        changed = True
    if spacing.get(qn("w:line")) != line_240:
        spacing.set(qn("w:line"), line_240)
        changed = True
    if spacing.get(qn("w:lineRule")) != "auto":
        spacing.set(qn("w:lineRule"), "auto")
        changed = True
    if spacing.get(qn("w:before")) not in (None, "0"):
        spacing.set(qn("w:before"), "0")
        changed = True
    if spacing.get(qn("w:after")) not in (None, "0"):
        spacing.set(qn("w:after"), "0")
        changed = True

    rPr = style_el.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        style_el.append(rPr)
        changed = True
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
        changed = True
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        if rFonts.get(qn(attr)) != font_name:
            rFonts.set(qn(attr), font_name)
            changed = True
    for theme_attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        if rFonts.get(qn(theme_attr)) is not None:
            del rFonts.attrib[qn(theme_attr)]
            changed = True
    for sz_tag in ("w:sz", "w:szCs"):
        sz = rPr.find(qn(sz_tag))
        if sz is None:
            sz = OxmlElement(sz_tag)
            rPr.append(sz)
            changed = True
        if sz.get(qn("w:val")) != size_half:
            sz.set(qn("w:val"), size_half)
            changed = True
    # Drop bold inherited from the «Heading 1»-based template some
    # docs use for TOC levels — TOC entries must render in regular
    # weight per the requested layout.
    for tag in ("w:b", "w:bCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            el.set(qn("w:val"), "0")
            rPr.append(el)
            changed = True
        elif el.get(qn("w:val")) != "0":
            el.set(qn("w:val"), "0")
            changed = True

    return changed


def _normalize_toc_entry(p_elem: OxmlElement, font_name: str | None,
                         font_size_pt: float | None, line_spacing: float | None) -> bool:
    """Apply the full «plain body text» normalization to a single TOC entry
    paragraph: clear bold/underline, set body font/size, set line spacing.
    """
    changed = _clear_bold_underline_in_paragraph(p_elem)
    if font_name and font_size_pt:
        if _normalize_toc_entry_run_font(p_elem, font_name, font_size_pt):
            changed = True
    if line_spacing:
        if _normalize_toc_entry_paragraph_format(p_elem, line_spacing):
            changed = True
    return changed


import re as _re

_INTRO_RE = _re.compile(
    r"^\s*(?:введение|introduction|вступлени)",
    _re.IGNORECASE | _re.UNICODE,
)


def lock_toc_fields(doc, details: list[str]) -> bool:
    """Add ``w:fldLock="true"`` to every PRE-EXISTING TOC field's
    ``<w:fldChar w:fldCharType="begin"/>`` so Word won't auto-recalculate
    the TOC on open / on print and overwrite our pre-rendered non-bold
    entries with bold-from-source-heading runs.

    Auto-inserted TOC fields (the ones we build via
    :func:`insert_toc_field`) carry ``w:dirty="true"`` on their begin
    marker — those MUST be refreshed by Word on first open so the
    placeholder entries get real page numbers. We deliberately skip
    fields that are already marked dirty.

    The lock attribute is honoured by Word and LibreOffice; users can
    still refresh the TOC manually via *Update Field*, but the default
    "auto-update on save / on print" path is disabled, which is enough
    to preserve our formatting when the customer simply opens the file.
    """
    body = doc.element.body
    fld_tag = qn("w:fldChar")
    fld_type_attr = qn("w:fldCharType")
    fld_lock_attr = qn("w:fldLock")
    fld_dirty_attr = qn("w:dirty")
    instr_tag = qn("w:instrText")

    locked_count = 0
    field_stack: list[tuple[str, object]] = []
    for p_elem in body.iter(qn("w:p")):
        for elem in p_elem.iter():
            tag = elem.tag
            if tag == fld_tag:
                ftype = elem.get(fld_type_attr)
                if ftype == "begin":
                    field_stack.append(("unknown", elem))
                elif ftype == "separate":
                    if field_stack and field_stack[-1][0] == "unknown":
                        kind, fld_begin = field_stack[-1]
                        field_stack[-1] = ("other", fld_begin)
                elif ftype == "end" and field_stack:
                    field_stack.pop()
            elif tag == instr_tag:
                text = (elem.text or "").upper()
                if (
                    field_stack
                    and field_stack[-1][0] == "unknown"
                    and "TOC" in text
                ):
                    kind, fld_begin = field_stack[-1]
                    field_stack[-1] = ("TOC", fld_begin)
                    dirty_val = (fld_begin.get(fld_dirty_attr) or "").lower()
                    if dirty_val in {"true", "1", "on"}:
                        # Auto-inserted TOC \u2014 leave dirty so Word refreshes it
                        # and fills in real page numbers on open.
                        continue
                    if fld_begin.get(fld_lock_attr) != "true":
                        fld_begin.set(fld_lock_attr, "true")
                        locked_count += 1

    if locked_count:
        details.append(
            f"Оглавление: поля TOC заблокированы от автопересчёта ({locked_count})"
        )
        return True
    return False


def ensure_page_break_after_toc(doc, details: list[str]) -> bool:
    """Force a page break after the TOC when the next body paragraph
    is not «Введение».

    Customer rule: «После содержания чтобы никакого текста не было
    Сразу с новой страницы, если нет слова введение». Translates to:
        * If the document starts the next chapter with «Введение»,
          leave the layout as-is — that paragraph already carries the
          chapter-level page-break-before via ``enforce_chapter_page_breaks``.
        * If there is no «Введение» (e.g. реферат / Plet нева doc),
          push the next non-empty paragraph onto a fresh page so the
          TOC sits on its own sheet.
    """
    body = doc.element.body
    p_tag = qn("w:p")

    # Locate the LAST element belonging to the TOC. Two flavours: an SDT
    # wrapping ``<w:sdtContent>`` (auto-TOC field) or a manual TOC that
    # starts with a paragraph matching ``_TOC_HEADING_RE`` followed by
    # entry paragraphs. We anchor on the SDT first because that is what
    # ``insert_toc_field`` produces.
    toc_anchor = None
    for sdt in body.iter(qn("w:sdt")):
        sdt_pr = sdt.find(qn("w:sdtPr"))
        is_toc_sdt = False
        if sdt_pr is not None:
            for child in sdt_pr.iter():
                tag_lower = child.tag.lower()
                if tag_lower.endswith("}docparttypes") or tag_lower.endswith("}docpartlist") or tag_lower.endswith("}docparts"):
                    txt = (etree_text_of_descendants(child) or "").lower()
                    if "table of contents" in txt or "toc" in txt:
                        is_toc_sdt = True
                        break
        # Fallback: SDT whose first content paragraph matches «Содержание».
        if not is_toc_sdt:
            content = sdt.find(qn("w:sdtContent"))
            if content is not None:
                first_p = next((c for c in content.iterchildren() if c.tag == p_tag), None)
                if first_p is not None:
                    txt = "".join((t.text or "") for t in first_p.iter(qn("w:t"))).strip()
                    if _TOC_HEADING_RE.match(txt):
                        is_toc_sdt = True
        if is_toc_sdt:
            toc_anchor = sdt
            break

    if toc_anchor is None:
        # Inline TOC field: walk body paragraphs tracking <w:fldChar> begin/
        # separate/end markers. The TOC field can span MANY paragraphs;
        # the anchor we want is the paragraph carrying the closing
        # ``<w:fldChar w:fldCharType="end"/>``.
        fld_tag = qn("w:fldChar")
        instr_tag = qn("w:instrText")
        fld_type_attr = qn("w:fldCharType")
        field_stack: list[str] = []
        last_toc_end_para = None
        for p_elem in body.iterchildren(p_tag):
            for elem in p_elem.iter():
                if elem.tag == fld_tag:
                    ftype = elem.get(fld_type_attr)
                    if ftype == "begin":
                        field_stack.append("unknown")
                    elif ftype == "separate":
                        if field_stack and field_stack[-1] == "unknown":
                            field_stack[-1] = "other"
                    elif ftype == "end" and field_stack:
                        ended = field_stack.pop()
                        if ended == "TOC":
                            last_toc_end_para = p_elem
                elif elem.tag == instr_tag:
                    text = elem.text or ""
                    if (
                        field_stack
                        and field_stack[-1] == "unknown"
                        and "TOC" in text.upper()
                    ):
                        field_stack[-1] = "TOC"
        if last_toc_end_para is not None:
            toc_anchor = last_toc_end_para

    if toc_anchor is None:
        # Manual TOC: find «Содержание» / «Оглавление» heading and walk
        # forward through paragraphs that look like TOC entries.
        for el in body:
            if el.tag != p_tag:
                continue
            text = "".join((t.text or "") for t in el.iter(qn("w:t"))).strip()
            if _TOC_HEADING_RE.match(text):
                tail = el
                cur = el.getnext()
                while cur is not None and cur.tag == p_tag:
                    txt2 = "".join((t.text or "") for t in cur.iter(qn("w:t"))).strip()
                    if not txt2:
                        tail = cur
                        cur = cur.getnext()
                        continue
                    # Heuristic: TOC entries are short rows ending with a
                    # page number tail OR using a TOC-linked paragraph
                    # style (``TOC1``/``TOC2``/…). Stop at the first
                    # non-TOC-looking paragraph.
                    pPr2 = cur.find(qn("w:pPr"))
                    sty_val = ""
                    if pPr2 is not None:
                        ps2 = pPr2.find(qn("w:pStyle"))
                        if ps2 is not None:
                            sty_val = (ps2.get(qn("w:val")) or "").lower()
                    is_toc_style = sty_val.startswith("toc") or sty_val in {
                        "tableofcontents", "tocheading"
                    }
                    import re as _re2
                    has_page_tail = bool(_re2.search(
                        r"(?:\.{2,}|\t|\u2026|\s{3,})\s*\d{1,4}\s*$", txt2
                    ))
                    if is_toc_style or has_page_tail:
                        tail = cur
                        cur = cur.getnext()
                        continue
                    break
                toc_anchor = tail
                break

    if toc_anchor is None:
        return False

    nxt = toc_anchor.getnext()
    while nxt is not None and nxt.tag == p_tag:
        text = "".join((t.text or "") for t in nxt.iter(qn("w:t"))).strip()
        if text:
            break
        nxt = nxt.getnext()

    if nxt is None or nxt.tag != p_tag:
        return False

    text = "".join((t.text or "") for t in nxt.iter(qn("w:t"))).strip()
    if _INTRO_RE.match(text):
        # Body already starts with «Введение», nothing to do.
        return False

    # Ensure ``<w:pageBreakBefore/>`` is set on *nxt* so it starts a
    # fresh page. Add to ``<w:pPr>`` if missing.
    pPr = nxt.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        nxt.insert(0, pPr)
    if pPr.find(qn("w:pageBreakBefore")) is None:
        pPr.append(OxmlElement("w:pageBreakBefore"))
        details.append(
            "Оглавление: после содержания добавлена пустая страница (нет «Введение»)"
        )
        return True
    return False


def etree_text_of_descendants(el) -> str:
    out: list[str] = []
    for t in el.itertext():
        out.append(t)
    return "".join(out)


def remove_duplicate_toc_heading_inside_sdt(doc, details: list[str]) -> bool:
    """Remove extra «Содержание»/«Оглавление» line inside TOC SDT when the same
    heading already exists as the preceding body paragraph (common Word
    layout: centered heading + duplicate inside the TOC control, wrong font).
    """
    body = doc.element.body
    changed = False
    prev_el = None
    for el in body:
        if el.tag != qn("w:sdt"):
            prev_el = el
            continue
        content = el.find(qn("w:sdtContent"))
        if content is None or prev_el is None or prev_el.tag != qn("w:p"):
            prev_el = el
            continue
        first_p = None
        for sub in content.iterchildren():
            if sub.tag == qn("w:p"):
                first_p = sub
                break
        if first_p is None:
            prev_el = el
            continue
        inner = Paragraph(first_p, doc).text.strip()
        outer = Paragraph(prev_el, doc).text.strip()
        if _TOC_HEADING_RE.match(inner) and _TOC_HEADING_RE.match(outer):
            content.remove(first_p)
            changed = True
            details.append(
                "Оглавление: удалён дублирующий заголовок внутри поля оглавления"
            )
        prev_el = el
    return changed


def normalize_toc_heading_formatting(doc, details: list[str], *, cfg=None) -> bool:
    """Standalone pass: ensure the «Содержание»/«Оглавление» heading is
    centered/non-bold AND every TOC entry paragraph (auto field or manual)
    is rendered without bold/underline, with body font/size and line spacing.

    Runs even when ``generate_toc`` skips the document (because it already
    has a TOC field), so the client's requirement «убрать жирное выделение
    в содержании» is always applied to both the heading and the entries.

    Also descends into ``<w:sdt>`` content controls (auto-TOC fields stored
    by Word) — those paragraphs are invisible to ``doc.paragraphs`` and
    were the reason TOC entries kept appearing in the heading/theme font
    after autofix.
    """
    changed = remove_duplicate_toc_heading_inside_sdt(doc, details)
    processed_elems: set = set()

    font_name = getattr(cfg, "font_name", None) if cfg is not None else None
    font_size_pt = getattr(cfg, "font_size_pt", None) if cfg is not None else None
    line_spacing = getattr(cfg, "line_spacing", None) if cfg is not None else None

    # Make sure the linked TOC paragraph styles inherit the body font
    # so post-refresh page numbers and dot leaders stay in TNR 14.
    if font_name and font_size_pt:
        if _ensure_toc_styles_font(doc, font_name, font_size_pt, line_spacing):
            changed = True
            details.append(
                f"Стили: TOC1/TOC2/TOC3 → {font_name} {font_size_pt:.0f} пт"
            )

    heading_elem = None
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if _TOC_HEADING_RE.match(text):
            heading_elem = para._element
            if _normalize_existing_toc_heading(para, details):
                changed = True
            if font_name and font_size_pt:
                if _normalize_toc_entry_run_font(heading_elem, font_name, font_size_pt):
                    changed = True
            if line_spacing:
                if _normalize_toc_entry_paragraph_format(heading_elem, line_spacing):
                    changed = True
            break

    sdt_entries_changed = 0
    for sdt_para, is_first in _iter_sdt_toc_paragraphs(doc):
        elem = sdt_para._element
        if elem in processed_elems:
            continue
        processed_elems.add(elem)
        text = (sdt_para.text or "").strip()
        if is_first and _TOC_HEADING_RE.match(text):
            if elem is not heading_elem:
                heading_elem = elem
                if _normalize_existing_toc_heading(sdt_para, details):
                    changed = True
                if font_name and font_size_pt:
                    if _normalize_toc_entry_run_font(heading_elem, font_name, font_size_pt):
                        changed = True
                if line_spacing:
                    if _normalize_toc_entry_paragraph_format(heading_elem, line_spacing):
                        changed = True
            continue
        if _normalize_toc_entry(elem, font_name, font_size_pt, line_spacing):
            sdt_entries_changed += 1
            changed = True

    entry_paragraphs_changed = 0
    toc_indices = detect_toc_paragraph_indices(doc)
    if toc_indices:
        paragraphs = doc.paragraphs
        for idx in toc_indices:
            if idx >= len(paragraphs):
                continue
            p_elem = paragraphs[idx]._element
            if p_elem is heading_elem or p_elem in processed_elems:
                continue
            if _normalize_toc_entry(p_elem, font_name, font_size_pt, line_spacing):
                entry_paragraphs_changed += 1
                changed = True

    total_entries = entry_paragraphs_changed + sdt_entries_changed
    if total_entries:
        details.append(
            f"Оглавление: записи приведены к шрифту/интервалу основного текста ({total_entries} шт.)"
        )
    return changed
