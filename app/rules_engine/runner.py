from pathlib import Path

from docx import Document


def _mm_from_emu(value: int | None) -> float | None:
    if value is None:
        return None
    return round(value / 36000, 2)


def _severity_from_block(rules: dict, key: str, default: str = "warning") -> str:
    blocks = rules.get("blocks", [])
    for b in blocks:
        if b.get("key") == key:
            return b.get("severity", default)
    return default


def _add_finding(
    findings: list[dict],
    title: str,
    category: str,
    severity: str,
    expected: str,
    found: str,
    location: str,
    recommendation: str,
) -> None:
    findings.append(
        {
            "title": title,
            "category": category,
            "severity": severity,
            "expected": expected,
            "found": found,
            "location": location,
            "recommendation": recommendation,
            "auto_fixed": False,
        }
    )


async def run_document_checks(file_path: str, rules: dict) -> dict:
    path = Path(file_path)
    findings: list[dict] = []

    if not path.exists():
        _add_finding(
            findings,
            "Файл не найден",
            "file",
            "error",
            "Файл должен быть доступен",
            "Отсутствует",
            "input",
            "Повторно загрузите файл",
        )
        return {
            "summary": {"errors": 1, "warnings": 0, "fixed": 0, "size": 0},
            "findings": findings,
            "rules_meta": {"blocks_count": len(rules.get("blocks", []))},
        }

    ext = path.suffix.lower()
    size = path.stat().st_size
    if ext != ".docx":
        _add_finding(
            findings,
            "Формат файла",
            "file",
            "warning",
            "DOCX",
            ext,
            "input",
            "Для полной проверки используйте DOCX",
        )
    else:
        doc = Document(str(path))
        sec = doc.sections[0] if doc.sections else None

        if sec:
            left = _mm_from_emu(sec.left_margin)
            right = _mm_from_emu(sec.right_margin)
            top = _mm_from_emu(sec.top_margin)
            bottom = _mm_from_emu(sec.bottom_margin)
            severity = _severity_from_block(rules, "layout", "warning")

            def check_margin(name: str, expected: float, found: float | None) -> None:
                if found is None:
                    return
                if abs(found - expected) > 1:
                    _add_finding(
                        findings,
                        f"Поле страницы: {name}",
                        "layout",
                        severity,
                        f"{expected} мм",
                        f"{found} мм",
                        "раздел 1",
                        "Исправьте параметры страницы в макете документа",
                    )

            check_margin("левое", 30, left)
            check_margin("правое", 15, right)
            check_margin("верхнее", 20, top)
            check_margin("нижнее", 25, bottom)

        typo_severity = _severity_from_block(rules, "typography", "warning")
        sample_checked = 0
        for p in doc.paragraphs:
            text = (p.text or "").strip()
            if not text:
                continue
            if p.style and "Heading" in p.style.name:
                continue
            for r in p.runs:
                if not r.text.strip():
                    continue
                sample_checked += 1
                if r.font.name and r.font.name.lower() != "times new roman":
                    _add_finding(
                        findings,
                        "Шрифт основного текста",
                        "typography",
                        typo_severity,
                        "Times New Roman",
                        r.font.name,
                        "абзац",
                        "Приведите основной текст к единому шрифту",
                    )
                    break
                if r.font.size and abs(r.font.size.pt - 14) > 0.2:
                    _add_finding(
                        findings,
                        "Размер шрифта",
                        "typography",
                        typo_severity,
                        "14 pt",
                        f"{round(r.font.size.pt, 2)} pt",
                        "абзац",
                        "Установите размер 14 pt для основного текста",
                    )
                    break
                break
            if sample_checked >= 40:
                break

        struct_severity = _severity_from_block(rules, "structure", "error")
        heading_titles = [p.text.strip().lower() for p in doc.paragraphs if p.style and "Heading" in p.style.name]
        for required in ["введение", "заключение"]:
            if required not in heading_titles:
                _add_finding(
                    findings,
                    f"Раздел «{required.capitalize()}»",
                    "structure",
                    struct_severity,
                    f"Раздел «{required.capitalize()}» должен присутствовать",
                    "Не найден",
                    "структура",
                    "Добавьте обязательный раздел в документ",
                )

    errors = sum(1 for f in findings if f["severity"] == "error")
    warnings = sum(1 for f in findings if f["severity"] == "warning")
    fixed = sum(1 for f in findings if f.get("auto_fixed") is True)

    if not findings:
        _add_finding(
            findings,
            "Проверка оформления",
            "summary",
            "advice",
            "Соответствие базовым правилам",
            "Нарушений не обнаружено",
            "документ",
            "Можно переходить к следующему этапу",
        )

    return {
        "summary": {"errors": errors, "warnings": warnings, "fixed": fixed, "size": size},
        "findings": findings,
        "rules_meta": {"blocks_count": len(rules.get("blocks", []))},
    }
