"""Auto-generate a TOC field when the document has none.

The field uses cached display populated from actual document headings so the
TOC is immediately visible.  When opened in Word, pressing Ctrl+A → F9
updates page numbers.
"""
from __future__ import annotations

import logging
import re

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.rules_engine.heading_detection import (
    TOC_LINE_TAIL_RE,
    KNOWN_SECTION_TITLES,
    detect_heading_candidate,
)
from app.rules_engine.style_resolve import detect_toc_paragraph_indices

logger = logging.getLogger(__name__)

_TOC_HEADING_RE = re.compile(r"^(содержание|оглавление)$", re.IGNORECASE)
_MAX_TOC_LEVEL = 3


def _has_auto_toc(doc) -> bool:
    body = doc.element.body
    for fld in body.iter(qn("w:fldChar")):
        if fld.get(qn("w:fldCharType")) != "begin":
            continue
        r_elem = fld.getparent()
        if r_elem is None:
            continue
        p_elem = r_elem.getparent()
        if p_elem is None:
            continue
        for sib_r in p_elem.iter(qn("w:r")):
            instr = sib_r.find(qn("w:instrText"))
            if instr is not None and instr.text and "TOC" in instr.text.upper():
                return True
    for fld in body.iter(qn("w:fldSimple")):
        instr = fld.get(qn("w:instr")) or ""
        if "TOC" in instr.upper():
            return True
    return False


# ── heading extraction ────────────────────────────────────────────────

def _get_heading_level(para) -> int | None:
    """Return 1-based heading level or *None*."""
    sid = getattr(para.style, "style_id", "") or ""
    for i in range(1, 10):
        if sid == f"Heading{i}":
            return i
    sname = (getattr(para.style, "name", "") or "").lower()
    for i in range(1, 10):
        if f"heading {i}" in sname or f"заголовок {i}" in sname:
            return i
    pPr = para._element.find(qn("w:pPr"))
    if pPr is not None:
        ol = pPr.find(qn("w:outlineLvl"))
        if ol is not None:
            try:
                val = int(ol.get(qn("w:val"), "9"))
                if val < 9:
                    return val + 1
            except (TypeError, ValueError):
                pass
    return None


_APPENDIX_PARENT_RE = re.compile(
    r"^приложени[еяй]\s*[:.\u2014\u2013-]?\s*$",
    re.IGNORECASE,
)
_APPENDIX_CHILD_RE = re.compile(
    r"^приложени[еяй]\s+[\dА-ЯЁA-Z][\)\.\s\u2014\u2013-]?",
    re.IGNORECASE,
)


def _filter_appendix_children_after_parent(
    headings: list[tuple[str, int]],
) -> list[tuple[str, int]]:
    """Drop every «Приложение N» entry that follows a plural «Приложения»
    parent so the table of contents shows the single canonical
    «ПРИЛОЖЕНИЯ» line and not a numbered list under it.

    The customer-approved layout collapses the appendix block into one
    TOC entry. This filter is the last line of defence: it runs even
    when the body-side ``consolidate_appendix_block`` pass could not
    demote the children (e.g. they were already real Heading 1
    paragraphs with custom formatting we did not want to touch).
    """
    seen_parent = False
    out: list[tuple[str, int]] = []
    for text, level in headings:
        stripped = text.strip()
        if _APPENDIX_PARENT_RE.match(stripped):
            seen_parent = True
            out.append((text, level))
            continue
        if seen_parent and _APPENDIX_CHILD_RE.match(stripped):
            # Suppressed: collapsed into the «ПРИЛОЖЕНИЯ» entry above.
            continue
        out.append((text, level))
    return out


def _collect_headings(doc) -> list[tuple[str, int]]:
    """Return ``[(text, level), ...]`` for headings up to ``_MAX_TOC_LEVEL``."""
    headings: list[tuple[str, int]] = []
    for para in doc.paragraphs:
        level = _get_heading_level(para)
        if level is None or level > _MAX_TOC_LEVEL:
            continue
        text = (para.text or "").strip()
        if not text or _TOC_HEADING_RE.match(text):
            continue
        headings.append((text, level))
    return _filter_appendix_children_after_parent(headings)


# ── XML builders ──────────────────────────────────────────────────────

def _append_no_bold_rpr(parent: OxmlElement) -> None:
    """Append <w:rPr><w:b w:val="0"/><w:bCs w:val="0"/><w:u w:val="none"/></w:rPr>."""
    rPr = OxmlElement("w:rPr")
    b = OxmlElement("w:b")
    b.set(qn("w:val"), "0")
    rPr.append(b)
    bCs = OxmlElement("w:bCs")
    bCs.set(qn("w:val"), "0")
    rPr.append(bCs)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "none")
    rPr.append(u)
    parent.append(rPr)


def _force_no_bold_no_underline(rPr: OxmlElement) -> bool:
    """Ensure rPr explicitly disables bold/bCs/underline. Return True if changed."""
    changed = False
    for tag in ("w:b", "w:bCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
            changed = True
        if el.get(qn("w:val")) != "0":
            el.set(qn("w:val"), "0")
            changed = True
    u = rPr.find(qn("w:u"))
    if u is None:
        u = OxmlElement("w:u")
        rPr.append(u)
        changed = True
    if u.get(qn("w:val")) != "none":
        u.set(qn("w:val"), "none")
        changed = True
    return changed


def _clear_bold_underline_in_paragraph(p_elem: OxmlElement) -> bool:
    """Force every <w:r> in *p_elem* (incl. inside <w:hyperlink>) to non-bold,
    non-underlined. Returns True if anything changed."""
    changed = False
    for r_elem in p_elem.iter(qn("w:r")):
        rPr = r_elem.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            r_elem.insert(0, rPr)
            changed = True
        if _force_no_bold_no_underline(rPr):
            changed = True
    return changed


def _build_toc_entry(text: str, level: int) -> OxmlElement:
    """Build a cached TOC entry paragraph with TOCx style, non-bold, no indent."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), f"TOC{level}")
    pPr.append(pStyle)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:firstLine"), "0")
    pPr.append(ind)
    _append_no_bold_rpr(pPr)
    p.append(pPr)
    r = OxmlElement("w:r")
    _append_no_bold_rpr(r)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def _make_fld_run(fld_type: str, *, dirty: bool = False) -> OxmlElement:
    r = OxmlElement("w:r")
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), fld_type)
    if dirty and fld_type == "begin":
        # ``w:dirty="true"`` tells Word "please refresh this field on
        # the next open" — used for AUTO-INSERTED TOC fields that
        # carry placeholder text instead of real page numbers. We do
        # NOT add this to pre-existing TOC fields (those already have
        # cached page numbers and we lock them via ``lock_toc_fields``
        # to prevent Word from copying run-level bold from headings
        # back into the TOC entries).
        fld.set(qn("w:dirty"), "true")
    r.append(fld)
    return r


def _make_instr_run(instr_text: str) -> OxmlElement:
    r = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    r.append(instr)
    return r


def _prepend_runs_after_pPr(p_elem: OxmlElement, runs: list[OxmlElement]) -> None:
    """Insert *runs* right after ``<w:pPr>`` (or at index 0) preserving order."""
    pPr = p_elem.find(qn("w:pPr"))
    insert_idx = (list(p_elem).index(pPr) + 1) if pPr is not None else 0
    for offset, r in enumerate(runs):
        p_elem.insert(insert_idx + offset, r)


def _build_toc_elements(
    headings: list[tuple[str, int]],
) -> list[OxmlElement]:
    """Return TOC paragraphs with field markers folded into the entries.

    Word allows the begin / separate / end markers of a TOC field to live
    inside the same paragraphs as the TOC entries. We deliberately do NOT
    emit standalone empty paragraphs for the markers because such "ghost"
    paragraphs are not removable by the empty-paragraph cleanup (they
    contain ``<w:fldChar>``, not ``<w:t>``) and frequently get pushed
    onto a page of their own when the TOC fills the first page,
    producing a blank page right after the table of contents.
    """
    entries = [_build_toc_entry(text, level) for text, level in headings]

    if not entries:
        fallback = OxmlElement("w:p")
        r_fb = OxmlElement("w:r")
        t_fb = OxmlElement("w:t")
        t_fb.text = "Обновите оглавление (Ctrl+A, затем F9)"
        r_fb.append(t_fb)
        fallback.append(r_fb)
        entries = [fallback]

    # Field begin / instr / separate go BEFORE the first entry's text run,
    # but after its <w:pPr>, so the entry text still renders correctly.
    _prepend_runs_after_pPr(
        entries[0],
        [
            _make_fld_run("begin", dirty=True),
            _make_instr_run(r' TOC \o "1-3" \h \z \u '),
            _make_fld_run("separate"),
        ],
    )

    # Field end goes at the very end of the last entry.
    entries[-1].append(_make_fld_run("end"))

    return entries


def _build_toc_heading_paragraph(text: str = "Содержание") -> OxmlElement:
    """Build a plain centered TOC heading paragraph (no bold, no underline).

    Matches the format requested by the user: single word "Содержание"
    centered on the line, regular (non-bold) font.
    """
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLine"), "0")
    ind.set(qn("w:left"), "0")
    pPr.append(ind)
    _append_no_bold_rpr(pPr)
    p.append(pPr)
    r = OxmlElement("w:r")
    _append_no_bold_rpr(r)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def _has_visible_content_before(para_elem) -> bool:
    """True iff there's any non-empty paragraph, table or manual page-break
    before *para_elem* in the document body.

    Used to decide whether «Содержание» should force a new page: when the
    document has a title page above the TOC, we want ``pageBreakBefore``
    on the heading; if «Содержание» is the very first paragraph, adding a
    break would create a blank first page.
    """
    cur = para_elem.getprevious()
    while cur is not None:
        tag = cur.tag
        if tag == qn("w:tbl"):
            return True
        if tag == qn("w:p"):
            text = "".join((t.text or "") for t in cur.iter(qn("w:t")))
            if text.strip():
                return True
            for br in cur.iter(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    return True
        cur = cur.getprevious()
    return False


def _ensure_page_break_before_toc_heading(para, details: list[str]) -> bool:
    """Force ``page_break_before`` on «Содержание»/«Оглавление» when there's
    a title page (or any other content) above it.

    Required because ``normalize_title_page_spacing`` strips manual
    ``<w:br w:type="page"/>`` between the title page and the TOC heading,
    and the heading itself is later detached from ``Heading 1`` style by
    :func:`_normalize_existing_toc_heading` — so neither path leaves the
    TOC on a fresh page without an explicit pageBreakBefore here.
    """
    if not _has_visible_content_before(para._element):
        return False
    pf = para.paragraph_format
    if pf.page_break_before:
        return False
    pf.page_break_before = True
    details.append(
        "Оглавление: «Содержание» вынесено на новую страницу"
    )
    return True


def _normalize_existing_toc_heading(para, details: list[str]) -> bool:
    """Re-format an existing «Содержание» paragraph: center + non-bold + no underline.

    The client explicitly wants «Содержание» to look like a plain regular
    paragraph centered on the line (see the reference photo). If the source
    document styled it as ``Heading N`` (or our earlier passes promoted it
    there), the heading style cascades bold onto every run. To ensure the
    final rendering matches the requirement we:
      1. Detach the paragraph from ``Heading N`` by switching to ``Normal``.
      2. Explicitly set ``<w:b w:val="0"/>`` / ``<w:bCs w:val="0"/>`` /
         ``<w:u w:val="none"/>`` on every run (inserting the element when
         missing) so no inherited property can re-enable bold/underline.
      3. Center alignment and wipe any red-line / left indent.
      4. Force ``pageBreakBefore`` so the TOC always starts on its own page,
         even after the title-page page-break stripper runs.
    """
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Mm

    changed = False

    doc = para.part.document
    try:
        normal = doc.styles["Normal"]
        current_sid = getattr(para.style, "style_id", "") or ""
        if current_sid != getattr(normal, "style_id", ""):
            para.style = normal
            changed = True
    except KeyError:
        pass

    if para.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        changed = True

    pf = para.paragraph_format
    if pf.first_line_indent is not None and int(pf.first_line_indent) != 0:
        pf.first_line_indent = Mm(0)
        changed = True
    if pf.left_indent is not None and int(pf.left_indent) != 0:
        pf.left_indent = Mm(0)
        changed = True

    pPr = para._element.find(qn("w:pPr"))
    if pPr is not None:
        pPr_rPr = pPr.find(qn("w:rPr"))
        if pPr_rPr is not None and _force_no_bold_no_underline(pPr_rPr):
            changed = True

    if _clear_bold_underline_in_paragraph(para._element):
        changed = True

    if _ensure_page_break_before_toc_heading(para, details):
        changed = True

    if changed:
        details.append("Оглавление: заголовок «Содержание» оформлен по центру без жирного")
    return changed


_TOC_MAX_ENTRIES = 60
_TOC_MAX_TRAILING_BLANKS = 4
# «Глава N», «Раздел N», roman numerals, section numbers, appendix markers —
# typical TOC line starts even when the author didn't type page numbers.
_TOC_ENTRY_START_RE = re.compile(
    r"^(?:"
    r"глава\s+\d"
    r"|раздел\s+\d"
    r"|chapter\s+\d"
    r"|\d+\s+глава"
    r"|\d{1,2}\.\d{1,2}(?:\.\d{1,2}){0,2}\.?\s"
    r"|[IVXLCM]{1,6}\s+[А-ЯЁA-Zа-яёa-z]"
    r"|приложени[еяй]\s"
    r")",
    re.IGNORECASE,
)


def _looks_like_toc_line(text: str) -> bool:
    """Heuristic: does *text* look like a manual TOC entry?

    Accepts three families of lines:
      1) entries with an explicit page number tail (``\\s{2,}\\d`` / dots+num /
         tab+num);
      2) entries that start with a recognizable heading marker — chapter
         keyword, section number «1.2», roman numeral chapter, «Приложение»;
      3) short lines matching a well-known section title
         («Введение», «Заключение», «Список литературы», …).
    """
    stripped = text.strip()
    if not stripped or len(stripped) > 200:
        return False

    if TOC_LINE_TAIL_RE.search(stripped):
        return True
    if re.search(r"\.{2,}\s*\d+\s*$", stripped):
        return True
    if re.search(r"\t+\d+\s*$", stripped):
        return True
    if _TOC_ENTRY_START_RE.match(stripped):
        return True

    low = stripped.lower().rstrip(":.;")
    if low in KNOWN_SECTION_TITLES:
        return True
    # Known section title followed by a short inline suffix, e.g.
    # «Список использованных источников информации».
    for title in KNOWN_SECTION_TITLES:
        if low.startswith(title + " ") and len(low) <= len(title) + 40:
            return True

    return False


def _is_strong_toc_break(para) -> bool:
    """Return True when *para* is clearly the start of real body content, so
    the TOC block must stop at this paragraph.
    """
    style_name = (getattr(para.style, "name", "") or "").lower()
    style_id = (getattr(para.style, "style_id", "") or "")
    if "heading" in style_name or "заголов" in style_name:
        return True
    if style_id.startswith("Heading"):
        return True
    pf = para.paragraph_format
    if pf.page_break_before:
        return True
    return False


def _normalize_for_dup(text: str) -> str:
    """Collapse whitespace + lowercase + strip trailing punctuation/numbers
    so that ``"Введение"`` and the body heading ``"ВВЕДЕНИЕ"`` (or
    ``"Введение  3"``) compare equal. Used to detect the moment when the
    manual TOC block ends and the body starts repeating its own titles."""
    base = re.sub(r"\s+", " ", text).strip().lower()
    base = re.sub(r"[\.\u2026]+\s*\d+\s*$", "", base)  # «… 12»
    base = re.sub(r"\s{2,}\d+\s*$", "", base)         # «  12»
    base = re.sub(r"\t+\d+\s*$", "", base)
    base = base.rstrip(":.;\u2014\u2013- ")
    return base


def _scan_manual_toc_block(
    paragraphs, heading_idx: int,
) -> tuple[list[int], list[object]]:
    """Walk the paragraphs after the TOC heading and collect TOC entries.

    Stops when:
      * a paragraph styled as a real heading or with a page break is met
        (``_is_strong_toc_break``), or
      * a non-TOC-looking paragraph appears, or
      * the same normalized title has been seen earlier in the block —
        this is the typical sign that the body has started repeating the
        section name (``"Введение"`` in the TOC vs ``"ВВЕДЕНИЕ"`` in the
        body), or
      * we exhausted ``_TOC_MAX_ENTRIES``.

    Returns (indices, elements) — the indices include the heading itself.
    """
    indices: list[int] = [heading_idx]
    elements: list[object] = []
    seen_norm: set[str] = set()
    blank_streak = 0
    entries_seen = 0

    for offset in range(heading_idx + 1, len(paragraphs)):
        para = paragraphs[offset]
        if entries_seen >= _TOC_MAX_ENTRIES:
            break
        text = (para.text or "").strip()
        if not text:
            blank_streak += 1
            if blank_streak > _TOC_MAX_TRAILING_BLANKS and entries_seen == 0:
                break
            indices.append(offset)
            elements.append(para._element)
            continue
        if _is_strong_toc_break(para):
            break
        if not _looks_like_toc_line(text):
            break
        norm = _normalize_for_dup(text)
        if norm and norm in seen_norm:
            # Same title appeared earlier in the TOC list → this is the
            # body section repeating its name. Stop here without absorbing.
            break
        if norm:
            seen_norm.add(norm)
        indices.append(offset)
        elements.append(para._element)
        blank_streak = 0
        entries_seen += 1

    if not elements:
        return [heading_idx], []

    # Trim purely-blank tail entries we may have accumulated.
    while elements and not (elements[-1].itertext() and any(
        (t or "").strip() for t in elements[-1].itertext()
    )):
        elements.pop()
        indices.pop()

    return indices, elements


def detect_manual_toc_entry_indices(doc) -> set[int]:
    """Return paragraph indices of a manual (plain-text) TOC block.

    Use this alongside ``detect_toc_paragraph_indices`` so passes that
    normally skip TOCs don't accidentally touch manual TOC text (e.g. by
    promoting a TOC line to a ``Heading N`` style, which would later
    prevent the TOC block from being removed before inserting the field).
    """
    paragraphs = doc.paragraphs
    for idx, p in enumerate(paragraphs):
        text = (p.text or "").strip()
        if _TOC_HEADING_RE.match(text):
            indices, _ = _scan_manual_toc_block(paragraphs, idx)
            return set(indices)
    return set()


_TOC_TABLE_MIN_HITS = 3


def _table_looks_like_toc(table) -> bool:
    """Heuristic: does *table* contain a manual TOC?

    A TOC table typically consists of two columns where the left column
    holds section titles (``Введение``, ``Глава 1.…``, ``1.1.…``,
    ``Заключение``…) and the right column holds page numbers. We mark
    the table as a TOC when at least :data:`_TOC_TABLE_MIN_HITS` rows
    contain a recognisable TOC entry (heading marker, well-known
    section title or short numeric tail).
    """
    rows = list(table.rows)
    if not rows or len(rows) > 80:
        return False

    hits = 0
    for row in rows:
        cells = list(row.cells)
        if not cells:
            continue
        left_text = (cells[0].text or "").strip()
        if not left_text:
            continue
        if _looks_like_toc_line(left_text):
            hits += 1
        if hits >= _TOC_TABLE_MIN_HITS:
            return True
    return False


def _remove_table_after_heading(doc, heading_para, details: list[str]) -> bool:
    """If the next sibling element after *heading_para* is a TOC table,
    detach it and drop any empty paragraphs that sat between the heading
    and the table.

    Returns True if a TOC table was removed.
    """
    body = doc.element.body
    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")

    sibling = heading_para._element.getnext()
    interim_blanks: list = []
    while sibling is not None:
        if sibling.tag == tbl_tag:
            break
        if sibling.tag == p_tag:
            text = "".join(
                (t.text or "") for t in sibling.iter(qn("w:t"))
            ).strip()
            if text:
                return False
            interim_blanks.append(sibling)
            sibling = sibling.getnext()
            continue
        return False

    if sibling is None or sibling.tag != tbl_tag:
        return False

    target_table = None
    for table in doc.tables:
        if table._element is sibling:
            target_table = table
            break
    if target_table is None or not _table_looks_like_toc(target_table):
        return False

    for blank in interim_blanks:
        if blank.getparent() is not None:
            blank.getparent().remove(blank)

    sibling.getparent().remove(sibling)
    details.append(
        "Оглавление: удалена таблица с ручным оглавлением"
    )
    return True


def _remove_manual_toc_entries(doc, heading_idx: int, details: list[str]) -> bool:
    """Remove manually typed TOC lines that follow the TOC heading.

    A single student-written TOC frequently contains a mix of entry styles:
    the first few lines have page numbers (``Введение  3-4``), while the
    subsequent subsection lines do not (``1.1. Определение и классификация
    заболевания``). The previous implementation stopped at the first entry
    without a number tail and left the rest of the TOC behind — the final
    document then ended up with two tables of contents glued together.

    This version treats any paragraph that *looks like a TOC line* (entry
    start marker, numbered subsection, appendix, well-known section title
    or number-tail) as still belonging to the TOC block. Up to four blank
    or inline paragraphs in a row are also absorbed so stray empty lines
    inside the block don't interrupt it. The block ends as soon as we see
    a real heading, a page-break-before paragraph, or a paragraph that
    clearly belongs to the body.
    """
    paragraphs = doc.paragraphs
    _, to_remove = _scan_manual_toc_block(paragraphs, heading_idx)

    for elem in to_remove:
        if elem.getparent() is not None:
            elem.getparent().remove(elem)

    if to_remove:
        details.append(f"Оглавление: удалено {len(to_remove)} ручных записей содержания")
    return len(to_remove) > 0


_RESIDUAL_TOC_BULLET_RE = re.compile(
    r"^[\s\xa0]*[\-\u2013\u2014\u2022\u00b7][\s\xa0]+",
    re.UNICODE,
)
_RESIDUAL_TOC_KNOWN_TITLE_RE = re.compile(
    r"^[\s\xa0]*(?:список\s+(?:использован|используем|литератур|источник|терминов|сокращ)"
    r"|библиограф|приложени[еяй]|список\s+терминов|список\s+сокращ"
    r"|термин(?:ы|ов)?|сокращ(?:ения|ений)?|глоссарий|abbreviation|terminolog)",
    re.IGNORECASE | re.UNICODE,
)
# Heuristic: short Title-cased paragraph that ends with «.» — almost
# certainly a one-word TOC entry leftover («Термины.», «Глоссарий.»,
# «Список сокращений.»). Body paragraphs are virtually never that short.
_RESIDUAL_TOC_SHORT_TITLE_RE = re.compile(
    r"^[\s\xa0]*[А-ЯЁA-Z][А-ЯЁа-яёA-Za-z\s\xa0\-]{1,40}\.?[\s\xa0]*$",
    re.UNICODE,
)
_BODY_CHAPTER_DECOR_RE = re.compile(
    r"^[\s\xa0]*[\u2580-\u259f]",  # block-element symbols ▌▍▎▏█▓▒░◀◆… that
    re.UNICODE,                     # students use to mark body chapter starts
)


def _looks_like_residual_toc_line(text: str) -> bool:
    """More permissive than ``_looks_like_toc_line`` — used for cleaning
    leftover TOC garbage that appears between the inserted auto-TOC
    field and the first real body chapter.

    Catches three shapes that ``_scan_manual_toc_block`` is too strict to
    absorb: bullet-prefixed continuations («– Значение …»),
    short well-known section titles even without a page number tail
    («Список терминов и сокращений 26-30», «\xa0Список использованной
    литературы 31»), and lines that mostly look like a TOC entry tail.
    """
    if not text:
        return True
    if len(text) > 200:
        return False
    if _looks_like_toc_line(text):
        return True
    if _RESIDUAL_TOC_BULLET_RE.match(text):
        # Bullet continuation — only when the rest of the line is short
        # (TOC tail, not a normal listed paragraph in body text).
        return len(text) <= 120
    if _RESIDUAL_TOC_KNOWN_TITLE_RE.match(text):
        return True
    if len(text) <= 40 and _RESIDUAL_TOC_SHORT_TITLE_RE.match(text):
        # Single short Title-cased fragment with optional trailing period —
        # treat as TOC residue ONLY when followed by no body content
        # (the caller already enforces this by stopping at the first
        # real heading / long body paragraph).
        return True
    return False


def _remove_residual_toc_after_heading(
    doc, heading_para, details: list[str], max_scan: int = 25,
) -> int:
    """Walk forward from the TOC heading and drop any paragraph that
    still looks like a manual TOC entry (page-number tail, bullet
    continuation, short known section title) up until the first real
    body chapter.

    The first paragraph that *clearly* belongs to the body — a real
    Heading style, a paragraph that starts with a chapter-decoration
    block character (▌, ◀, …) typical of student manuscripts, or a long
    body paragraph (> 200 chars without TOC traits) — terminates the
    scan and stays untouched.

    Returns the number of paragraphs removed.
    """
    body = doc.element.body
    heading_elem = heading_para._element if hasattr(heading_para, "_element") else heading_para
    if heading_elem.getparent() is not body:
        return 0

    siblings = list(body)
    try:
        start_pos = siblings.index(heading_elem) + 1
    except ValueError:
        return 0

    removed = 0
    scanned = 0
    pos = start_pos
    while pos < len(body) and scanned < max_scan:
        el = body[pos]
        if el.tag != qn("w:p"):
            # Hit a table or sectPr — stop, body content begins.
            break
        text = "".join((t.text or "") for t in el.iter(qn("w:t"))).strip()
        # Skip fldChar paragraphs that belong to the just-inserted TOC
        # field (they may be empty text-wise but carry the field).
        has_field = any(True for _ in el.iter(qn("w:fldChar")))
        if has_field:
            pos += 1
            scanned += 1
            continue
        # A real Word heading or a chapter-decoration body line ends the
        # cleanup window.
        pPr = el.find(qn("w:pPr"))
        sval = ""
        if pPr is not None:
            ps = pPr.find(qn("w:pStyle"))
            if ps is not None:
                sval = ps.get(qn("w:val")) or ""
        if sval and sval.startswith("Heading"):
            break
        if _BODY_CHAPTER_DECOR_RE.match(text):
            break
        if not text:
            pos += 1
            scanned += 1
            continue
        if _looks_like_residual_toc_line(text):
            body.remove(el)
            removed += 1
            # Don't advance pos — list shifted.
            continue
        # Real body paragraph reached.
        break

    if removed:
        details.append(
            f"Оглавление: удалено {removed} остаточных строк ручного содержания"
        )
    return removed


def _insert_toc_after(anchor, doc, details: list[str]) -> bool:
    """Insert a multi-paragraph TOC field right after *anchor* element."""
    headings = _collect_headings(doc)
    entries = _build_toc_elements(headings)

    last = anchor
    for elem in entries:
        last.addnext(elem)
        last = elem

    n_entries = len(headings)
    if n_entries:
        details.append(
            f"Оглавление: вставлено поле TOC ({n_entries} записей из заголовков)"
        )
    else:
        details.append("Оглавление: вставлено поле TOC (обновите в Word: Ctrl+A → F9)")
    return True


def insert_toc_field(doc, toc_indices: set[int], details: list[str]) -> bool:
    if _has_auto_toc(doc):
        return False

    body = doc.element.body
    paragraphs = doc.paragraphs

    for idx, para in enumerate(paragraphs):
        text = (para.text or "").strip()
        if _TOC_HEADING_RE.match(text):
            _remove_manual_toc_entries(doc, idx, details)
            _remove_table_after_heading(doc, para, details)
            _normalize_existing_toc_heading(para, details)
            inserted = _insert_toc_after(para._element, doc, details)
            # Find the LAST paragraph of the inserted TOC field so cleanup
            # starts AFTER the field. Intermediate TOC entries do not
            # carry their own ``<w:fldChar>`` markers — only the first
            # entry has ``begin/separate`` and the last entry has
            # ``end``. We therefore track the begin/end nesting depth
            # while walking siblings, advancing through every paragraph
            # until depth returns to zero (which marks the end of the
            # TOC field, regardless of how many entry paragraphs sit in
            # between).
            anchor_elem = para._element
            tail = anchor_elem
            depth = 0
            saw_begin = False
            cur = anchor_elem.getnext()
            while cur is not None and cur.tag == qn("w:p"):
                fld_types = [fc.get(qn("w:fldCharType")) for fc in cur.iter(qn("w:fldChar"))]
                for ft in fld_types:
                    if ft == "begin":
                        depth += 1
                        saw_begin = True
                    elif ft == "end":
                        depth -= 1
                tail = cur
                if saw_begin and depth <= 0:
                    break
                if not saw_begin:
                    # Auto-TOC field MUST start in the very first inserted
                    # paragraph; if we walked past it without finding a
                    # ``begin`` marker, something is off — bail out and let
                    # the cleanup pass start from the heading itself.
                    if cur.getnext() is None:
                        tail = anchor_elem
                        break
                cur = cur.getnext()
            from docx.text.paragraph import Paragraph as _P
            _remove_residual_toc_after_heading(doc, _P(tail, doc), details)
            return inserted

    for para in paragraphs:
        style_name = (getattr(para.style, "name", "") or "").lower()
        if "heading" in style_name or "заголов" in style_name:
            heading_p = _build_toc_heading_paragraph("Содержание")
            para._element.addprevious(heading_p)
            details.append("Оглавление: создан заголовок «Содержание» (по центру)")
            return _insert_toc_after(heading_p, doc, details)

    if len(body) > 0:
        heading_p = _build_toc_heading_paragraph("Содержание")
        body.insert(0, heading_p)
        details.append("Оглавление: создан заголовок «Содержание» (по центру)")
        return _insert_toc_after(heading_p, doc, details)

    return False

