"""Распределение вертикальных интервалов на титульном листе (блоки → город/год внизу)."""
from __future__ import annotations

import logging
import re

from docx.oxml.ns import qn
from docx.shared import Emu, Pt

logger = logging.getLogger(__name__)

_PERSONNEL_RE = re.compile(
    r"\bвыполнил[аи]?|\bпроверил",
    re.IGNORECASE,
)
_REFERAT_WORD_RE = re.compile(r"\bреферат\b", re.IGNORECASE)
_BLOCK1_STRONG_HINT = re.compile(
    r"министерство|федеральн|университет|институт|школа\s|департамент|кафедр|"
    r"\bдвфу\b|фгаоу|образовательн(?:ое|ого)\s+учрежден",
    re.IGNORECASE,
)
_YEAR_STRICT = re.compile(r"^\d{4}\s*$")
_CITY_YEAR_ONE = re.compile(
    r"^(?:г\.\s*)?([А-ЯЁа-яЁё][А-ЯЁа-яЁё\-\s\w]*)\s*,\s*(\d{4})\s*$",
    re.IGNORECASE,
)
_PERSONNEL_SATELLITE_RE = re.compile(
    r"\d{4}|«|»|_{3,}|апрел|январ|феврал|март|ма[яй]|июн|июл|август|сентяб|октяб|нояб|декаб",
    re.IGNORECASE,
)


def _collapse(text: str) -> str:
    return re.sub(r"[\n\r\t\xa0]+", " ", (text or "")).strip()


def _mostly_uppercase_ratio(s: str) -> float:
    letters = [c for c in s if c.isalpha()]
    if not letters:
        return 0.0
    return sum(1 for c in letters if c.isupper()) / len(letters)


def _is_topic_caps_line(c: str) -> bool:
    if not c or _BLOCK1_STRONG_HINT.search(c):
        return False
    if len(c) < 10 or len(c) > 190:
        return False
    if _REFERAT_WORD_RE.search(c):
        return False
    return _mostly_uppercase_ratio(c) >= 0.78


def _one_line_city_year(collapsed: str) -> bool:
    return _CITY_YEAR_ONE.match(collapsed.strip()) is not None


def _find_footer_city_idx(paras, body_start: int) -> int | None:
    for i in range(body_start - 1, -1, -1):
        raw = (paras[i].text or "").strip()
        if not raw:
            continue
        col = _collapse(paras[i].text or "")
        if _one_line_city_year(col):
            return i
        if _YEAR_STRICT.match(raw):
            j = i - 1
            while j >= 0 and not (paras[j].text or "").strip():
                j -= 1
            if j < 0:
                return None
            return j
    return None


def _find_personnel_range(paras, city_idx: int) -> tuple[int, int] | None:
    """Первый и последний абзацы блока исполнителя (строго выше city_idx)."""
    end = city_idx - 1
    while end >= 0 and not (paras[end].text or "").strip():
        end -= 1
    if end < 0:
        return None
    marker_at: int | None = None
    for u in range(end, -1, -1):
        c = _collapse(paras[u].text or "")
        if _PERSONNEL_RE.search(c):
            marker_at = u
            break
    if marker_at is None:
        return None
    start = marker_at
    while start - 1 >= 0:
        c = _collapse(paras[start - 1].text or "")
        if not c:
            start -= 1
            continue
        if _REFERAT_WORD_RE.search(c):
            break
        if _PERSONNEL_RE.search(c):
            start -= 1
            continue
        if len(c) < 130 and _PERSONNEL_SATELLITE_RE.search(c):
            start -= 1
            continue
        break
    return start, end


def _find_referat_span(paras, personnel_start: int) -> tuple[int, int] | None:
    referat_idx: int | None = None
    for k in range(personnel_start):
        t = (paras[k].text or "").strip()
        if not t:
            continue
        if _REFERAT_WORD_RE.search(_collapse(t)):
            referat_idx = k
            break
    if referat_idx is None:
        return None
    referat_start = referat_idx
    while referat_start - 1 >= 0:
        c = _collapse(paras[referat_start - 1].text or "")
        if not c:
            referat_start -= 1
            continue
        if _PERSONNEL_RE.search(c):
            break
        if _is_topic_caps_line(c):
            break
        if _BLOCK1_STRONG_HINT.search(c) and _mostly_uppercase_ratio(c) < 0.55:
            break
        low = c.lower()
        if any(
            x in low
            for x in (
                "образовательной",
                "направлению",
                "направлен",
                "программ",
                "бакалавр",
                "магистр",
                "специалитет",
            )
        ):
            referat_start -= 1
            continue
        break
    referat_end = referat_idx
    while referat_end + 1 < personnel_start:
        c = _collapse(paras[referat_end + 1].text or "")
        if not c:
            break
        if _PERSONNEL_RE.search(c):
            break
        low = c.lower()
        if any(
            x in low
            for x in (
                "образовательной",
                "направлению",
                "направлен",
                "программ",
                "бакалавр",
                "магистр",
                "специалитет",
            )
        ):
            referat_end += 1
            continue
        break
    if referat_end >= personnel_start or referat_start >= personnel_start:
        return None
    return referat_start, referat_end


def _find_topic_before_referat(paras, referat_start: int) -> int | None:
    t = referat_start - 1
    while t >= 0 and not (paras[t].text or "").strip():
        t -= 1
    if t < 0:
        return None
    c = _collapse(paras[t].text or "")
    if _is_topic_caps_line(c):
        return t
    return None


def _rough_title_content_pt(doc, body_start: int) -> float:
    tot = 0.0
    for i in range(body_start):
        p = doc.paragraphs[i]
        t = (p.text or "").strip()
        pf = p.paragraph_format
        if pf.space_before is not None:
            tot += pf.space_before.pt
        if pf.space_after is not None:
            tot += pf.space_after.pt
        if not t:
            continue
        lines = max(1, len(t.split("\n")))
        fs = 14.0
        for r in p.runs:
            if r.font.size:
                fs = r.font.size.pt
                break
        ls = pf.line_spacing
        if ls is None:
            mult = 1.5
        elif isinstance(ls, float):
            mult = max(1.0, float(ls))
        else:
            mult = 1.5
        tot += lines * fs * mult * 1.06
    return tot


def _length_like_to_pt(val) -> float:
    """python-docx: вычитание Length иногда даёт int (EMU) без атрибута .pt."""
    if val is None:
        return 0.0
    if hasattr(val, "pt"):
        return float(val.pt)
    return float(Emu(val).pt)


def _usable_page_inner_pt(doc) -> float:
    sec = doc.sections[0]
    return (
        _length_like_to_pt(sec.page_height)
        - _length_like_to_pt(sec.top_margin)
        - _length_like_to_pt(sec.bottom_margin)
    )


def _text_width_pt(doc) -> float:
    sec = doc.sections[0]
    return (
        _length_like_to_pt(sec.page_width)
        - _length_like_to_pt(sec.left_margin)
        - _length_like_to_pt(sec.right_margin)
    )


def _para_has_page_break(p) -> bool:
    pPr = p._element.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:pageBreakBefore")) is not None:
        return True
    for br in p._element.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def _cover_has_table(doc, fp_end: int) -> bool:
    """True if a ``<w:tbl>`` appears before the last cover paragraph.

    Table-based covers (info blocks in a borderless table) have their own
    geometry that a paragraph-height estimate cannot model, so we leave
    them untouched.
    """
    if fp_end <= 0:
        return False
    end_elem = doc.paragraphs[fp_end - 1]._element
    for child in doc.element.body:
        if child is end_elem:
            return False
        if child.tag == qn("w:tbl"):
            return True
    return False


def _para_font_pt(p) -> float:
    for r in p.runs:
        if r.font.size:
            return float(r.font.size.pt)
    return 14.0


def _para_line_mult(p) -> float:
    ls = p.paragraph_format.line_spacing
    return float(ls) if isinstance(ls, float) else 1.5


def _para_height_pt(p, text_width_pt: float, *, empty_override: float | None = None) -> float:
    """Rough rendered height of a paragraph in points.

    Crucially counts EMPTY paragraphs (one blank line) — the original
    ``_rough_title_content_pt`` skipped them, which is why title covers
    stuffed with blank lines were never detected as overflowing.
    """
    pf = p.paragraph_format
    h = _length_like_to_pt(pf.space_before) + _length_like_to_pt(pf.space_after)
    fs = _para_font_pt(p)
    mult = _para_line_mult(p)
    text = p.text or ""
    if not text.strip():
        m = empty_override if empty_override is not None else mult
        return h + fs * m * 1.15
    lines = 0
    for seg in text.split("\n"):
        seg = seg.strip()
        if not seg:
            lines += 1
            continue
        chars_per_line = max(1, int(text_width_pt / (fs * 0.48)))
        lines += max(1, -(-len(seg) // chars_per_line))
    return h + lines * fs * mult * 1.15


def fit_title_cover_to_page(doc, body_start: int, details: list[str]) -> bool:
    """Compress the title cover so its last line (city/year) stays on page 1.

    Many covers use a stack of empty 1.5-spaced paragraphs as vertical
    spacing; the total exceeds one page and Word pushes «Город, 2026» onto
    a second sheet. We measure the cover height (including blank lines and
    text wrapping) and, only when it overflows, scale the empty-paragraph
    spacing down proportionally so everything fits on a single page. Text
    paragraphs and table-based covers are never modified.
    """
    if body_start < 4:
        return False
    paras = doc.paragraphs

    # First-page region ends at the first hard page break inside the front
    # matter (so multi-page front matter — задание/реферат — is not forced
    # onto one sheet), otherwise at the body start.
    fp_end = body_start
    for i in range(1, body_start):
        if _para_has_page_break(paras[i]):
            fp_end = i
            break
    if fp_end < 4 or _cover_has_table(doc, fp_end):
        return False

    last_ne = -1
    for i in range(fp_end):
        if (paras[i].text or "").strip():
            last_ne = i
    if last_ne < 4:
        return False

    usable = _usable_page_inner_pt(doc)
    tw = _text_width_pt(doc)
    region = range(last_ne + 1)

    est = sum(_para_height_pt(paras[i], tw) for i in region)
    if est <= usable * 1.02:
        return False  # cover already fits — leave the title exactly as is

    text_pt = 0.0
    empty_pt = 0.0
    empties: list = []
    for i in region:
        p = paras[i]
        if (p.text or "").strip():
            text_pt += _para_height_pt(p, tw)
        else:
            empty_pt += _para_height_pt(p, tw)
            empties.append(p)
    if not empties or empty_pt <= 0:
        return False

    target = usable * 0.96
    avail_empty = target - text_pt
    k = max(0.0, avail_empty / empty_pt)
    if k >= 0.999:
        return False  # text alone already needs all the room; nothing to gain

    changed = False
    for p in empties:
        pf = p.paragraph_format
        if pf.space_before is not None and _length_like_to_pt(pf.space_before) != 0:
            pf.space_before = Pt(0)
            changed = True
        if pf.space_after is not None and _length_like_to_pt(pf.space_after) != 0:
            pf.space_after = Pt(0)
            changed = True
        cur_mult = _para_line_mult(p)
        new_mult = round(max(0.6, cur_mult * k), 2)
        if abs(cur_mult - new_mult) > 0.02:
            pf.line_spacing = new_mult
            changed = True

    if changed:
        details.append(
            "Титульный лист: уплотнены пустые интервалы, чтобы город/год "
            "не уходили на вторую страницу"
        )
    return changed


def distribute_title_page_vertical_blocks(
    doc, body_start: int, details: list[str],
) -> bool:
    """Задать space_before на границах блоков, чтобы заполнить страницу и опустить город/год."""
    if body_start < 4:
        return False
    paras = doc.paragraphs
    city_idx = _find_footer_city_idx(paras, body_start)
    if city_idx is None:
        return False
    pr = _find_personnel_range(paras, city_idx)
    if pr is None:
        return False
    personnel_start, personnel_end = pr
    if personnel_end >= city_idx:
        return False
    span = _find_referat_span(paras, personnel_start)
    if span is None:
        return False
    referat_start, referat_end = span
    if referat_end >= personnel_start:
        return False
    topic_idx = _find_topic_before_referat(paras, referat_start)

    boundaries: list[int] = []
    if topic_idx is not None:
        boundaries.append(topic_idx)
    boundaries.extend([referat_start, personnel_start, city_idx])

    usable = _usable_page_inner_pt(doc)
    est = _rough_title_content_pt(doc, body_start)
    reserve_pt = 64.0
    slack = usable - est - reserve_pt
    if slack < 72.0:
        logger.info(
            "title layout: slack %.1f pt too small (usable %.1f, est %.1f)",
            slack,
            usable,
            est,
        )
        return False

    n = len(boundaries)
    if n == 4:
        weights = [1.0, 1.0, 1.0, 2.0]
    elif n == 3:
        weights = [1.0, 1.0, 2.0]
    else:
        return False
    wsum = sum(weights[:n])
    raw_gaps = [slack * weights[i] / wsum for i in range(n)]
    max_gap = 220.0
    gaps = [min(g, max_gap) for g in raw_gaps]

    changed = False
    for idx, gap in zip(boundaries, gaps):
        p = paras[idx]
        pf = p.paragraph_format
        cur = pf.space_before.pt if pf.space_before is not None else 0.0
        if abs(cur - gap) > 0.75:
            pf.space_before = Pt(round(gap, 1))
            changed = True
    if changed:
        details.append(
            "Титульный лист: интервалы между блоками (шапка — тема — реферат — исполнитель — город)"
        )
    return changed
