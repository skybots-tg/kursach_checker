from __future__ import annotations

import logging
import re
import unicodedata
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from lxml import etree

from app.rules_engine.style_resolve import effective_first_line_indent_mm

logger = logging.getLogger(__name__)

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_BINARY_PREFIXES = ("word/media/", "word/embeddings/")

_BULLET_CHARS = frozenset("\u2022\u25cf\u25cb\u25e6\u2023\u2043\u25aa\u25ab\u00b7")
_EM_DASH = "\u2014"
_EN_DASH = "\u2013"
_HYPHEN = "-"
_BLACK = RGBColor(0, 0, 0)

_ENUM_LETTER_RE = re.compile(r"^[а-яёa-z]\)\s", re.IGNORECASE)
_ENUM_DIGIT_PAREN_RE = re.compile(r"^\d{1,2}\)\s")

_CAPTION_RE = re.compile(
    r"^(?:"
    r"\u0420\u0438\u0441\u0443\u043d\u043e\u043a"
    r"|\u0420\u0438\u0441\.?"
    r"|\u0422\u0430\u0431\u043b\u0438\u0446\u0430"
    r")\s+\d+"
)
_THEME_COLOR_ATTRS = (qn("w:themeColor"), qn("w:themeTint"), qn("w:themeShade"))

def _set_color_to_black(color_el) -> bool:
    val = color_el.get(qn("w:val"))
    needs_fix = (val and val.lower() != "000000") or color_el.get(qn("w:themeColor")) is not None
    if not needs_fix:
        return False
    color_el.set(qn("w:val"), "000000")
    for attr in _THEME_COLOR_ATTRS:
        if color_el.get(attr) is not None:
            del color_el.attrib[attr]
    return True

def is_field_code_run(run) -> bool:
    elem = run._element
    return (
        elem.find(qn("w:fldChar")) is not None
        or elem.find(qn("w:instrText")) is not None
    )

def is_manual_list_para(text: str) -> bool:
    stripped = text.lstrip()
    if not stripped:
        return False
    if stripped[0] in _BULLET_CHARS:
        return True
    if stripped[0] == "*" and len(stripped) > 1 and stripped[1] in (" ", "\t", "\xa0"):
        return True
    if stripped[0] in (_HYPHEN, _EN_DASH, _EM_DASH) and len(stripped) > 1:
        if not stripped[1].isdigit():
            return True
    if _ENUM_LETTER_RE.match(stripped) or _ENUM_DIGIT_PAREN_RE.match(stripped):
        return True
    return False

def fix_font_color_styles(doc: Document, details: list[str]) -> bool:
    changed = False
    for style in doc.styles:
        name = (getattr(style, "name", "") or "").lower()
        if not any(k in name for k in (
            "hyperlink", "heading", "toc",
            "\u0437\u0430\u0433\u043e\u043b\u043e\u0432",
        )):
            continue
        try:
            c = style.font.color
            if (c.rgb is not None and c.rgb != _BLACK) or c.theme_color is not None:
                c.rgb = _BLACK
                changed = True
        except (AttributeError, TypeError):
            pass
    if changed:
        details.append("\u0421\u0442\u0438\u043b\u0438: \u0446\u0432\u0435\u0442 \u0448\u0440\u0438\u0444\u0442\u0430 -> \u0447\u0451\u0440\u043d\u044b\u0439")
    return changed

def fix_font_color_runs(paragraph, para_label: str, details: list[str]) -> bool:
    p_elem = paragraph._element
    changed = False
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is not None:
        rPr_def = pPr.find(qn("w:rPr"))
        if rPr_def is not None:
            c_el = rPr_def.find(qn("w:color"))
            if c_el is not None and _set_color_to_black(c_el):
                changed = True
    for r_elem in p_elem.iter(qn("w:r")):
        rPr = r_elem.find(qn("w:rPr"))
        if rPr is None:
            continue
        c_el = rPr.find(qn("w:color"))
        if c_el is not None and _set_color_to_black(c_el):
            changed = True
    if changed:
        details.append(f"{para_label}: \u0446\u0432\u0435\u0442 \u0448\u0440\u0438\u0444\u0442\u0430 -> \u0447\u0451\u0440\u043d\u044b\u0439")
    return changed

def fix_italic_styles(doc: Document, details: list[str]) -> bool:
    changed = False
    for style in doc.styles:
        try:
            if style.font.italic is True:
                style.font.italic = False
                changed = True
        except (AttributeError, TypeError):
            pass
    if changed:
        details.append("\u0421\u0442\u0438\u043b\u0438: \u043a\u0443\u0440\u0441\u0438\u0432 \u0443\u0431\u0440\u0430\u043d")
    return changed

def fix_remove_italic(paragraph, para_label: str, details: list[str]) -> bool:
    p_elem = paragraph._element
    changed = False
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is not None:
        rPr_default = pPr.find(qn("w:rPr"))
        if rPr_default is not None:
            for tag in (qn("w:i"), qn("w:iCs")):
                el = rPr_default.find(tag)
                if el is not None:
                    rPr_default.remove(el)
                    changed = True
    for r_elem in p_elem.iter(qn("w:r")):
        rPr = r_elem.find(qn("w:rPr"))
        if rPr is None:
            continue
        for tag in (qn("w:i"), qn("w:iCs")):
            el = rPr.find(tag)
            if el is not None:
                rPr.remove(el)
                changed = True
    if changed:
        details.append(f"{para_label}: \u043a\u0443\u0440\u0441\u0438\u0432 \u0443\u0431\u0440\u0430\u043d")
    return changed

def fix_list_indent(paragraph, para_label: str, details: list[str]) -> bool:
    pf = paragraph.paragraph_format
    changed = False
    if abs(effective_first_line_indent_mm(paragraph)) > 0.5:
        pf.first_line_indent = Mm(0)
        changed = True
    if pf.left_indent is not None and int(pf.left_indent) > int(Mm(0.5)):
        pf.left_indent = Mm(0)
        changed = True
    pPr_el = paragraph._element.find(qn("w:pPr"))
    if pPr_el is not None:
        ind = pPr_el.find(qn("w:ind"))
        if ind is not None:
            for a in (qn("w:left"), qn("w:start"), qn("w:hanging")):
                v = ind.get(a)
                if v is None:
                    continue
                try:
                    if int(v) > 0:
                        if a == qn("w:hanging"):
                            ind.attrib.pop(a, None)
                        else:
                            ind.set(a, "0")
                        changed = True
                except (ValueError, TypeError):
                    pass
    if changed:
        details.append(f"{para_label}: \u043e\u0442\u0441\u0442\u0443\u043f \u0441\u043f\u0438\u0441\u043a\u0430 \u043e\u0431\u043d\u0443\u043b\u0451\u043d")
    return changed

_ALL_MARKER_CHARS = _BULLET_CHARS | frozenset((_EN_DASH, _EM_DASH, "*"))

def fix_markers_text(
    paragraph, para_label: str, details: list[str],
    marker_char: str = _HYPHEN,
) -> bool:
    if not paragraph.runs:
        return False
    full = paragraph.text.lstrip()
    if not full or full[0] not in _ALL_MARKER_CHARS:
        return False
    for run in paragraph.runs:
        rt = run.text
        if not rt.strip():
            continue
        stripped = rt.lstrip()
        if stripped and stripped[0] in _ALL_MARKER_CHARS:
            ws = rt[: len(rt) - len(stripped)]
            rest = stripped[1:].lstrip()
            new_text = ws + marker_char + " " + rest
            if new_text == rt:
                return False
            run.text = new_text
            details.append(f"{para_label}: \u043c\u0430\u0440\u043a\u0435\u0440 -> {marker_char}")
            return True
        break
    return False

def fix_numbering_bullets(
    doc: Document, body_font: str, details: list[str],
    marker_char: str = _HYPHEN,
) -> bool:
    try:
        npart = doc.part.numbering_part
    except Exception:
        return False
    if npart is None:
        return False
    elem = npart._element
    changed = False
    for lvl in elem.iter(qn("w:lvl")):
        fmt = lvl.find(qn("w:numFmt"))
        if fmt is None or fmt.get(qn("w:val")) != "bullet":
            continue
        lt = lvl.find(qn("w:lvlText"))
        if lt is None:
            continue
        val = lt.get(qn("w:val")) or ""
        if val and val != marker_char:
            lt.set(qn("w:val"), marker_char)
            rPr = lvl.find(qn("w:rPr"))
            if rPr is not None:
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is not None:
                    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                        fq = qn(a)
                        if rFonts.get(fq) is not None:
                            rFonts.set(fq, body_font)
            changed = True
    if changed:
        details.append("\u041d\u0443\u043c\u0435\u0440\u0430\u0446\u0438\u044f: \u043c\u0430\u0440\u043a\u0435\u0440\u044b -> \u0434\u0435\u0444\u0438\u0441")
    return changed

def fix_dashes_in_text(paragraph, para_label: str, details: list[str]) -> bool:
    changed = False
    for run in paragraph.runs:
        if is_field_code_run(run):
            continue
        t = run.text
        if not t:
            continue
        new_t = t
        if _EM_DASH in new_t:
            new_t = new_t.replace(_EM_DASH, _EN_DASH)
        if " - " in new_t:
            new_t = new_t.replace(" - ", f" {_EN_DASH} ")
        if new_t != t:
            run.text = new_t
            changed = True
    if changed:
        details.append(f"{para_label}: \u0434\u0435\u0444\u0438\u0441/\u0442\u0438\u0440\u0435 \u2192 \u0441\u0440\u0435\u0434\u043d\u0435\u0435 \u0442\u0438\u0440\u0435 (\u2013)")
    return changed


# Markdown emphasis markers that frequently leak in when students paste
# answers from ChatGPT / other LLMs. We strip the surrounding markers but
# keep the inner text intact (pure cleanup — we do NOT try to convert the
# emphasis to actual Word bold/italic, because GPT often over-bolds and the
# resulting body text would look messy in academic work).
_MD_BOLD_RE = re.compile(r"\*\*([^*\n]{1,500}?)\*\*")
_MD_BOLD_ALT_RE = re.compile(r"__([^_\n]{1,500}?)__")
_MD_STRIKE_RE = re.compile(r"~~([^~\n]{1,500}?)~~")
# Single ``*foo*`` is risky — a lone ``*`` is also a common bullet glyph
# and footnote marker. We strip it only when:
#   * the opening ``*`` is NOT touching a word/asterisk on its left (so
#     ``Сноска*1`` and ``a*b`` survive), and there is NO whitespace right
#     after it (so ``* пункт`` and ``5 * 3`` survive);
#   * the closing ``*`` is NOT touching a word/asterisk on its right and
#     has NO whitespace right before it.
_MD_ITALIC_STAR_RE = re.compile(r"(?<![\w*])\*(?!\s)([^*\n]{1,300}?)(?<!\s)\*(?![\w*])")


def fix_strip_markdown_artifacts(
    paragraph, para_label: str, details: list[str],
) -> bool:
    """Strip leftover Markdown emphasis markers (``**bold**`` / ``__bold__``
    / ``~~strike~~`` / a conservative ``*italic*``) from each run's text.

    Operates per-run, which covers the dominant copy-from-GPT scenario
    where the entire pasted block lands as a single run with the markers
    inline. The inner text is preserved verbatim; only the surrounding
    delimiters disappear.
    """
    changed = False
    removed_kinds: list[str] = []
    for run in paragraph.runs:
        if is_field_code_run(run):
            continue
        original = run.text
        if not original:
            continue
        if "**" not in original and "__" not in original and "~~" not in original and "*" not in original:
            continue

        new_text = original
        if "**" in new_text:
            replaced = _MD_BOLD_RE.sub(r"\1", new_text)
            if replaced != new_text:
                if "**" not in removed_kinds:
                    removed_kinds.append("**")
                new_text = replaced
        if "__" in new_text:
            replaced = _MD_BOLD_ALT_RE.sub(r"\1", new_text)
            if replaced != new_text:
                if "__" not in removed_kinds:
                    removed_kinds.append("__")
                new_text = replaced
        if "~~" in new_text:
            replaced = _MD_STRIKE_RE.sub(r"\1", new_text)
            if replaced != new_text:
                if "~~" not in removed_kinds:
                    removed_kinds.append("~~")
                new_text = replaced
        if "*" in new_text:
            replaced = _MD_ITALIC_STAR_RE.sub(r"\1", new_text)
            if replaced != new_text:
                if "*" not in removed_kinds:
                    removed_kinds.append("*")
                new_text = replaced

        if new_text != original:
            run.text = new_text
            changed = True

    # Cross-run pass: markers split across run boundaries (e.g. run1="**text",
    # run2="more**") are invisible to per-run regex. Concatenate, clean, redistribute.
    full_text = "".join(r.text or "" for r in paragraph.runs if not is_field_code_run(r))
    if "**" in full_text or "__" in full_text or "~~" in full_text:
        cleaned = full_text
        if "**" in cleaned:
            rep = _MD_BOLD_RE.sub(r"\1", cleaned)
            if rep != cleaned:
                if "**" not in removed_kinds:
                    removed_kinds.append("**")
                cleaned = rep
        if "__" in cleaned:
            rep = _MD_BOLD_ALT_RE.sub(r"\1", cleaned)
            if rep != cleaned:
                if "__" not in removed_kinds:
                    removed_kinds.append("__")
                cleaned = rep
        if "~~" in cleaned:
            rep = _MD_STRIKE_RE.sub(r"\1", cleaned)
            if rep != cleaned:
                if "~~" not in removed_kinds:
                    removed_kinds.append("~~")
                cleaned = rep
        if cleaned != full_text:
            _redistribute_text(paragraph, cleaned)
            changed = True

    if changed:
        kinds = ", ".join(removed_kinds) if removed_kinds else "Markdown"
        details.append(f"{para_label}: убраны Markdown-выделения ({kinds})")
    return changed


def _redistribute_text(paragraph, cleaned_text: str) -> None:
    """Distribute *cleaned_text* back across existing runs, preserving formatting."""
    runs = [r for r in paragraph.runs if not is_field_code_run(r)]
    if not runs:
        return
    pos = 0
    for i, run in enumerate(runs):
        orig_len = len(run.text or "")
        if i < len(runs) - 1:
            run.text = cleaned_text[pos:pos + orig_len] if pos + orig_len <= len(cleaned_text) else cleaned_text[pos:]
            pos += orig_len
        else:
            run.text = cleaned_text[pos:]


def fix_caption_trailing_dot(paragraph, para_label: str, details: list[str]) -> bool:
    text = paragraph.text.strip()
    if not text.endswith(".") or text.endswith(".."):
        return False
    if not _CAPTION_RE.match(text):
        return False
    t_elements = list(paragraph._element.iter(qn("w:t")))
    for t_el in reversed(t_elements):
        s = (t_el.text or "").rstrip()
        if s and s.endswith(".") and not s.endswith(".."):
            tail = (t_el.text or "")[len(s):]
            t_el.text = s[:-1] + tail
            details.append(f"{para_label}: \u0442\u043e\u0447\u043a\u0430 \u0432 \u043a\u043e\u043d\u0446\u0435 \u043f\u043e\u0434\u043f\u0438\u0441\u0438 \u0443\u0431\u0440\u0430\u043d\u0430")
            return True
        if s:
            break
    return False

def postprocess_fixed_docx(original: Path, output: Path) -> None:
    orig_bins: dict[str, tuple[zipfile.ZipInfo, bytes]] = {}
    with zipfile.ZipFile(str(original), "r") as orig_zf:
        for info in orig_zf.infolist():
            if any(info.filename.startswith(p) for p in _BINARY_PREFIXES) and not info.filename.endswith("/"):
                orig_bins[info.filename] = (info, orig_zf.read(info.filename))

    has_toc = _zip_has_toc(output)

    if not orig_bins and not has_toc:
        _validate_docx_zip(output)
        return

    temp_path = output.with_name(output.stem + ".tmp.docx")
    settings_injected = False
    with zipfile.ZipFile(str(output), "r") as out_zf:
        with zipfile.ZipFile(str(temp_path), "w", zipfile.ZIP_DEFLATED) as new_zf:
            for item in out_zf.infolist():
                if item.filename in orig_bins:
                    orig_info, orig_data = orig_bins[item.filename]
                    restored = _clone_zipinfo(orig_info)
                    restored.compress_type = zipfile.ZIP_STORED
                    new_zf.writestr(restored, orig_data)
                elif has_toc and item.filename == "word/settings.xml":
                    data = _inject_update_fields(out_zf.read(item.filename))
                    new_zf.writestr(item, data)
                    settings_injected = True
                else:
                    new_zf.writestr(item, out_zf.read(item.filename))

    if has_toc and not settings_injected:
        logger.debug("Autofix: word/settings.xml not found, cannot inject updateFields")

    _validate_docx_zip(temp_path)
    temp_path.replace(output)

def _zip_has_toc(path: Path) -> bool:
    with zipfile.ZipFile(str(path), "r") as zf:
        if "word/document.xml" not in zf.namelist():
            return False
        data = zf.read("word/document.xml")
        return b"TOC " in data or b"Table of Contents" in data

def _clone_zipinfo(src: zipfile.ZipInfo) -> zipfile.ZipInfo:
    zi = zipfile.ZipInfo(src.filename)
    for a in ("date_time", "compress_type", "comment", "extra", "create_system", "external_attr"):
        setattr(zi, a, getattr(src, a))
    zi.flag_bits = src.flag_bits & 0x800
    return zi

def _inject_update_fields(settings_bytes: bytes) -> bytes:
    root = etree.fromstring(settings_bytes)
    tag = f"{{{_W_NS}}}updateFields"
    el = root.find(tag)
    if el is None:
        el = etree.SubElement(root, tag)
    el.set(f"{{{_W_NS}}}val", "true")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)

def _validate_docx_zip(path: Path) -> None:
    with zipfile.ZipFile(str(path), "r") as zf:
        bad = zf.testzip()
        if bad is not None:
            raise ValueError(f"Corrupted ZIP entry: {bad}")
        names = set(zf.namelist())
        for required in ("[Content_Types].xml", "word/document.xml"):
            if required not in names:
                raise ValueError(f"Missing required entry: {required}")
        etree.fromstring(zf.read("word/document.xml"))
        if "word/settings.xml" in names:
            etree.fromstring(zf.read("word/settings.xml"))

def preflight_margins_safe(doc: Document, target_margins_mm: dict) -> bool:
    lt = Mm(target_margins_mm.get("left", 30)).twips
    rt = Mm(target_margins_mm.get("right", 15)).twips
    mn = None
    for sec in doc.sections:
        try: pw = sec.page_width.twips
        except (AttributeError, TypeError): continue
        content = pw - lt - rt
        if mn is None or content < mn:
            mn = content
    if mn is None or mn < 400:
        return False

    slop = 72
    for table in doc.tables:
        el = table._tbl
        tp = el.tblPr
        if tp is None:
            continue
        tw = tp.find(qn("w:tblW"))
        if tw is None:
            continue
        if tw.get(qn("w:type")) != "dxa":
            continue
        wr = tw.get(qn("w:w"))
        try:
            wv = int(wr) if wr is not None else 0
        except (TypeError, ValueError):
            continue
        if wv > mn + slop:
            return False

    return True

def set_table_width_pct100(tbl) -> None:
    tp = tbl.tblPr
    if tp is None:
        tp = OxmlElement("w:tblPr")
        tbl.insert(0, tp)
    for old in tp.findall(qn("w:tblW")):
        tp.remove(old)
    tw = OxmlElement("w:tblW")
    tw.set(qn("w:w"), "5000")
    tw.set(qn("w:type"), "pct")
    tp.append(tw)

def clamp_overflow_table_widths(doc: Document, details: list[str]) -> bool:
    """Force every table to fit within the page content area.

    Handles three separate overflow shapes:
      1. ``<w:tblW w:type="dxa">`` with a value larger than the content
         width — switch the table to 100 % preferred width.
      2. ``<w:tblW w:type="auto">`` with ``<w:tblGrid>`` columns whose
         total width exceeds the content area. Word would render that
         as a fixed grid sticking out past the right margin, so we
         reset the preferred width to 100 % and scale the grid columns
         proportionally down into the content area.
      3. Negative or positive ``<w:tblInd>`` that shifts the table
         outside of the left margin — force the indent back to zero.
    """
    limit = min_content_width_twips(doc)
    slop = 72
    changed_any = False
    seen: set[int] = set()
    for table in doc.tables:
        el = table._tbl
        tid = id(el)
        if tid in seen:
            continue
        seen.add(tid)
        tp = el.tblPr
        if tp is None:
            continue

        if _reset_negative_tbl_ind(tp):
            changed_any = True

        tw = tp.find(qn("w:tblW"))
        wv = 0
        wtype = None
        if tw is not None:
            wtype = tw.get(qn("w:type"))
            try:
                wv = int(tw.get(qn("w:w")) or 0)
            except (TypeError, ValueError):
                wv = 0

        grid_total = _sum_grid_cols(el)

        overflow_dxa = wtype == "dxa" and wv > limit + slop
        overflow_auto = (
            wtype in (None, "auto")
            and grid_total > 0
            and grid_total > limit + slop
        )

        if overflow_dxa:
            set_table_width_pct100(el)
            changed_any = True
        elif overflow_auto:
            set_table_width_pct100(el)
            if _scale_tbl_grid(el, limit):
                changed_any = True
            else:
                changed_any = True

    if changed_any:
        details.append(
            "Таблицы: ширина приведена к области текста (100%)"
        )
    return changed_any


def _reset_negative_tbl_ind(tbl_pr) -> bool:
    """Zero out ``<w:tblInd>`` when it would push the table past the left
    margin (negative value). Positive indents ≥ 1 cm usually are intentional
    and get left alone.
    """
    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        return False
    try:
        val = int(ind.get(qn("w:w")) or 0)
    except (TypeError, ValueError):
        return False
    if val >= 0:
        return False
    ind.set(qn("w:w"), "0")
    ind.set(qn("w:type"), "dxa")
    return True


def _sum_grid_cols(tbl) -> int:
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        return 0
    total = 0
    for col in grid.findall(qn("w:gridCol")):
        try:
            total += int(col.get(qn("w:w")) or 0)
        except (TypeError, ValueError):
            pass
    return total


def _scale_tbl_grid(tbl, target_twips: int) -> bool:
    """Scale every ``<w:gridCol>`` so the grid fits ``target_twips``.

    Also scales cell-level ``<w:tcW type="dxa">`` widths in every row so
    column widths match. Returns True on changes.
    """
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        return False
    cols = grid.findall(qn("w:gridCol"))
    widths: list[int] = []
    for col in cols:
        try:
            widths.append(int(col.get(qn("w:w")) or 0))
        except (TypeError, ValueError):
            widths.append(0)
    total = sum(widths)
    if total <= 0:
        return False
    factor = target_twips / total
    if factor >= 1.0:
        return False
    changed = False
    new_widths = [max(1, int(w * factor)) for w in widths]
    for col, nw in zip(cols, new_widths):
        if col.get(qn("w:w")) != str(nw):
            col.set(qn("w:w"), str(nw))
            changed = True
    for row in tbl.findall(qn("w:tr")):
        for tc, nw in _iter_cells_with_width(row, new_widths):
            if tc is None:
                continue
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                continue
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                continue
            if tcW.get(qn("w:type")) != "dxa":
                continue
            if tcW.get(qn("w:w")) != str(nw):
                tcW.set(qn("w:w"), str(nw))
                changed = True
    return changed


def _iter_cells_with_width(row, new_widths: list[int]):
    """Yield ``(tc_element, target_width)`` pairs for every ``<w:tc>`` in
    *row*, matching them to the scaled grid columns by occupation order.

    If a cell spans multiple columns (``<w:gridSpan>``), its target width
    is the sum of the spanned column widths.
    """
    tcs = row.findall(qn("w:tc"))
    col_idx = 0
    for tc in tcs:
        tcPr = tc.find(qn("w:tcPr"))
        span = 1
        if tcPr is not None:
            gs = tcPr.find(qn("w:gridSpan"))
            if gs is not None:
                try:
                    span = max(1, int(gs.get(qn("w:val")) or 1))
                except (TypeError, ValueError):
                    span = 1
        if col_idx + span <= len(new_widths):
            width = sum(new_widths[col_idx:col_idx + span])
        else:
            width = 0
        yield tc, width
        col_idx += span

def min_content_width_twips(doc: Document) -> int:
    w = []
    for s in doc.sections:
        try: w.append(max(int(s.page_width.twips - s.left_margin.twips - s.right_margin.twips), 400))
        except (AttributeError, TypeError, ValueError): pass
    return min(w) if w else 8640

def iter_table_cell_paragraphs(doc: Document):
    def _collect(tbl, out):
        for row in tbl.rows:
            for cell in row.cells:
                out.extend(cell.paragraphs)
                for nested in cell.tables:
                    _collect(nested, out)
    all_paras: list = []
    for table in doc.tables:
        _collect(table, all_paras)
    seen: set[int] = set()
    for p in all_paras:
        eid = id(p._element)
        if eid not in seen:
            seen.add(eid)
            yield p


def fix_table_cell_spacing(
    paragraph, target_line_spacing: float, label: str, details: list[str],
) -> bool:
    """Force single (or configured) line spacing on a table-cell paragraph.

    Body text uses 1.5 line spacing by spec, but table cells must stay
    compact — the customer explicitly asked for «одинарный интервал в
    таблице». The same call also wipes any inherited ``space_after`` and
    first-line indent so cell content does not break the row layout.

    We always force ``first_line_indent`` and ``left_indent`` to zero at
    the XML level, not just when already explicitly set, because cell
    paragraphs frequently inherit the 1.25 cm red-line from the default
    body style and keeping it collapses two-column tables into an
    unreadable mess.
    """
    pf = paragraph.paragraph_format
    changed = False
    eff_ls = pf.line_spacing
    if eff_ls is None or abs(float(eff_ls) - target_line_spacing) > 0.05:
        pf.line_spacing = target_line_spacing
        changed = True
    if pf.space_after is None or int(pf.space_after) != 0:
        pf.space_after = Pt(0)
        changed = True
    if pf.space_before is not None and int(pf.space_before) != 0:
        pf.space_before = Pt(0)
        changed = True
    if _force_table_cell_zero_indent(paragraph):
        changed = True
    if changed:
        details.append(f"{label}: интервал {target_line_spacing}, без отступов")
    return changed


def _force_table_cell_zero_indent(paragraph) -> bool:
    """Hard-zero first-line and left indents on the paragraph XML.

    Rewrites ``<w:ind>`` directly so style-inherited indents (``1.25 cm``
    from Normal) cannot seep through and re-add the red line inside a
    table cell. Returns ``True`` when anything changed.
    """
    p_elem = paragraph._element
    pPr = p_elem.find(qn("w:pPr"))
    changed = False
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_elem.insert(0, pPr)
        changed = True
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
        changed = True
    for attr in ("w:firstLine", "w:left", "w:start"):
        if ind.get(qn(attr)) != "0":
            ind.set(qn(attr), "0")
            changed = True
    for attr in ("w:hanging", "w:firstLineChars", "w:leftChars", "w:startChars"):
        if ind.get(qn(attr)) is not None:
            del ind.attrib[qn(attr)]
            changed = True
    pf = paragraph.paragraph_format
    if pf.first_line_indent is not None and int(pf.first_line_indent) != 0:
        pf.first_line_indent = Mm(0)
        changed = True
    if pf.left_indent is not None and int(pf.left_indent) != 0:
        pf.left_indent = Mm(0)
        changed = True
    return changed

def fix_remove_highlight(paragraph, para_label: str, details: list[str]) -> bool:
    _HL, _SHD = qn("w:highlight"), qn("w:shd")
    _SAFE = ("", "auto", "ffffff")
    def _strip(parent):
        c = False
        for tag in (_HL, _SHD):
            el = parent.find(tag)
            if el is not None and (tag != _SHD or (el.get(qn("w:fill")) or "").lower() not in _SAFE):
                parent.remove(el); c = True
        return c
    changed = False
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is not None:
        changed |= _strip(pPr)
    for r in paragraph._element.iter(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is not None:
            changed |= _strip(rPr)
    if changed:
        details.append(f"{para_label}: \u0437\u0430\u043b\u0438\u0432\u043a\u0430/\u0432\u044b\u0434\u0435\u043b\u0435\u043d\u0438\u0435 \u0443\u0431\u0440\u0430\u043d\u044b")
    return changed

def _is_safe_to_strip(ch: str) -> bool:
    """Дополнительный предохранитель: даже если whitelist окажется неполным,
    автоисправление не должно вырезать буквы и иероглифы. Удаляем только
    действительно мусорные символы — управляющие, приватные области,
    суррогаты, символы-форматирования и непечатные пробелы.
    """
    cat = unicodedata.category(ch)
    return cat[0] in ("C", "Z") and ch not in (" ", "\t", "\n")


def fix_remove_strange_chars(paragraph, para_label: str, details: list[str], allowed_re) -> bool:
    changed = False
    for run in paragraph.runs:
        if is_field_code_run(run) or not run.text:
            continue
        cleaned = "".join(
            c for c in run.text
            if allowed_re.match(c) or not _is_safe_to_strip(c)
        )
        if cleaned != run.text:
            run.text = cleaned
            changed = True
    if changed:
        for run in paragraph.runs:
            if is_field_code_run(run) or not run.text:
                continue
            stripped = run.text.lstrip()
            if stripped != run.text:
                run.text = stripped
            break
        details.append(f"{para_label}: \u043f\u043e\u0441\u0442\u043e\u0440\u043e\u043d\u043d\u0438\u0435 \u0441\u0438\u043c\u0432\u043e\u043b\u044b \u0443\u0431\u0440\u0430\u043d\u044b")
    return changed

def remove_manual_page_breaks(paragraph) -> bool:
    changed = False
    for br in list(paragraph._element.iter(qn("w:br"))):
        if br.get(qn("w:type")) == "page":
            br.getparent().remove(br)
            changed = True
    return changed

def fix_page_break_before(paragraph, para_label: str, details: list[str]) -> bool:
    pf = paragraph.paragraph_format
    if pf.page_break_before:
        return False
    pf.page_break_before = True
    details.append(f"{para_label}: \u0440\u0430\u0437\u0440\u044b\u0432 \u0441\u0442\u0440\u0430\u043d\u0438\u0446\u044b \u0434\u043e\u0431\u0430\u0432\u043b\u0435\u043d")
    return True

def _is_removable_empty_para(para) -> bool:
    if (para.text or "").strip():
        return False
    for child in para._element:
        tag = child.tag
        if tag == qn("w:pPr"):
            continue
        if tag in (qn("w:bookmarkStart"), qn("w:bookmarkEnd")):
            continue
        if tag == qn("w:r"):
            for rc in child:
                if rc.tag == qn("w:rPr"):
                    continue
                if rc.tag == qn("w:t") and not (rc.text or "").strip():
                    continue
                return False
        else:
            return False
    return True

def remove_empty_paras_before_page_breaks(doc, details: list[str]) -> bool:
    paragraphs = list(doc.paragraphs)
    body = doc.element.body
    to_remove = []
    seen: set[int] = set()

    for i, para in enumerate(paragraphs):
        pf = para.paragraph_format
        if not pf.page_break_before:
            continue
        j = i - 1
        while j >= 0:
            prev = paragraphs[j]
            if not _is_removable_empty_para(prev):
                break
            eid = id(prev._element)
            if eid not in seen:
                seen.add(eid)
                to_remove.append(prev._element)
            j -= 1

    for elem in to_remove:
        body.remove(elem)

    if to_remove:
        details.append(f"\u0423\u0434\u0430\u043b\u0435\u043d\u043e {len(to_remove)} \u043f\u0443\u0441\u0442\u044b\u0445 \u0430\u0431\u0437\u0430\u0446\u0435\u0432 \u043f\u0435\u0440\u0435\u0434 \u0440\u0430\u0437\u0440\u044b\u0432\u0430\u043c\u0438 \u0441\u0442\u0440\u0430\u043d\u0438\u0446")
    return len(to_remove) > 0

def fix_section_margins(
    section, margins_mm: dict, sec_idx: int, details: list[str],
) -> bool:
    changed = False
    for key in ("left", "right", "top", "bottom"):
        target_mm = margins_mm.get(key)
        if target_mm is None:
            continue
        attr = f"{key}_margin"
        current = getattr(section, attr, None)
        target = Mm(target_mm)
        if current is None or abs(int(current) - int(target)) > int(Mm(0.5)):
            setattr(section, attr, target)
            changed = True
    if changed:
        details.append(f"\u0421\u0435\u043a\u0446\u0438\u044f #{sec_idx + 1}: \u043f\u043e\u043b\u044f \u0438\u0441\u043f\u0440\u0430\u0432\u043b\u0435\u043d\u044b")
    return changed
