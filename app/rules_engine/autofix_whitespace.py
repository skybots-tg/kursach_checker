"""Whitespace and indent normalization for autofix.

Handles leading whitespace stripping, left indent normalization for body text,
collapsing of excessive consecutive empty paragraphs, title page spacing,
and source-line single-spacing.
"""
from __future__ import annotations

import logging
import re

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from app.rules_engine.autofix_title_layout import distribute_title_page_vertical_blocks

logger = logging.getLogger(__name__)

_WS_CHARS = " \t\xa0"


def normalize_doc_defaults_spacing(
    doc, line_spacing: float, space_after_pt: float, details: list[str],
) -> bool:
    """Force ``<w:docDefaults>/<w:pPrDefault>/<w:pPr>/<w:spacing>`` to body values.

    Word's default ``pPrDefault`` from a fresh template typically declares
    ``after="200" line="276"`` (10 pt below every paragraph and ~1.15 line
    height). Empty paragraphs inherit those defaults regardless of any
    paragraph-level normalization that autofix performs only on non-empty
    text. The accumulated 10 pt blocks at section ends produce the visible
    «extra spacing after text» the customer reported. Rewriting the document
    defaults removes the gap once for the whole document and keeps non-empty
    paragraphs unchanged because they still carry their own explicit
    ``<w:spacing>`` overrides set by the rest of autofix.
    """
    from docx.oxml import OxmlElement

    doc_defaults = doc.styles.element.find(qn("w:docDefaults"))
    if doc_defaults is None:
        return False
    pPrDefault = doc_defaults.find(qn("w:pPrDefault"))
    if pPrDefault is None:
        pPrDefault = OxmlElement("w:pPrDefault")
        doc_defaults.append(pPrDefault)
    pPr = pPrDefault.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        pPrDefault.append(pPr)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)

    target_after = str(int(round(space_after_pt * 20)))
    target_line = str(int(round(line_spacing * 240)))

    changed = False
    if spacing.get(qn("w:after")) != target_after:
        spacing.set(qn("w:after"), target_after)
        changed = True
    if spacing.get(qn("w:line")) != target_line:
        spacing.set(qn("w:line"), target_line)
        changed = True
    if spacing.get(qn("w:lineRule")) != "auto":
        spacing.set(qn("w:lineRule"), "auto")
        changed = True
    before_attr = spacing.get(qn("w:before"))
    if before_attr is not None and before_attr != "0":
        spacing.set(qn("w:before"), "0")
        changed = True

    if changed:
        details.append(
            f"Стили: интервал по умолчанию -> {line_spacing}, "
            f"после абзаца {space_after_pt} пт"
        )
    return changed


def fix_strip_leading_whitespace(
    paragraph, para_label: str, details: list[str],
) -> bool:
    """Strip leading spaces, tabs, and non-breaking spaces from paragraph runs."""
    if not paragraph.runs:
        return False

    changed = False
    for run in paragraph.runs:
        if not run.text:
            continue
        stripped = run.text.lstrip(_WS_CHARS)
        if stripped == run.text:
            break
        if stripped:
            run.text = stripped
            changed = True
            break
        run.text = ""
        changed = True

    if changed:
        details.append(f"{para_label}: ведущие пробелы убраны")
    return changed


def fix_normalize_left_indent(
    paragraph, para_label: str, details: list[str],
) -> bool:
    """Reset left indent to 0 for body text paragraphs."""
    pf = paragraph.paragraph_format
    changed = False

    if pf.left_indent is not None and int(pf.left_indent) > int(Mm(0.5)):
        pf.left_indent = Mm(0)
        changed = True

    pPr_el = paragraph._element.find(qn("w:pPr"))
    if pPr_el is not None:
        ind = pPr_el.find(qn("w:ind"))
        if ind is not None:
            for attr in (qn("w:left"), qn("w:start")):
                val = ind.get(attr)
                if val is None:
                    continue
                try:
                    if int(val) > 0:
                        ind.set(attr, "0")
                        changed = True
                except (ValueError, TypeError):
                    pass

    if changed:
        details.append(f"{para_label}: левый отступ обнулён")
    return changed


def _is_removable_empty(para) -> bool:
    if (para.text or "").strip():
        return False
    for child in para._element:
        tag = child.tag
        if tag == qn("w:pPr"):
            continue
        if tag in (qn("w:bookmarkStart"), qn("w:bookmarkEnd")):
            continue
        if tag == qn("w:r"):
            for rc in child:
                if rc.tag == qn("w:rPr"):
                    continue
                if rc.tag == qn("w:t") and not (rc.text or "").strip():
                    continue
                return False
        else:
            return False
    return True


def _is_heading_para(p) -> bool:
    """True if paragraph is a Word heading (style name or outlineLvl)."""
    sname = (getattr(p.style, "name", "") or "").lower()
    sid = (getattr(p.style, "style_id", "") or "")
    if "heading" in sname or "заголов" in sname or sid.startswith("Heading"):
        return True
    pPr = p._element.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:outlineLvl")) is not None:
        return True
    return False


def collapse_excessive_empty_paras(
    doc, max_consecutive: int, details: list[str],
) -> bool:
    """Collapse runs of empty paragraphs.

    Policy:
        * Between two body (non-heading) paragraphs — keep **0** empty
          paragraphs. Random blank lines that students leave between
          their introduction sentences are pure noise and the customer
          flags them as «лишние пробелы».
        * Before a heading (chapter/subsection) — keep up to
          ``max_consecutive`` empty paragraphs so the heading visually
          separates from the preceding text. This also preserves the
          blanks that :func:`ensure_blank_before_subheadings` inserts.
        * Trailing blanks at the very end of the document are removed
          entirely.
    """
    paragraphs = list(doc.paragraphs)
    body = doc.element.body
    to_remove: list = []
    n = len(paragraphs)

    i = 0
    while i < n:
        if not _is_removable_empty(paragraphs[i]):
            i += 1
            continue
        j = i
        while j < n and _is_removable_empty(paragraphs[j]):
            j += 1
        # Blanks span paragraphs[i:j]. Decide how many to keep.
        next_is_heading = j < n and _is_heading_para(paragraphs[j])
        allow = max_consecutive if next_is_heading else 0
        for k in range(i + allow, j):
            to_remove.append(paragraphs[k]._element)
        i = j

    for elem in to_remove:
        try:
            body.remove(elem)
        except ValueError:
            pass

    if to_remove:
        details.append(f"Удалено {len(to_remove)} лишних пустых абзацев")
    return len(to_remove) > 0


_SOURCE_LINE_RE = re.compile(r"^источник", re.IGNORECASE)

# Подвал титула (типовая методичка): город без «г.»; в одну строку — «Город, год»
# с запятой; в две строки — запятая после города не ставится. По центру.
# Шапка и блок РЕФЕРАТ/исполнитель не нормализуем — шаблоны ВУЗов расходятся.
# Интервалы до/после на титуле не трогаем; checks_advanced их не проверяет.
_TITLE_CITY_RE = re.compile(
    r"^г\.\s*[А-ЯЁа-яЁё][А-ЯЁа-яЁё\-\s\w]*\s*,?\s*$",
    re.IGNORECASE | re.UNICODE,
)
_YEAR_ONLY_RE = re.compile(r"^\d{4}\s*$")
_CITY_YEAR_ONE_LINE_RE = re.compile(
    r"^(?:г\.\s*)?([А-ЯЁа-яЁё][А-ЯЁа-яЁё\-\s\w]*)\s*,\s*(\d{4})\s*$",
    re.IGNORECASE | re.UNICODE,
)
_CITY_FOOTER_NOISE_RE = re.compile(
    r"программ|направлен|дисциплин|кафедр|факультет|университет|институт|"
    r"реферат|выполн|проверил|студент|групп|министерство|образован",
    re.IGNORECASE,
)


def _collapse_title_text_for_match(text: str) -> str:
    return re.sub(r"[\n\r\t\xa0]+", " ", (text or "")).strip()


def _replace_paragraph_plain_text(paragraph, new_text: str) -> None:
    runs = paragraph.runs
    if not runs:
        return
    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


def _strip_title_city_g_prefix(collapsed: str) -> str:
    return re.sub(r"^г\.\s*", "", (collapsed or "").strip(), flags=re.IGNORECASE).strip()


def _normalize_city_comma_year_one_line(collapsed: str) -> str | None:
    m = _CITY_YEAR_ONE_LINE_RE.match((collapsed or "").strip())
    if not m:
        return None
    return f"{m.group(1).strip()}, {m.group(2).strip()}"


def _is_probable_footer_city_line(collapsed: str) -> bool:
    if not collapsed or len(collapsed) > 80:
        return False
    if _CITY_FOOTER_NOISE_RE.search(collapsed):
        return False
    if _TITLE_CITY_RE.match(collapsed):
        return True
    stripped = _strip_title_city_g_prefix(collapsed)
    if not stripped or _CITY_FOOTER_NOISE_RE.search(stripped):
        return False
    return bool(
        re.match(r"^[А-ЯЁа-яЁё][А-ЯЁа-яЁё\-\s]{0,78}$", stripped, re.UNICODE),
    )


def _run_is_only_soft_vertical_space(r_elem) -> bool:
    """True if the run only draws blank lines / soft breaks (not page/column)."""
    br_tag = qn("w:br")
    for br in r_elem.findall(br_tag):
        bt = br.get(qn("w:type"))
        if bt in ("page", "column"):
            return False
    texts = "".join((t.text or "") for t in r_elem.findall(qn("w:t")))
    if any(c not in "\n\r \t\xa0" for c in texts):
        return False
    return True


def _strip_leading_soft_vertical_runs_from_paragraph(p_element) -> int:
    """Remove leading w:r that only add line breaks / vertical whitespace."""
    r_tag = qn("w:r")
    removed = 0
    while True:
        stripped = False
        for child in list(p_element.iterchildren()):
            if child.tag != r_tag:
                continue
            if _run_is_only_soft_vertical_space(child):
                p_element.remove(child)
                removed += 1
                stripped = True
                break
        if not stripped:
            break
    return removed


def normalize_title_page_city_footer(doc, body_start: int, details: list[str]) -> bool:
    """Подвал: город без «г.», год; одна строка «Город, год» или две строки; по центру."""
    changed = False
    if body_start <= 0:
        return False
    paras = doc.paragraphs

    footer_idx: int | None = None
    one_line = False
    for i in range(body_start - 1, -1, -1):
        raw = (paras[i].text or "").strip()
        if not raw:
            continue
        collapsed = _collapse_title_text_for_match(paras[i].text or "")
        if _normalize_city_comma_year_one_line(collapsed) is not None:
            footer_idx = i
            one_line = True
            break
        if _YEAR_ONLY_RE.match(raw):
            footer_idx = i
            one_line = False
            break

    if footer_idx is not None and one_line:
        p = paras[footer_idx]
        collapsed = _collapse_title_text_for_match(p.text or "")
        want = _normalize_city_comma_year_one_line(collapsed)
        if want is None:
            return False
        if _strip_leading_soft_vertical_runs_from_paragraph(p._element):
            changed = True
        cur = _collapse_title_text_for_match(p.text or "")
        if cur.replace("\xa0", " ").strip() != want:
            _replace_paragraph_plain_text(p, want)
            changed = True
        if p.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            changed = True
        if changed:
            details.append(
                "Титульный лист: подвал — «город, год» в одну строку без «г.», по центру"
            )
        return changed

    if footer_idx is not None and not one_line:
        yp = paras[footer_idx]
        if _strip_leading_soft_vertical_runs_from_paragraph(yp._element):
            changed = True
        if yp.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
            yp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            changed = True
        city_idx = footer_idx - 1
        if city_idx >= 0:
            cp = paras[city_idx]
            c_collapsed = _collapse_title_text_for_match(cp.text or "")
            if _is_probable_footer_city_line(c_collapsed):
                if _strip_leading_soft_vertical_runs_from_paragraph(cp._element):
                    changed = True
                new_city = _strip_title_city_g_prefix(c_collapsed).rstrip(",").strip()
                if new_city and _collapse_title_text_for_match(cp.text or "") != new_city:
                    _replace_paragraph_plain_text(cp, new_city)
                    changed = True
                if cp.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
                    cp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    changed = True
        if changed:
            details.append(
                "Титульный лист: подвал — город без «г.», год на следующей строке, по центру"
            )
        return changed

    for i in range(body_start):
        p = paras[i]
        collapsed = _collapse_title_text_for_match(p.text or "")
        if not _TITLE_CITY_RE.match(collapsed):
            continue
        if _strip_leading_soft_vertical_runs_from_paragraph(p._element):
            changed = True
        new_city = _strip_title_city_g_prefix(collapsed).rstrip(",").strip()
        if new_city and collapsed != new_city:
            _replace_paragraph_plain_text(p, new_city)
            changed = True
        if p.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            changed = True
        if i + 1 < body_start:
            p2 = paras[i + 1]
            if _YEAR_ONLY_RE.match((p2.text or "").strip()):
                if p2.alignment != WD_PARAGRAPH_ALIGNMENT.CENTER:
                    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    changed = True
        break
    if changed:
        details.append(
            "Титульный лист: подвал — убраны лишние переносы, «г.» перед городом снято, по центру"
        )
    return changed


def strip_paragraph_page_breaks(paragraph) -> bool:
    """Remove page breaks on a paragraph.

    python-docx sets ``paragraph_format.page_break_before = False`` but often
    leaves ``<w:pageBreakBefore/>`` in ``<w:pPr>`` — Word still breaks the
    page. Also strips explicit ``<w:br w:type="page"/>`` in runs.
    """
    changed = False
    pf = paragraph.paragraph_format
    if pf.page_break_before:
        pf.page_break_before = False
        changed = True
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is not None:
        pb = pPr.find(qn("w:pageBreakBefore"))
        if pb is not None:
            pPr.remove(pb)
            changed = True
    r_tag = qn("w:r")
    br_tag = qn("w:br")
    for child in paragraph._element.iterchildren():
        if child.tag != r_tag:
            continue
        for br in list(child.findall(br_tag)):
            if br.get(qn("w:type")) == "page":
                child.remove(br)
                changed = True
    return changed


def normalize_title_page_spacing(doc, body_start: int, details: list[str]) -> bool:
    """Титул: снять разрывы страницы; распределить интервалы между блоками; подвал город/год.

    Межблочные интервалы (шапка — тема — реферат — исполнитель — город) задаются через
    ``space_before`` на первых абзацах блоков — оценка свободного места по полям страницы.
    Остальные интервалы/межстрочный интервал у неграничных абзацев не трогаем.
    """
    if body_start <= 0:
        return False

    loop_changed = False
    for idx in range(body_start):
        p = doc.paragraphs[idx]
        if strip_paragraph_page_breaks(p):
            loop_changed = True
    layout_changed = distribute_title_page_vertical_blocks(doc, body_start, details)
    footer_changed = normalize_title_page_city_footer(doc, body_start, details)
    changed = loop_changed or layout_changed or footer_changed
    if loop_changed:
        details.append("Титульный лист: убраны разрывы страницы внутри блока титула")
    return changed


def normalize_source_line_spacing(doc, details: list[str]) -> bool:
    """Set single (1.0) line spacing on 'Источник:...' paragraphs under tables."""
    changed = False
    count = 0
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if _SOURCE_LINE_RE.match(text):
            pf = p.paragraph_format
            cur_ls = pf.line_spacing
            if cur_ls is None or abs(float(cur_ls) - 1.0) > 0.05:
                pf.line_spacing = 1.0
                changed = True
                count += 1
    if changed:
        details.append(f"«Источник…»: одинарный интервал ({count} шт.)")
    return changed
