from __future__ import annotations

"""
Вспомогательные функции для работы с DOCX:
- загрузка документа из байтов;
- доступ к "сырым" XML для технических проверок;
- упрощённые утилиты для структуры и объёма.

Это не универсальный парсер DOCX, а минимальный слой под задачи rules engine.
"""

from dataclasses import dataclass
from io import BytesIO
from typing import Iterable, Sequence
from zipfile import ZipFile

from docx import Document  # type: ignore[import-untyped]
from docx.document import Document as DocxDocument  # type: ignore[import-untyped]
from docx.text.paragraph import Paragraph  # type: ignore[import-untyped]


@dataclass
class LoadedDocx:
    doc: DocxDocument
    main_xml: str
    settings_xml: str | None
    comments_xml: str | None


def load_docx_from_bytes(file_bytes: bytes) -> LoadedDocx:
    """
    Загружает DOCX и вытаскивает важные XML‑части.
    """
    bio = BytesIO(file_bytes)
    doc = Document(bio)

    with ZipFile(BytesIO(file_bytes)) as zf:
        main_xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
        settings_xml = None
        comments_xml = None
        if "word/settings.xml" in zf.namelist():
            settings_xml = zf.read("word/settings.xml").decode("utf-8", errors="ignore")
        if "word/comments.xml" in zf.namelist():
            comments_xml = zf.read("word/comments.xml").decode("utf-8", errors="ignore")

    return LoadedDocx(
        doc=doc, main_xml=main_xml, settings_xml=settings_xml, comments_xml=comments_xml
    )


def iter_paragraphs(doc: DocxDocument) -> Iterable[Paragraph]:
    for p in doc.paragraphs:
        # Игнорируем полностью пустые параграфы
        if not p.text or not p.text.strip():
            continue
        yield p


def get_layout_margins_mm(doc: DocxDocument) -> dict[str, float]:
    """
    Возвращает поля для первой секции документа в миллиметрах.
    DOCX хранит размеры в твипах (1/20 поинта), 1 pt = 1/72 inch.
    Для целей проверки достаточно приблизительного перевода.
    """

    def twips_to_mm(value: int) -> float:
        # 1 inch = 25.4 mm; 1 inch = 72 pt; 1 pt = 20 twips.
        return (value / 20.0 / 72.0) * 25.4

    section = doc.sections[0]
    return {
        "top": twips_to_mm(section.top_margin),
        "bottom": twips_to_mm(section.bottom_margin),
        "left": twips_to_mm(section.left_margin),
        "right": twips_to_mm(section.right_margin),
    }


def count_chars_between_titles(
    doc: DocxDocument,
    start_titles: Sequence[str],
    stop_titles: Sequence[str],
) -> int:
    """
    Подсчёт символов (с пробелами) между разделами с заданными заголовками.

    Используется для подсчёта объёма: от "Введение" до "Список литературы"/"Приложения".
    """
    start_idx: int | None = None
    stop_idx: int | None = None

    titles_normalized_start = {t.strip().lower() for t in start_titles}
    titles_normalized_stop = {t.strip().lower() for t in stop_titles}

    for idx, p in enumerate(doc.paragraphs):
        title = (p.text or "").strip().lower()
        if not title:
            continue
        if start_idx is None and title in titles_normalized_start:
            start_idx = idx
        if start_idx is not None and title in titles_normalized_stop:
            stop_idx = idx
            break

    if start_idx is None or stop_idx is None or stop_idx <= start_idx:
        return 0

    text_parts: list[str] = []
    for p in doc.paragraphs[start_idx:stop_idx]:
        if not p.text:
            continue
        text_parts.append(p.text)
    full_text = "\n".join(text_parts)
    return len(full_text)


def xml_has_track_changes(main_xml: str) -> bool:
    # Поиск типичных маркеров режима правок.
    markers = ["<w:ins", "<w:del", "w:trackRevisions"]
    return any(m in main_xml for m in markers)


def xml_has_comments(main_xml: str, comments_xml: str | None) -> bool:
    if comments_xml and "<w:comment " in comments_xml:
        return True
    # Также смотрим диапазоны комментариев в основном документе.
    return "<w:commentRangeStart" in main_xml


def xml_has_password_protection(settings_xml: str | None) -> bool:
    if not settings_xml:
        return False
    return "<w:documentProtection" in settings_xml


def xml_has_linked_media(main_xml: str) -> bool:
    # Грубая эвристика: ищем ссылки на внешние объекты/картинки.
    # Для наших целей достаточно обнаружения факта наличия ссылок.
    markers = ["r:link", "o:link", "v:imagedata"]
    return any(m in main_xml for m in markers)




