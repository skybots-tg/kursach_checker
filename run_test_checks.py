"""
Тестовый скрипт: генерирует 3 DOCX-файла и прогоняет через rules engine.
  - good_doc.docx   — почти идеальный документ по ГОСТ
  - bad_doc.docx    — документ с множеством нарушений
  - medium_doc.docx — документ с частичными проблемами
"""
from __future__ import annotations

import asyncio
import json
import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm, Pt

BASE_DIR = Path(__file__).resolve().parent
TEST_DIR = BASE_DIR / ".test_files"
TEST_DIR.mkdir(exist_ok=True)

sys.path.insert(0, str(BASE_DIR))

GOST_RULES = {
    "blocks": [
        {
            "key": "file_intake",
            "title": "Приём файлов",
            "enabled": True,
            "severity": "error",
            "params": {
                "allowed_extensions": ["docx"],
                "max_size_mb": 15,
            },
        },
        {
            "key": "integrity",
            "title": "Техническая чистота",
            "enabled": True,
            "severity": "error",
            "params": {
                "forbid_track_changes": True,
                "forbid_comments": True,
                "forbid_password_protection": True,
            },
        },
        {
            "key": "context_extraction",
            "title": "Контекст работы",
            "enabled": True,
            "severity": "warning",
            "params": {
                "detect_course_year": True,
                "detect_authors_count": True,
                "course_year_regex": r"(?i)\b([2-6])\s*курс\b",
                "authors_labels": ["Автор", "Студент", "Выполнил"],
            },
        },
        {
            "key": "work_formats",
            "title": "Формат работы",
            "enabled": True,
            "severity": "warning",
            "params": {
                "allowed_formats": ["academic"],
            },
        },
        {
            "key": "layout",
            "title": "Поля страницы (ГОСТ 7.32-2017)",
            "enabled": True,
            "severity": "warning",
            "params": {
                "margins_mm": {"top": 20, "bottom": 25, "left": 30, "right": 15},
                "tolerance_mm": 1,
            },
        },
        {
            "key": "typography",
            "title": "Основной текст (ГОСТ 7.32-2017)",
            "enabled": True,
            "severity": "warning",
            "params": {
                "body": {
                    "font": "Times New Roman",
                    "size_pt": 14,
                    "line_spacing": 1.5,
                    "first_line_indent_mm": 12.5,
                },
            },
        },
        {
            "key": "heading_formatting",
            "title": "Заголовки",
            "enabled": True,
            "severity": "warning",
            "params": {
                "font": "Times New Roman",
                "size_pt": 14,
                "require_bold": True,
                "level1_alignment": "CENTER",
            },
        },
        {
            "key": "structure",
            "title": "Обязательные разделы",
            "enabled": True,
            "severity": "error",
            "params": {
                "required_sections_in_order": [
                    {"id": "введение", "titles_any_of": ["введение"]},
                    {"id": "основная_часть", "titles_any_of": ["основная часть"]},
                    {"id": "заключение", "titles_any_of": ["заключение"]},
                    {
                        "id": "список_литературы",
                        "titles_any_of": [
                            "список использованных источников",
                            "список литературы",
                        ],
                    },
                ],
            },
        },
        {
            "key": "volume",
            "title": "Объём",
            "enabled": True,
            "severity": "warning",
            "params": {
                "author_sheet_chars_with_spaces": 40000,
                "min_author_sheets_default": 0.3,
            },
        },
        {
            "key": "bibliography",
            "title": "Источники",
            "enabled": True,
            "severity": "warning",
            "params": {
                "min_total_sources": 10,
                "require_foreign_sources": True,
            },
        },
        {
            "key": "page_numbering",
            "title": "Нумерация страниц",
            "enabled": True,
            "severity": "warning",
            "params": {"require": True},
        },
        {
            "key": "toc",
            "title": "Оглавление",
            "enabled": True,
            "severity": "warning",
            "params": {"require": True},
        },
        {
            "key": "captions",
            "title": "Подписи",
            "enabled": True,
            "severity": "warning",
            "params": {
                "figure_pattern": r"^Рисунок\s+\d+\s*[—–-]\s*\S",
                "table_pattern": r"^Таблица\s+\d+\s*[—–-]\s*\S",
                "check_sequential_numbering": True,
            },
        },
        {
            "key": "objects",
            "title": "Объекты",
            "enabled": True,
            "severity": "warning",
            "params": {"forbid_linked_media": True},
        },
        {
            "key": "autofix",
            "title": "Автоисправления",
            "enabled": True,
            "severity": "advice",
            "params": {
                "normalize_alignment": True,
                "normalize_line_spacing": True,
                "normalize_first_line_indent": True,
                "normalize_spacing_before_after": True,
                "normalize_font": True,
                "normalize_margins": True,
                "normalize_headings": True,
                "space_before_pt": 0,
                "space_after_pt": 0,
            },
        },
    ]
}

FILLER_TEXT = (
    "Данная курсовая работа посвящена исследованию методов анализа данных "
    "в контексте современных информационных технологий. Актуальность темы "
    "определяется ростом объёмов данных и необходимостью разработки "
    "эффективных алгоритмов их обработки. В работе рассматриваются "
    "основные подходы к машинному обучению, включая методы классификации, "
    "кластеризации и регрессии. Особое внимание уделяется нейронным сетям "
    "и их применению для решения прикладных задач."
)

BIB_ITEMS_RU = [
    "[1]. Иванов И.И. Основы информатики. — М.: Наука, 2022. — 350 с.",
    "[2]. Петров П.П. Базы данных. — СПб.: Питер, 2021. — 480 с.",
    "[3]. Сидоров С.С. Алгоритмы и структуры данных. — М.: МЦНМО, 2020. — 512 с.",
    "[4]. Козлов К.К. Программная инженерия. — М.: Юрайт, 2023. — 280 с.",
    "[5]. Михайлов М.М. Нейронные сети. — М.: ДМК Пресс, 2022. — 390 с.",
    "[6]. Николаев Н.Н. Статистический анализ. — М.: Статистика, 2021. — 256 с.",
    "[7]. Орлов О.О. Теория вероятностей. — М.: Наука, 2020. — 310 с.",
    "[8]. Фёдоров Ф.Ф. Линейная алгебра. — М.: Высш. школа, 2019. — 224 с.",
]

BIB_ITEMS_EN = [
    "[9]. Smith J. Introduction to Machine Learning. — Springer, 2022. — 450 p.",
    "[10]. Brown A. Deep Learning Fundamentals. — O'Reilly, 2021. — 380 p.",
    "[11]. Johnson B. Data Science Handbook. — Cambridge, 2023. — 500 p.",
    "[12]. Williams C. Statistical Learning. — MIT Press, 2020. — 420 p.",
]


def _set_margins(doc: Document, top=20, bottom=25, left=30, right=15):
    for section in doc.sections:
        section.top_margin = Mm(top)
        section.bottom_margin = Mm(bottom)
        section.left_margin = Mm(left)
        section.right_margin = Mm(right)


def _add_body_paragraph(doc: Document, text: str, *, font="Times New Roman",
                        size_pt=14, spacing=1.5, indent_mm=12.5):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = spacing
    pf.first_line_indent = Mm(indent_mm)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size_pt)
    return p


def _add_heading(doc: Document, text: str, level=1, *, font="Times New Roman",
                 size_pt=14, bold=True, center=True):
    h = doc.add_heading(text, level=level)
    if center:
        h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in h.runs:
        run.font.name = font
        run.font.size = Pt(size_pt)
        run.bold = bold
    return h


def create_good_document() -> Path:
    """Почти идеальный документ по ГОСТ."""
    doc = Document()
    _set_margins(doc)

    _add_body_paragraph(doc, "КУРСОВАЯ РАБОТА", font="Times New Roman", size_pt=14)
    _add_body_paragraph(doc, "Студент 3 курс, группа ИТ-301")
    _add_body_paragraph(doc, "Выполнил: Иванов Иван Иванович")

    _add_heading(doc, "Введение", level=1)
    for _ in range(5):
        _add_body_paragraph(doc, FILLER_TEXT)

    _add_heading(doc, "Основная часть", level=1)
    _add_heading(doc, "Обзор литературы", level=2, center=False)
    for _ in range(8):
        _add_body_paragraph(doc, FILLER_TEXT)

    _add_heading(doc, "Методология исследования", level=2, center=False)
    for _ in range(6):
        _add_body_paragraph(doc, FILLER_TEXT)

    _add_body_paragraph(doc, "Рисунок 1 — Схема алгоритма обработки данных")
    _add_body_paragraph(doc, "Таблица 1 — Сравнительный анализ методов")

    _add_heading(doc, "Заключение", level=1)
    for _ in range(3):
        _add_body_paragraph(doc, FILLER_TEXT)

    _add_heading(doc, "Список использованных источников", level=1)
    for item in BIB_ITEMS_RU + BIB_ITEMS_EN:
        _add_body_paragraph(doc, item, indent_mm=0)

    path = TEST_DIR / "good_doc.docx"
    doc.save(str(path))
    return path


def create_bad_document() -> Path:
    """Документ со множеством нарушений."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
        section.left_margin = Mm(15)
        section.right_margin = Mm(30)

    p = doc.add_paragraph("Курсовая работа")
    run = p.add_run("")
    run.font.name = "Arial"
    run.font.size = Pt(16)

    h = doc.add_heading("Что-то вроде введения", level=1)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.size = Pt(18)
        run.bold = False

    for _ in range(3):
        p = doc.add_paragraph()
        run = p.add_run(FILLER_TEXT[:150])
        run.font.name = "Arial"
        run.font.size = Pt(12)
        pf = p.paragraph_format
        pf.line_spacing = 1.0
        pf.first_line_indent = Mm(0)

    h2 = doc.add_heading("Выводы", level=1)
    for run in h2.runs:
        run.font.name = "Comic Sans MS"
        run.font.size = Pt(20)

    p = doc.add_paragraph()
    run = p.add_run("Работа завершена.")
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    _add_body_paragraph(doc, "Рис. 3 — Без нумерации с 1")
    _add_body_paragraph(doc, "Таблица 5 — Нумерация не с 1")

    path = TEST_DIR / "bad_doc.docx"
    doc.save(str(path))
    return path


def create_medium_document() -> Path:
    """Документ с частичными нарушениями."""
    doc = Document()
    _set_margins(doc, top=20, bottom=25, left=30, right=15)

    _add_body_paragraph(doc, "Курсовая работа по дисциплине «Информатика»")
    _add_body_paragraph(doc, "Студент 4 курс")
    _add_body_paragraph(doc, "Автор: Петров Пётр Петрович")

    _add_heading(doc, "Введение", level=1)
    for _ in range(4):
        _add_body_paragraph(doc, FILLER_TEXT)

    _add_heading(doc, "Основная часть", level=1)
    for _ in range(3):
        p = doc.add_paragraph()
        run = p.add_run(FILLER_TEXT)
        run.font.name = "Arial"
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.0

    for _ in range(3):
        _add_body_paragraph(doc, FILLER_TEXT)

    _add_heading(doc, "Заключение", level=1)
    _add_body_paragraph(doc, FILLER_TEXT)

    _add_heading(doc, "Список литературы", level=1)
    for item in BIB_ITEMS_RU[:5]:
        _add_body_paragraph(doc, item, indent_mm=0)

    path = TEST_DIR / "medium_doc.docx"
    doc.save(str(path))
    return path


SEVERITY_ICONS = {"error": "❌", "warning": "⚠️", "advice": "💡"}
SEVERITY_COLORS = {"error": "\033[91m", "warning": "\033[93m", "advice": "\033[96m"}
RESET = "\033[0m"


def print_report(file_name: str, report: dict):
    summary = report["summary"]
    findings = report["findings"]
    check_errors = report.get("check_errors", [])
    autofix_path = report.get("output_docx_path")

    print(f"\n{'='*72}")
    print(f"  Файл: {file_name}")
    print(f"{'='*72}")
    print(f"  Ошибок: {summary['errors']}  |  Предупреждений: {summary['warnings']}  "
          f"|  Автоисправлений: {summary['fixed']}  |  Размер: {summary['size']} байт")

    if autofix_path:
        print(f"  Исправленный файл: {autofix_path}")

    if check_errors:
        print(f"\n  Внутренние ошибки проверки: {', '.join(check_errors)}")

    print(f"{'─'*72}")

    for i, f in enumerate(findings, 1):
        sev = f["severity"]
        icon = SEVERITY_ICONS.get(sev, "•")
        color = SEVERITY_COLORS.get(sev, "")
        fixed_tag = " [ИСПРАВЛЕНО]" if f.get("auto_fixed") else ""
        print(f"  {color}{icon} [{sev.upper()}]{RESET} {f['title']}{fixed_tag}")
        print(f"     Ожидается: {f['expected']}")
        print(f"     Найдено:   {f['found']}")
        print(f"     Где:       {f['location']}")
        if f.get("auto_fix_details"):
            details_short = f["auto_fix_details"][:120]
            print(f"     Детали:    {details_short}...")
        print()

    print(f"{'='*72}")
    print(f"  ИТОГО: {len(findings)} замечаний")
    print(f"{'='*72}\n")


async def main():
    print("\n📝 Генерация тестовых DOCX-файлов...\n")

    good_path = create_good_document()
    print(f"  ✅ {good_path.name} — создан (правильный документ)")

    bad_path = create_bad_document()
    print(f"  ✅ {bad_path.name} — создан (документ с ошибками)")

    medium_path = create_medium_document()
    print(f"  ✅ {medium_path.name} — создан (документ со средними проблемами)")

    from app.rules_engine.runner import run_document_checks

    test_files = [
        ("good_doc.docx", str(good_path)),
        ("bad_doc.docx", str(bad_path)),
        ("medium_doc.docx", str(medium_path)),
    ]

    print(f"\n{'#'*72}")
    print(f"  ЗАПУСК ПРОВЕРОК ПО ГОСТ-ШАБЛОНУ ({len(GOST_RULES['blocks'])} блоков правил)")
    print(f"{'#'*72}")

    all_reports = {}
    for name, filepath in test_files:
        report = await run_document_checks(filepath, GOST_RULES)
        all_reports[name] = report
        print_report(name, report)

    print("\n" + "=" * 72)
    print("  СВОДНАЯ ТАБЛИЦА")
    print("=" * 72)
    print(f"  {'Файл':<25} {'Ошибки':>8} {'Предупр.':>10} {'Авто-fix':>10} {'Замечаний':>10}")
    print(f"  {'─'*25} {'─'*8} {'─'*10} {'─'*10} {'─'*10}")
    for name, report in all_reports.items():
        s = report["summary"]
        total = len(report["findings"])
        print(f"  {name:<25} {s['errors']:>8} {s['warnings']:>10} {s['fixed']:>10} {total:>10}")
    print("=" * 72)

    report_path = TEST_DIR / "test_report.json"
    with open(report_path, "w", encoding="utf-8") as fp:
        json.dump(all_reports, fp, ensure_ascii=False, indent=2)
    print(f"\n  JSON-отчёт сохранён: {report_path}\n")


if __name__ == "__main__":
    asyncio.run(main())
