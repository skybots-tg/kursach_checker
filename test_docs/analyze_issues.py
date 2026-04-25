"""Analyze specific issues reported by the customer in both documents."""
from __future__ import annotations
import json
import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))


def analyze(path: str) -> dict:
    doc = Document(path)
    result = {
        "total_paragraphs": len(doc.paragraphs),
        "total_tables": len(doc.tables),
        "paragraphs": [],
    }

    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "")
        style_name = p.style.name if p.style else "?"
        style_id = getattr(p.style, "style_id", "") or ""
        pf = p.paragraph_format
        ls = pf.line_spacing
        sb = pf.space_before
        sa = pf.space_after
        fl = pf.first_line_indent
        align = str(p.alignment) if p.alignment is not None else None

        has_page_break = False
        has_soft_break = False
        has_page_break_before = False
        for r in p.runs:
            for br in r._element.findall(qn("w:br")):
                btype = br.get(qn("w:type"))
                if btype == "page":
                    has_page_break = True
                elif btype is None:
                    has_soft_break = True

        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None:
            pb_before = pPr.find(qn("w:pageBreakBefore"))
            if pb_before is not None:
                val = pb_before.get(qn("w:val"))
                if val is None or val not in ("0", "false"):
                    has_page_break_before = True

        runs_info = []
        for r in p.runs[:6]:
            f = r.font
            runs_info.append({
                "text": r.text[:80] if r.text else "",
                "font": f.name,
                "size_pt": f.size.pt if f.size else None,
                "bold": r.bold,
                "italic": r.italic,
            })

        info = {
            "idx": i,
            "text": text[:250],
            "style_name": style_name,
            "style_id": style_id,
            "alignment": align,
            "line_spacing": float(ls) if ls else None,
            "space_before_pt": float(sb.pt) if sb else None,
            "space_after_pt": float(sa.pt) if sa else None,
            "first_line_indent_pt": float(fl.pt) if fl else None,
            "has_page_break": has_page_break,
            "has_soft_break": has_soft_break,
            "has_page_break_before": has_page_break_before,
            "runs": runs_info,
        }
        result["paragraphs"].append(info)

    # Table info
    result["tables_info"] = []
    for ti, tbl in enumerate(doc.tables):
        rows_text = []
        for ri, row in enumerate(tbl.rows[:4]):
            cells = [c.text[:60].replace("\n", " | ") for c in row.cells]
            rows_text.append(cells)
        # paragraphs after table
        result["tables_info"].append({
            "index": ti,
            "rows": len(tbl.rows),
            "cols": len(tbl.columns),
            "sample_rows": rows_text,
        })

    return result


if __name__ == "__main__":
    path = sys.argv[1]
    out_path = sys.argv[2]
    data = analyze(path)
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"OK: {len(data['paragraphs'])} paragraphs written to {out_path}")
