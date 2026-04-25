"""Quick inspector to print the structure of a DOCX (paragraphs, runs, formatting).

Writes UTF-8 directly to a file (avoids Windows console encoding pain).
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))


def _emu_to_pt(v):
    if v is None:
        return None
    try:
        return float(v) / 12700.0
    except Exception:
        return None


def _dump_paragraphs(doc, out, limit=80, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if i >= start + limit:
            break
        text = (p.text or "").replace("\n", " <BR> ")
        style = p.style.name if p.style else "?"
        pf = p.paragraph_format
        ls = pf.line_spacing
        sb = pf.space_before
        sa = pf.space_after
        fl = pf.first_line_indent
        align = p.alignment
        has_pb = False
        has_break = False
        for r in p.runs:
            for br in r._element.findall(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    has_pb = True
                else:
                    has_break = True
        runs_info = []
        for r in p.runs:
            f = r.font
            runs_info.append(
                f"<{r.text!r} {f.name or ''} {f.size.pt if f.size else ''}pt b={r.bold} i={r.italic}>"
            )
        page_marker = " [PB]" if has_pb else ""
        soft = " [SOFT-BREAK]" if has_break else ""
        out.write(
            f"#{i:03d} [{style}] align={align} ls={ls} sb={sb.pt if sb else None} sa={sa.pt if sa else None} fl={fl.pt if fl else None}{page_marker}{soft}\n"
        )
        out.write(f"     TEXT: {text[:200]!r}\n")
        if len(runs_info) <= 8:
            for ri in runs_info:
                out.write(f"      run: {ri}\n")


def _dump_tables(doc, out):
    for ti, tbl in enumerate(doc.tables):
        out.write(f"--- Table #{ti} ({len(tbl.rows)}x{len(tbl.columns)}) ---\n")
        for ri, row in enumerate(tbl.rows[:3]):
            cells = " | ".join((c.text[:40] or "").replace("\n", "<BR>") for c in row.cells)
            out.write(f"  row{ri}: {cells}\n")


def main():
    if len(sys.argv) < 3:
        print("Usage: inspect_doc.py <docx> <out_txt> [start] [limit]")
        sys.exit(1)
    path = sys.argv[1]
    out_path = sys.argv[2]
    start = int(sys.argv[3]) if len(sys.argv) >= 4 else 0
    limit = int(sys.argv[4]) if len(sys.argv) >= 5 else 1000
    doc = Document(path)

    with open(out_path, "w", encoding="utf-8") as out:
        out.write(f"Total paragraphs: {len(doc.paragraphs)}\n")
        out.write(f"Total tables: {len(doc.tables)}\n")
        _dump_paragraphs(doc, out, start=start, limit=limit)
        out.write("\n")
        _dump_tables(doc, out)


if __name__ == "__main__":
    main()
