"""Check specific issues in fixed documents."""
from __future__ import annotations
import json
import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))


def check_referat4(path: str):
    from docx import Document
    doc = Document(path)
    print("=" * 60)
    print(f"  CHECKING: {path}")
    print("=" * 60)

    # 1) Check for "Оглавление" occurrences
    ogl_count = 0
    ogl_paras = []
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip().lower()
        if text in ("оглавление", "содержание"):
            ogl_count += 1
            ogl_paras.append((i, p.text.strip(), p.style.name))
    print(f"\n1. 'Оглавление'/'Содержание' occurrences: {ogl_count}")
    for idx, txt, style in ogl_paras:
        print(f"   #{idx}: '{txt}' (style={style})")

    # 2) Check for markdown ** artifacts
    md_paras = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text or ""
        if "**" in text:
            md_paras.append((i, text[:100]))
    print(f"\n2. Paragraphs with '**' markdown: {len(md_paras)}")
    for idx, txt in md_paras[:5]:
        print(f"   #{idx}: {txt!r}")

    # 3) Check for single * artifacts
    star_paras = []
    for i, p in enumerate(doc.paragraphs):
        for run in p.runs:
            t = run.text or ""
            if re.search(r'(?<!\w)\*(?!\s)', t) or re.search(r'(?<!\s)\*(?!\w)', t):
                star_paras.append((i, p.text[:80] if p.text else ""))
                break
    print(f"\n3. Paragraphs with suspicious '*': {len(star_paras)}")
    for idx, txt in star_paras[:5]:
        print(f"   #{idx}: {txt!r}")

    # 4) Check Введение - empty paragraphs around it
    vved_idx = None
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip().lower() == "введение":
            vved_idx = i
            break
    if vved_idx is not None:
        print(f"\n4. 'Введение' at #{vved_idx}")
        # Count empty paras before and after
        empty_before = 0
        for j in range(vved_idx - 1, max(0, vved_idx - 10), -1):
            if not (doc.paragraphs[j].text or "").strip():
                empty_before += 1
            else:
                break
        empty_after = 0
        for j in range(vved_idx + 1, min(len(doc.paragraphs), vved_idx + 10)):
            if not (doc.paragraphs[j].text or "").strip():
                empty_after += 1
            else:
                break
        print(f"   Empty paras before: {empty_before}")
        print(f"   Empty paras after: {empty_after}")
        # Show next 3 non-empty paragraphs
        count = 0
        for j in range(vved_idx + 1, min(len(doc.paragraphs), vved_idx + 20)):
            text = (doc.paragraphs[j].text or "").strip()
            if text:
                count += 1
                pf = doc.paragraphs[j].paragraph_format
                print(f"   #{j}: ls={pf.line_spacing} sb={pf.space_before} sa={pf.space_after} fi={pf.first_line_indent}")
                print(f"         '{text[:80]}'")
                if count >= 3:
                    break

    # 5) Check title page - paragraphs 0-21 spacing
    print(f"\n5. Title page paragraphs:")
    for i in range(min(22, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        text = (p.text or "").strip()[:40]
        pf = p.paragraph_format
        sa = pf.space_after
        sb = pf.space_before
        ls = pf.line_spacing
        if text or sa or sb:
            print(f"   #{i}: sa={sa.pt if sa else None} sb={sb.pt if sb else None} ls={ls} '{text}'")


def check_referat_meo(path: str):
    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(path)
    print("\n" + "=" * 60)
    print(f"  CHECKING: {path}")
    print("=" * 60)

    # 1) Title page - check first ~20 paragraphs
    print(f"\n1. Title page paragraphs (spacing):")
    for i in range(min(25, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        text = (p.text or "").strip()[:50]
        pf = p.paragraph_format
        sa = pf.space_after
        sb = pf.space_before
        ls = pf.line_spacing
        has_sb = has_sa = False
        if sa and sa.pt > 0:
            has_sa = True
        if sb and sb.pt > 0:
            has_sb = True
        if text or has_sa or has_sb:
            print(f"   #{i}: sa={sa.pt if sa else '-'} sb={sb.pt if sb else '-'} ls={ls} '{text}'")

    # 2) Check headings - space_before
    print(f"\n2. Headings spacing:")
    for i, p in enumerate(doc.paragraphs):
        sn = (p.style.name or "").lower() if p.style else ""
        if "heading" in sn or "заголов" in sn:
            pf = p.paragraph_format
            text = (p.text or "").strip()[:60]
            print(f"   #{i} [{p.style.name}]: sb={pf.space_before} sa={pf.space_after} ls={pf.line_spacing} '{text}'")

    # 3) Check figure captions - "Рисунок"
    print(f"\n3. Figure captions:")
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if text.startswith("Рисунок") or text.startswith("рисунок"):
            print(f"   #{i}: '{text[:80]}'")

    # 4) Check "Источник" paragraphs
    print(f"\n4. 'Источник' paragraphs:")
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if text.lower().startswith("источник"):
            pf = p.paragraph_format
            print(f"   #{i}: ls={pf.line_spacing} sb={pf.space_before} sa={pf.space_after} '{text[:80]}'")

    # 5) Check bibliography entries for dashes
    print(f"\n5. Bibliography entries (checking dashes):")
    in_bib = False
    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        low = text.lower()
        if low in ("список литературы", "список использованных источников", "список использованной литературы"):
            in_bib = True
            continue
        if in_bib and text:
            # Check for en-dash vs hyphen
            has_endash = "\u2013" in text
            has_emdash = "\u2014" in text
            has_hyphen = " - " in text
            print(f"   #{i}: endash={has_endash} emdash={has_emdash} hyphen={has_hyphen} '{text[:80]}'")
            if i > 180:
                break


if __name__ == "__main__":
    test_dir = Path(__file__).parent
    check_referat4(str(test_dir / "referat_4.fixed.docx"))
    check_referat_meo(str(test_dir / "referat_meo.fixed.docx"))
