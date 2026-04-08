import asyncio
import io
import os
import sys

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from app.rules_engine.runner import run_document_checks
from app.rules_engine.template_schema import DEFAULT_TEMPLATE_BLOCKS


async def main():
    doc_path = r"c:\Users\Dmitry\Downloads\Telegram Desktop\Gosudarstvennoe_byudzhetnoe_professionalnoe_obrazovatelnoe_uchrezhdenie.docx"
    rules = {"blocks": [b.model_dump() for b in DEFAULT_TEMPLATE_BLOCKS]}
    result = await run_document_checks(doc_path, rules)

    s = result["summary"]
    print(f"Errors: {s['errors']}, Warnings: {s['warnings']}")

    categories = {}
    for f in result["findings"]:
        cat = f.get("category", "unknown")
        if cat not in categories:
            categories[cat] = []
        categories[cat].append(f)

    for cat in ("section_breaks", "text_cleanliness"):
        items = categories.get(cat, [])
        print(f"\n--- {cat} ({len(items)}) ---")
        for f in items:
            print(f"  [{f['severity']}] {f['title']} @ {f['location']}: {f['found']}")

    fixed = result.get("output_docx_path")
    if fixed and os.path.exists(fixed):
        print(f"\n=== FIXED FILE CHECK ===")
        doc_fixed = Document(fixed)
        u258c_count = sum(1 for p in doc_fixed.paragraphs if "\u258c" in (p.text or ""))
        print(f"Paragraphs with U+258C in fixed: {u258c_count}")

        hl_count = 0
        from docx.oxml.ns import qn
        for p in doc_fixed.paragraphs:
            for r in p._element.iter(qn("w:r")):
                rPr = r.find(qn("w:rPr"))
                if rPr is not None:
                    hl = rPr.find(qn("w:highlight"))
                    if hl is not None and hl.get(qn("w:val")) not in (None, "none"):
                        hl_count += 1
                        break
        print(f"Paragraphs with highlight in fixed: {hl_count}")


if __name__ == "__main__":
    asyncio.run(main())
